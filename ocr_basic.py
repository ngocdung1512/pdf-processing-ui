"""
Complete PDF processing pipeline: PDF → Single DOCX file
Combines all steps: Detection → VietOCR Batch Prediction → Top-to-Bottom Export
"""
import sys
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_BREAK

# Import from existing scripts
from test_yolo_detect_pdf import (
    pdf_to_images, detect_bboxes, crop_bbox
)
from doclayout_yolo import YOLOv10
import torch
import numpy as np
import cv2
from PIL import Image

try:
    from vietocr.tool.predictor import Predictor
    from vietocr.tool.config import Cfg
    VIETOCR_AVAILABLE = True
except ImportError:
    VIETOCR_AVAILABLE = False


def get_device() -> str:
    """
    Chọn device cho YOLO + VietOCR.
    - Ưu tiên GPU ('cuda' / 'cuda:0') nếu torch.cuda.available().
    - Nếu không có GPU thì mới rơi về 'cpu'.
    """
    if torch.cuda.is_available():
        return "cuda"
    return "cpu"


def sort_bboxes_by_position(bboxes):
    """Sort bboxes by natural reading order: top-to-bottom, left-to-right"""
    if not bboxes:
        return []

    # Calculate median bbox height for tolerance
    bbox_heights = [bbox['y2'] - bbox['y1'] for bbox in bboxes if 'y1' in bbox and 'y2' in bbox]
    median_height = sorted(bbox_heights)[len(bbox_heights) // 2] if bbox_heights else 30

    # Sort primarily by Y, then by X
    sorted_by_y = sorted(bboxes, key=lambda b: (b['y1'], b['center_y'], b['x1']))
    
    strips = []
    current_strip = []
    last_y = None
    
    # Tolerance for clustering into same row
    strict_tolerance = max(median_height * 0.4, 15)

    for bbox in sorted_by_y:
        bbox_y1 = bbox['y1']
        
        if last_y is None:
            current_strip.append(bbox)
            last_y = bbox_y1
        else:
            y_diff = abs(bbox_y1 - last_y)
            if y_diff <= strict_tolerance:
                current_strip.append(bbox)
                last_y = min(last_y, bbox_y1)
            else:
                if current_strip:
                    current_strip.sort(key=lambda b: (b['x1'], b['center_x']))
                    strips.append(current_strip)
                current_strip = [bbox]
                last_y = bbox_y1
    
    if current_strip:
        current_strip.sort(key=lambda b: (b['x1'], b['center_x']))
        strips.append(current_strip)
    
    # Flatten strips
    result = []
    for strip in strips:
        result.extend(strip)
    
    return result


def get_tight_text_boxes(image_np):
    """Find tight horizontal single-line bounding boxes using dynamic projection thresholds."""
    import cv2
    import numpy as np
    
    gray = cv2.cvtColor(image_np, cv2.COLOR_BGR2GRAY)
    _, thresh = cv2.threshold(gray, 200, 255, cv2.THRESH_BINARY_INV)
    
    # 1. Horizontal projection to separate lines
    proj = np.sum(thresh, axis=1)
    if len(proj) == 0:
        return []
        
    min_val = np.min(proj)
    max_val = np.max(proj)
    # Define threshold 5% above the baseline noise level
    cut_threshold = min_val + (max_val - min_val) * 0.05
    
    line_y_spans = []
    in_line = False
    start_y = 0
    for i, val in enumerate(proj):
        if val > cut_threshold and not in_line:
            in_line = True
            start_y = i
        elif val <= cut_threshold and in_line:
            in_line = False
            line_y_spans.append((max(0, start_y - 2), min(image_np.shape[0], i + 2)))
            
    if in_line:
        line_y_spans.append((max(0, start_y - 2), image_np.shape[0]))
        
    # Filter extraordinarily small noise lines
    line_y_spans = [span for span in line_y_spans if (span[1] - span[0]) > 8]
    if not line_y_spans:
        return [(0, 0, image_np.shape[1], image_np.shape[0])]
        
    # 2. Vertical projection to strip immense horizontal padding
    final_boxes = []
    for (ly1, ly2) in line_y_spans:
        line_strip = thresh[ly1:ly2, :]
        v_proj = np.sum(line_strip, axis=0)
        
        v_min = np.min(v_proj)
        v_max = np.max(v_proj)
        # 2% above noise floor to trim white spaces at sides
        v_thresh = v_min + (v_max - v_min) * 0.02
        
        text_cols = np.where(v_proj > v_thresh)[0]
        if len(text_cols) == 0:
            text_cols = np.where(v_proj > 0)[0]
            if len(text_cols) == 0:
                continue
            
        min_x = text_cols[0]
        max_x = text_cols[-1]
        
        # Add 3 pixels padding so we don't clip the edges of characters
        lx = max(0, min_x - 3)
        lw = min(image_np.shape[1], max_x + 3) - lx
        ly = ly1
        lh = ly2 - ly1
        
        final_boxes.append((lx, ly, lw, lh))
        
    return final_boxes


def process_pdf_to_docx(
    pdf_path: str,
    output_docx: str = None,
    model_path: str = "doclayout_yolo_docstructbench_imgsz1024.pt",
    imgsz: int = 1024,
    conf: float = 0.1,
    dpi: int = 300,
    enable_ocr: bool = True,
    max_pages: int = None,
    ocr_weight: str = None
):
    """Complete pipeline: PDF → Single DOCX using VietOCR"""
    
    pdf_path = Path(pdf_path)
    if not pdf_path.exists():
        print(f"Error: PDF file not found: {pdf_path}")
        return None
    
    if output_docx is None:
        output_docx = pdf_path.parent / f"{pdf_path.stem}_reconstructed.docx"
    else:
        output_docx = Path(output_docx)
    
    print("=" * 80)
    print("COMPLETE PDF PROCESSING PIPELINE (VietOCR Batch Mode)")
    print("=" * 80)
    print(f"PDF: {pdf_path.name}")
    print(f"Output: {output_docx.name}")
    print()
    
    # Step 1: Load YOLO model
    print("[Step 1] Loading YOLO model...")
    model = YOLOv10(str(model_path))
    device = get_device()
    print(f"  Using device: {device} (GPU ưu tiên, fallback CPU nếu không có)")
    print("  ✓ YOLO Model loaded")
    
    # Step 2: Load OCR models if enabled
    detector = None
    if enable_ocr:
        if not VIETOCR_AVAILABLE:
            print("  ⚠️  VietOCR not available, skipping OCR step")
            enable_ocr = False
        else:
            print("\n[Step 2] Loading VietOCR model...")
            config = Cfg.load_config_from_name('vgg_transformer')
            config['cnn']['pretrained'] = False
            # Dùng cùng device với YOLO: ưu tiên GPU, nếu không có thì CPU
            config['device'] = f"{device}:0" if device.startswith("cuda") else "cpu"
            if ocr_weight and Path(ocr_weight).exists():
                print(f"  Using local OCR weights: {ocr_weight}")
                config['weights'] = str(ocr_weight)
            else:
                config['weights'] = 'https://vocr.vn/data/vietocr/vgg_transformer.pth'
            detector = Predictor(config)
            print("  ✓ OCR model loaded")
    
    # Step 3: Convert PDF to images
    print(f"\n[Step 3] Converting PDF to images (DPI={dpi})...")
    images = pdf_to_images(str(pdf_path), dpi=dpi)
    total_pages = len(images)
    if max_pages is not None and max_pages > 0:
        images = images[:max_pages]
        print(f"  ✓ Converted {total_pages} pages, processing first {len(images)} pages")
    else:
        print(f"  ✓ Converted {len(images)} pages to images")
    
    # Create Word document early
    print("\n[Step 4] Processing pages & generating Word document...")
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    
    total_bboxes = 0
    
    for page_idx, image in enumerate(images):
        print(f"\n  Processing page {page_idx + 1}/{len(images)}...")
        
        # Detect
        print("    Detecting bboxes...")
        detections = detect_bboxes(model, image, imgsz=imgsz, conf=conf)
        boxes = detections['boxes']
        scores = detections['scores']
        class_names = detections['class_names']
        
        # Remove overlaps > 80%
        print("    Removing overlapping bboxes...")
        filtered_boxes = []
        filtered_scores = []
        filtered_class_names = []
        removed_indices = set()
        
        for i, (box1, score1, class1) in enumerate(zip(boxes, scores, class_names)):
            if i in removed_indices:
                continue
            
            x1_1, y1_1, x2_1, y2_1 = box1
            is_duplicate = False
            
            for j, (box2, score2, class2) in enumerate(zip(boxes, scores, class_names)):
                if i == j or j in removed_indices:
                    continue
                x1_2, y1_2, x2_2, y2_2 = box2
                
                overlap_x = max(0, min(x2_1, x2_2) - max(x1_1, x1_2))
                overlap_y = max(0, min(y2_1, y2_2) - max(y1_1, y1_2))
                overlap_area = overlap_x * overlap_y
                
                area1 = (x2_1 - x1_1) * (y2_1 - y1_1)
                area2 = (x2_2 - x1_2) * (y2_2 - y1_2)
                min_area = min(area1, area2)
                
                if min_area > 0:
                    overlap_ratio = overlap_area / min_area
                    if overlap_ratio > 0.8:
                        if score1 < score2:
                            is_duplicate = True
                            removed_indices.add(i)
                            break
                        else:
                            removed_indices.add(j)
            
            if not is_duplicate:
                filtered_boxes.append(box1)
                filtered_scores.append(score1)
                filtered_class_names.append(class1)
                
        print(f"    ✓ Kept {len(filtered_boxes)} bboxes after removing overlaps")
        
        valid_boxes = []
        all_line_images = []
        box_to_lines = []
        
        for idx, (box, score, class_name) in enumerate(zip(filtered_boxes, filtered_scores, filtered_class_names)):
            x1, y1, x2, y2 = box
            width = x2 - x1
            height = y2 - y1
            
            # Skip likely watermarks
            if class_name == 'abandon':
                is_in_header_zone = y1 < image.shape[0] * 0.20
                if (height > width * 3 or width > height * 3) and not is_in_header_zone:
                    continue
                    
            bbox_image = crop_bbox(image, box, padding=10)
            if bbox_image.shape[0] == 0 or bbox_image.shape[1] == 0:
                continue
                
            line_boxes = get_tight_text_boxes(bbox_image)
            lines_in_this_box = 0
            
            for (lx, ly, lw, lh) in line_boxes:
                # Add minor padding
                lx1 = max(0, lx - 2)
                ly1 = max(0, ly - 4)
                lx2 = min(bbox_image.shape[1], lx + lw + 2)
                ly2 = min(bbox_image.shape[0], ly + lh + 4)
                
                line_img = bbox_image[ly1:ly2, lx1:lx2]
                if line_img.shape[0] == 0 or line_img.shape[1] == 0:
                    continue
                    
                # Resize if incredibly small
                if line_img.shape[0] < 16:
                    line_img = cv2.resize(line_img, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
                    
                pil_image = Image.fromarray(line_img)
                all_line_images.append(pil_image)
                lines_in_this_box += 1
                
            if lines_in_this_box > 0:
                valid_boxes.append((box, score, class_name))
                box_to_lines.append(lines_in_this_box)

        page_bboxes = []
        if enable_ocr and all_line_images:
            print(f"    Running VietOCR in batch mode for {len(all_line_images)} lines across {len(valid_boxes)} bboxes...")
            try:
                # predict_batch with return_prob=True returns a tuple of lists: (texts, probs)
                batch_res = detector.predict_batch(all_line_images, return_prob=True)
                all_texts = list(zip(batch_res[0], batch_res[1]))
            except Exception as e:
                print(f"    ⚠️  Batch prediction failed ({e}), falling back to loop...")
                all_texts = []
                for img in all_line_images:
                    res_tuple = detector.predict(img, return_prob=True)
                    all_texts.append(res_tuple)
            
            curr_idx = 0
            for (box, score, class_name), num_lines in zip(valid_boxes, box_to_lines):
                box_texts = all_texts[curr_idx : curr_idx + num_lines]
                curr_idx += num_lines
                
                # Join text elements with a space
                # Filter out pure hallucinations using the prediction confidence score and noise logic
                filtered_texts = []
                for (text_raw, prob) in box_texts:
                    text_str = text_raw.strip()
                    if not text_str:
                        continue
                        
                    # VietOCR confidence drops significantly on watermarks/noise (usually < 0.65)
                    # Safe text is almost always > 0.75
                    if prob < 0.60:
                        # Too risky, probably reading a watermark or background noise
                        continue
                        
                    text_no_space = text_str.replace(' ', '').replace('.', '').replace(',', '')
                    # Aggressive filter for "0310...10111" purely digit watermarks or VVVVV gibberish
                    if len(text_no_space) > 5 and sum(c.isdigit() for c in text_no_space) / max(1, len(text_no_space)) > 0.7:
                        continue
                    if len(text_str) > 8 and ' ' not in text_str and text_str.isupper() and prob < 0.85:
                        # Typical uppercase English hallucinations (NEWSHOWER, TRANSFITTERING)
                        continue
                        
                    filtered_texts.append(text_str)

                text = " ".join(filtered_texts)
                if not text:
                    continue
                
                x1, y1, x2, y2 = box
                page_bboxes.append({
                    'class': class_name,
                    'confidence': float(score),
                    'text': text,
                    'x1': float(x1),
                    'y1': float(y1),
                    'x2': float(x2),
                    'y2': float(y2),
                    'center_x': (x1 + x2) / 2,
                    'center_y': (y1 + y2) / 2,
                })

        # Sort bboxes
        sorted_bboxes = sort_bboxes_by_position(page_bboxes)
        total_bboxes += len(sorted_bboxes)
        
        # Add to document
        if page_idx > 0:
            doc.add_page_break()
            
        for bbox in sorted_bboxes:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            
            run = p.add_run(bbox['text'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            
            if bbox['class'] in ['title', 'heading1', 'heading2']:
                run.font.bold = True
                run.font.size = Pt(14)
                
    # Save document
    print(f"\n[Step 5] Saving Word document...")
    try:
        doc.save(str(output_docx))
        print(f"  ✓ Saved to: {output_docx}")
    except PermissionError:
        import datetime
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        new_output = output_docx.parent / f"{output_docx.stem}_{timestamp}.docx"
        print(f"  ⚠️  Original file locked, saving as: {new_output.name}")
        doc.save(str(new_output))
        output_docx = new_output
    
    print("\n" + "=" * 80)
    print("PROCESSING COMPLETE")
    print("=" * 80)
    print(f"Output file: {output_docx}")
    print(f"Total pages: {len(images)}")
    print(f"Total bboxes: {total_bboxes}")
    
    return output_docx

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='PDF to DOCX using VietOCR Batch')
    parser.add_argument('pdf_path', type=str, help='Path to PDF file')
    parser.add_argument('--output', type=str, default=None, help='Output DOCX file path')
    parser.add_argument('--model', type=str, default='doclayout_yolo_docstructbench_imgsz1024.pt',
                        help='Path to YOLO model file')
    parser.add_argument('--imgsz', type=int, default=1024, help='Image size for inference')
    parser.add_argument('--conf', type=float, default=0.1, help='Confidence threshold')
    parser.add_argument('--dpi', type=int, default=300, help='DPI for PDF conversion')
    parser.add_argument('--no-ocr', action='store_true', help='Disable OCR')
    parser.add_argument('--max-pages', type=int, default=None, help='Maximum pages for testing')
    parser.add_argument('--ocr-weight', type=str, default='vgg_transformer.pth', help='Path to local VietOCR weight')
    
    args = parser.parse_args()
    
    process_pdf_to_docx(
        pdf_path=args.pdf_path,
        output_docx=args.output,
        model_path=args.model,
        imgsz=args.imgsz,
        conf=args.conf,
        dpi=args.dpi,
        enable_ocr=not args.no_ocr,
        max_pages=args.max_pages,
        ocr_weight=args.ocr_weight
    )

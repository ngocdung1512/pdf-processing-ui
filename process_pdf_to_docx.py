"""
Complete PDF processing pipeline: PDF → Single DOCX file
Combines all steps: Detection → OCR → Reconstruction → Word Export
"""
import sys
import argparse
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from test_yolo_detect_pdf import (
    pdf_to_images, detect_bboxes, 
    crop_bbox
)
from doclayout_yolo import YOLOv10
import torch
import cv2
import numpy as np
from PIL import Image
from transformers import Qwen2_5_VLForConditionalGeneration, AutoProcessor
from qwen_vl_utils import process_vision_info
import markdown
from htmldocx import HtmlToDocx


def pixels_to_cm(pixels: float, dpi: int = 200) -> float:
    """Convert pixels to centimeters"""
    return (pixels / dpi) * 2.54


def determine_alignment_by_position(bbox: dict, page_width: float, page_height: float = None):
    """
    Phân tích vị trí để xác định alignment như con người nhìn vào văn bản
    Dựa vào vị trí X (horizontal) để quyết định căn trái/giữa/phải
    """
    x1 = bbox['x1']
    x2 = bbox['x2']
    center_x = bbox['center_x']
    y1 = bbox['y1']
    
    page_center = page_width / 2
    left_margin_threshold = page_width * 0.25  # 25% từ trái
    right_margin_threshold = page_width * 0.75  # 75% từ trái
    center_tolerance = page_width * 0.15  # 15% tolerance cho center
    
    # Phân tích vị trí X để xác định alignment
    # Nếu bbox bắt đầu gần lề trái và kết thúc trước center -> LEFT
    if x1 < left_margin_threshold and x2 < page_center + center_tolerance:
        return WD_ALIGN_PARAGRAPH.LEFT
    # Nếu bbox kết thúc gần lề phải và bắt đầu sau center -> RIGHT
    elif x2 > right_margin_threshold and x1 > page_center - center_tolerance:
        return WD_ALIGN_PARAGRAPH.RIGHT
    # Nếu center_x gần với page_center -> CENTER
    elif abs(center_x - page_center) < center_tolerance:
        return WD_ALIGN_PARAGRAPH.CENTER
    # Mặc định: LEFT
    else:
        return WD_ALIGN_PARAGRAPH.LEFT


def add_header_table(doc: Document, header: dict):
    """
    Tạo header giống PDF bằng table vô hình (2 hàng x 2 cột):
    Hàng 1: [CHÍNH PHỦ] [CỘNG HÒA... + Độc lập...]
    Hàng 2: [Số: ...]   [Hà Nội, ngày...]
    
    CHỈ tạo header table khi thực sự có các phần tử header được phát hiện.
    Không hard-code fallback values - nếu không có thì không tạo.
    """
    if not header:
        return
    
    # Chỉ tạo header table nếu có ít nhất một trong các phần tử sau:
    # - chinh_phu (hoặc so_ky_hieu)
    # - quoc_hieu (hoặc tieu_ngu)
    # Nếu không có các phần tử này, không tạo header table (để các phần tử được xử lý như content bình thường)
    has_left_column = bool(header.get("chinh_phu") or header.get("so_ky_hieu"))
    has_right_column = bool(header.get("quoc_hieu") or header.get("tieu_ngu"))
    
    if not (has_left_column or has_right_column):
        # Không có phần tử header nào được phát hiện, không tạo header table
        return

    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Cố định độ rộng cột để quốc hiệu không bị xuống dòng
    try:
        table.autofit = False
        col_left_width = Cm(5)
        col_right_width = Cm(11)
        table.columns[0].width = col_left_width
        table.columns[1].width = col_right_width
        for cell in table.columns[0].cells:
            cell.width = col_left_width
        for cell in table.columns[1].cells:
            cell.width = col_right_width
    except Exception:
        # Nếu python-docx thay đổi API thì bỏ qua, vẫn dùng layout mặc định
        pass

    # Xóa border của table để trông như text bình thường
    tbl = table._tbl
    tblPr = tbl.tblPr
    # Lấy hoặc tạo node <w:tblBorders>
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tblBorders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tblBorders.append(el)
        el.set(qn("w:val"), "nil")

    # Hàng 1 - ô trái: CHÍNH PHỦ (chỉ hiển thị nếu có)
    cell_cp = table.cell(0, 0)
    para_cp = cell_cp.paragraphs[0]
    para_cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    chinh_phu_text = (header.get("chinh_phu") or "").strip()
    if chinh_phu_text:
        run_cp = para_cp.add_run(chinh_phu_text)
        run_cp.font.name = "Times New Roman"
        run_cp._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run_cp.font.size = Pt(13)
        run_cp.font.bold = True
        run_cp.font.underline = True
    para_cp.paragraph_format.space_after = Pt(0)

    # Hàng 1 - ô phải: CỘNG HÒA... (dòng 1) + Độc lập... (dòng 2, có gạch chân) trong CÙNG 1 paragraph
    # CHỈ hiển thị nếu thực sự có trong header dict
    cell_quoc = table.cell(0, 1)
    para_quoc = cell_quoc.paragraphs[0]
    para_quoc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quoc_hieu_text = (header.get("quoc_hieu") or "").strip()
    if quoc_hieu_text:
        # Nếu OCR gộp luôn cả "Độc lập - Tự do - Hạnh phúc" vào cùng bbox,
        # tách phần tiêu ngữ ra khỏi quốc hiệu để tránh lặp.
        q_lower = quoc_hieu_text.lower()
        if "độc lập" in q_lower and "tự do" in q_lower:
            idx = q_lower.index("độc lập")
            quoc_hieu_text = quoc_hieu_text[:idx].strip(" -\n\r\t")
        # Nếu vẫn còn nhiều dòng, chỉ lấy dòng đầu tiên
        if "\n" in quoc_hieu_text:
            quoc_hieu_text = quoc_hieu_text.splitlines()[0].strip()
        if quoc_hieu_text:
            run_quoc = para_quoc.add_run(quoc_hieu_text)
            run_quoc.font.name = "Times New Roman"
            run_quoc._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run_quoc.font.size = Pt(13)
            run_quoc.font.bold = True

    # Xuống dòng trong cùng paragraph cho tiêu ngữ (chỉ nếu có quốc hiệu hoặc tiêu ngữ)
    tieu_ngu_text = (header.get("tieu_ngu") or "").strip()
    if tieu_ngu_text:
        if quoc_hieu_text:
            run_quoc.add_break()
        run_tn = para_quoc.add_run(tieu_ngu_text)
        run_tn.font.name = "Times New Roman"
        run_tn._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run_tn.font.size = Pt(13)
        run_tn.font.bold = True
        run_tn.font.underline = True
    para_quoc.paragraph_format.space_after = Pt(0)

    # Hàng 2 - ô trái: Số: ...
    cell_so = table.cell(1, 0)
    para_so = cell_so.paragraphs[0]
    para_so.alignment = WD_ALIGN_PARAGRAPH.LEFT
    so_text = (header.get("so_ky_hieu") or "").strip()
    if so_text:
        run_so = para_so.add_run(so_text)
        run_so.font.name = "Times New Roman"
        run_so._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run_so.font.size = Pt(13)
        run_so.font.bold = False
    para_so.paragraph_format.space_after = Pt(0)

    # Hàng 2 - ô phải: Hà Nội, ngày...
    cell_dia = table.cell(1, 1)
    para_dia = cell_dia.paragraphs[0]
    para_dia.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dia_text = (header.get("dia_danh_ngay") or "").strip()
    if dia_text:
        run_dia = para_dia.add_run(dia_text)
        run_dia.font.name = "Times New Roman"
        run_dia._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run_dia.font.size = Pt(13)
        run_dia.font.italic = True
        run_dia.font.bold = False
    para_dia.paragraph_format.space_after = Pt(0)


def add_text_to_doc(doc: Document, bbox: dict, page_width: float, last_y: float = None, is_first_in_page: bool = False, page_height: float = None, quoc_hieu_info: dict = None, is_same_row: bool = False):
    """Add text to document with proper formatting - phân tích bố cục như con người"""
    text = bbox['text'].strip()
    if not text:
        return bbox['y2']
    
    bbox_class = bbox['class']
    x1 = bbox['x1']
    y1 = bbox['y1']
    y2 = bbox['y2']
    center_x = bbox['center_x']
    
    # Phân tích vị trí để xác định alignment (như con người nhìn)
    if page_height is None:
        page_height = 3000  # Default estimate
    
    # Xác định alignment dựa trên vị trí
    alignment = determine_alignment_by_position(bbox, page_width, page_height)
    
    # Detect header elements for first page (kết hợp với vị trí)
    text_lower = text.lower().strip()
    is_in_header_zone = y1 < page_height * 0.20  # Top 20% của trang (increased for better header detection)
    is_chinh_phu = 'chính phủ' in text_lower and len(text_lower) < 20
    is_so_ky_hieu = text_lower.startswith('số:') or text_lower.startswith('số ') or ('số:' in text_lower and '/' in text_lower and 'nđ-cp' in text_lower)
    is_quoc_hieu = 'cộng hòa' in text_lower and 'việt nam' in text_lower
    # Tiêu ngữ: bắt rộng hơn để tránh sót (độc / tự do / hạnh phúc)
    is_tieu_ngu = (
        ('độc lập' in text_lower and 'tự do' in text_lower)
        or ('độc' in text_lower and 'hạnh phúc' in text_lower)
    )
    is_dia_danh_ngay = ('hà nội' in text_lower or 'hà nội' in text) and 'ngày' in text_lower and 'tháng' in text_lower
    is_nghi_dinh = text_lower == 'nghị định' or (text_lower.startswith('nghị định') and len(text_lower) < 15)
    is_can_cu = text_lower.startswith('căn cứ') or text_lower.startswith('căn cứ luật') or text_lower.startswith('căn cứ nghị quyết')
    is_theo_de_nghi = text_lower.startswith('theo đề nghị')
    
    # Tiêu ngữ "Độc lập - Tự do - Hạnh phúc" đã được dựng bằng header table,
    # nên bỏ qua mọi bbox tiêu ngữ để tránh lặp lại lần nữa.
    if is_tieu_ngu:
        return last_y if last_y is not None else y2
    
    # Create paragraph
    para = doc.add_paragraph()
    
    # Override alignment nếu là header elements (ưu tiên vị trí thực tế)
    if is_chinh_phu or is_so_ky_hieu:
        # "CHÍNH PHỦ" và "Số: ...": luôn căn trái (header zone left column)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif is_tieu_ngu:
        # "Độc lập - Tự do - Hạnh phúc": căn giữa dưới "CỘNG HÒA..." trong block bên phải
        if quoc_hieu_info:
            # Tính toán left và right indent để tạo container căn giữa dưới "CỘNG HÒA..."
            quoc_hieu_x1 = quoc_hieu_info['x1']
            quoc_hieu_x2 = quoc_hieu_info['x2']
            
            # Convert pixels to cm
            # Page có left margin 3.0cm, right margin 2.0cm
            # Content area: từ 3.0cm đến (page_width_cm - 2.0cm)
            page_width_cm = pixels_to_cm(page_width)
            quoc_hieu_x1_cm = pixels_to_cm(quoc_hieu_x1)
            quoc_hieu_x2_cm = pixels_to_cm(quoc_hieu_x2)
            
            # Tính left indent: từ lề trái content (3.0cm) đến đầu "CỘNG HÒA..."
            left_indent = max(0, quoc_hieu_x1_cm - 3.0)
            # Tính right indent: từ cuối "CỘNG HÒA..." đến lề phải content
            right_indent = max(0, (page_width_cm - quoc_hieu_x2_cm) - 2.0)
            
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if left_indent > 0:
                para.paragraph_format.left_indent = Cm(left_indent)
            if right_indent > 0:
                para.paragraph_format.right_indent = Cm(right_indent)
        else:
            # Fallback: center-aligned nếu không có thông tin "CỘNG HÒA..."
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif is_quoc_hieu or is_dia_danh_ngay:
        # "CỘNG HÒA..." và "Hà Nội, ngày...": luôn căn phải (header zone right column)
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif is_nghi_dinh:
        # "NGHỊ ĐỊNH": luôn căn giữa
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif is_in_header_zone:
        # Trong header zone nhưng không phải các element đặc biệt, dùng alignment từ vị trí
        para.alignment = alignment
    elif bbox_class in ['title']:
        # Tiêu đề: căn giữa
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # Nội dung chính: dùng alignment từ vị trí
        para.alignment = alignment
    
    # Set left indent (không áp dụng cho header elements và centered text)
    if para.alignment == WD_ALIGN_PARAGRAPH.LEFT and not is_in_header_zone:
        left_margin_px = x1
        left_margin_cm = pixels_to_cm(left_margin_px)
        relative_indent = max(0, left_margin_cm - 3.0)
        if relative_indent > 0:
            para.paragraph_format.left_indent = Cm(relative_indent)
    
    # Format based on class
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    
    # Set line spacing to 1.0 for all paragraphs to avoid default spacing
    para.paragraph_format.line_spacing = 1.0
    
    # Format based on class and content type
    # IMPORTANT: Check header elements FIRST before checking bbox_class
    # This ensures header elements get correct formatting even if YOLO detects them as 'title'
    if is_chinh_phu:
        # "CHÍNH PHỦ": bold, size 13, underlined
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.underline = True
        para.paragraph_format.space_after = Pt(3)
    elif is_quoc_hieu:
        # Quốc hiệu: bold, size 13
        run.font.size = Pt(13)
        run.font.bold = True
        para.paragraph_format.space_after = Pt(6)
    elif is_tieu_ngu:
        # Tiêu ngữ: bold, size 13, underlined
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.underline = True
        para.paragraph_format.space_after = Pt(6)
    elif is_nghi_dinh:
        # "NGHỊ ĐỊNH": bold, size 18, centered, underlined
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.underline = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(12)
    elif bbox_class in ['title'] and not is_in_header_zone:
        # Only apply title formatting if NOT in header zone (header elements already handled above)
        run.font.size = Pt(18)
        run.font.bold = True
        para.paragraph_format.space_after = Pt(12)
    elif is_so_ky_hieu:
        # "Số: ...": không bold, size 13
        run.font.size = Pt(13)
        run.font.bold = False
        para.paragraph_format.space_after = Pt(3)
    elif is_dia_danh_ngay:
        # "Hà Nội, ngày...": không bold, size 13, italic
        run.font.size = Pt(13)
        run.font.bold = False
        run.font.italic = True
        para.paragraph_format.space_after = Pt(3)
    elif is_can_cu or is_theo_de_nghi:
        # "Căn cứ..." và "Theo đề nghị...": italic, size 13
        run.font.size = Pt(13)
        run.font.bold = False
        run.font.italic = True
        para.paragraph_format.space_after = Pt(0)
    elif bbox_class in ['heading1', 'heading2']:
        run.font.size = Pt(14)
        run.font.bold = True
        para.paragraph_format.space_after = Pt(6)
    else:
        run.font.size = Pt(13)  # Changed from 12 to 13
        run.font.bold = False
        para.paragraph_format.space_after = Pt(0)
    
    # Set spacing - IMPORTANT: first paragraph of new page should have NO space_before
    # Also, elements on the same row should have NO space_before
    if is_first_in_page or is_same_row:
        para.paragraph_format.space_before = Pt(0)
    elif last_y is not None:
        vertical_gap = y1 - last_y
        gap_cm = pixels_to_cm(vertical_gap)
        
        # If gap is very small (< 0.1cm), consider it same row
        if gap_cm < 0.1:
            para.paragraph_format.space_before = Pt(0)
        elif gap_cm > 0.5:
            para.paragraph_format.space_before = Pt(12)
        elif gap_cm > 0.2:
            para.paragraph_format.space_before = Pt(6)
        elif gap_cm > 0.1:
            para.paragraph_format.space_before = Pt(3)
        else:
            para.paragraph_format.space_before = Pt(0)
    else:
        # If last_y is None and not first in page, set to 0
        para.paragraph_format.space_before = Pt(0)
    
    return y2


def sort_bboxes_by_position(bboxes, page_width=None):
    """Sort bboxes by reading order - improved header handling for multi-column layout"""
    if not bboxes:
        return []
    
    if page_width is None:
        max_x = max(bbox['x2'] for bbox in bboxes if bbox.get('x2'))
        page_width = max_x if max_x > 0 else 2000
    
    max_y = max(bbox['y2'] for bbox in bboxes if bbox.get('y2'))
    page_height = max_y if max_y > 0 else 3000
    
    # Calculate median bbox height for better tolerance calculation
    bbox_heights = [bbox['y2'] - bbox['y1'] for bbox in bboxes if 'y1' in bbox and 'y2' in bbox]
    if bbox_heights:
        median_height = sorted(bbox_heights)[len(bbox_heights) // 2]
    else:
        median_height = 30  # Default fallback
    
    # Identify header zone (top 20% of page)
    header_zone_threshold = page_height * 0.20
    
    # Separate header bboxes from body bboxes
    header_bboxes = [b for b in bboxes if b['y1'] < header_zone_threshold]
    body_bboxes = [b for b in bboxes if b['y1'] >= header_zone_threshold]
    
    # For header zone: use smaller tolerance based on median height
    header_y_tolerance = max(median_height * 0.5, 20)  # Half of median height, min 20px
    # For body: use median height as tolerance (same row if within one line height)
    body_y_tolerance = max(median_height * 0.8, 25)  # 80% of median height, min 25px
    
    def sort_header_bboxes(header_bboxes):
        """Sort header bboxes: process left column first, then right column, then center"""
        if not header_bboxes:
            return []
        
        # First, sort by Y (row)
        sorted_by_y = sorted(header_bboxes, key=lambda b: (b['y1'], b['center_y']))
        
        # Group into strips (same row)
        strips = []
        current_strip = []
        last_y = None
        
        for bbox in sorted_by_y:
            if last_y is None or abs(bbox['y1'] - last_y) <= header_y_tolerance:
                current_strip.append(bbox)
            else:
                if current_strip:
                    def header_sort_key(b):
                        center_x = b['center_x']
                        left_threshold = page_width * 0.40
                        right_threshold = page_width * 0.60
                        
                        if center_x < left_threshold:
                            zone = 0
                        elif center_x > right_threshold:
                            zone = 2
                        else:
                            zone = 1
                        
                        return (zone, b['x1'], b['center_x'])
                    
                    current_strip.sort(key=header_sort_key)
                    strips.append(current_strip)
                current_strip = [bbox]
            last_y = bbox['y1']
        
        # Handle last strip
        if current_strip:
            def header_sort_key(b):
                center_x = b['center_x']
                left_threshold = page_width * 0.40
                right_threshold = page_width * 0.60
                
                if center_x < left_threshold:
                    zone = 0
                elif center_x > right_threshold:
                    zone = 2
                else:
                    zone = 1
                
                return (zone, b['x1'], b['center_x'])
            
            current_strip.sort(key=header_sort_key)
            strips.append(current_strip)
        
        # Flatten strips
        result = []
        for strip in strips:
            result.extend(strip)
        
        return result
    
    def sort_body_bboxes(body_bboxes):
        """
        Sort body bboxes: simple top-to-bottom, left-to-right (natural reading order)
        No priority system - just sort by Y position, then X position within each row
        """
        if not body_bboxes:
            return []
        
        # Step 1: Sort all bboxes by Y position (top to bottom), then X (left to right)
        # This ensures natural reading order: top-to-bottom, left-to-right
        sorted_by_y = sorted(body_bboxes, key=lambda b: (b['y1'], b['center_y'], b['x1']))
        
        # Step 2: Group into strips (same row) based on Y tolerance
        # Use stricter tolerance to avoid grouping different rows together
        strips = []
        current_strip = []
        last_y = None
        
        for bbox in sorted_by_y:
            bbox_y1 = bbox['y1']
            bbox_height = bbox.get('y2', bbox_y1) - bbox_y1
            
            if last_y is None:
                # First bbox
                current_strip.append(bbox)
                last_y = bbox_y1
            else:
                # Calculate Y difference
                y_diff = abs(bbox_y1 - last_y)
                
                # Use MUCH stricter tolerance: max(median_height * 0.3, 10px)
                # Chỉ nhóm vào cùng dòng nếu Y difference rất nhỏ (< 30% của median height)
                # Điều này đảm bảo các bboxes ở các dòng khác nhau không bị nhóm lại
                strict_tolerance = max(median_height * 0.3, 10)
                
                # Chỉ nhóm vào cùng dòng nếu:
                # 1. Y difference <= strict_tolerance (rất nhỏ)
                # 2. Y difference < bbox_height (nhỏ hơn chiều cao của bbox)
                if y_diff <= strict_tolerance and y_diff < bbox_height:
                    # Same row (within very strict tolerance) - add to current strip
                    current_strip.append(bbox)
                    # Keep the topmost Y for the strip
                    last_y = min(last_y, bbox_y1)
                else:
                    # New row - finalize current strip
                    if current_strip:
                        # Sort within strip: left to right
                        current_strip.sort(key=lambda b: (b['x1'], b['center_x']))
                        strips.append(current_strip)
                    # Start new strip
                    current_strip = [bbox]
                    last_y = bbox_y1
        
        # Handle last strip
        if current_strip:
            current_strip.sort(key=lambda b: (b['x1'], b['center_x']))
            strips.append(current_strip)
        
        # Step 3: Flatten strips (already sorted top-to-bottom, left-to-right within each strip)
        result = []
        for strip in strips:
            result.extend(strip)
        
        return result
    
    # Sort header and body separately
    sorted_header = sort_header_bboxes(header_bboxes)
    sorted_body = sort_body_bboxes(body_bboxes)
    
    # Combine: header first, then body
    return sorted_header + sorted_body


def process_pdf_to_docx(
    pdf_path: str,
    output_docx: str = None,
    model_path: str = None,
    ocr_model_path: str = None,
    imgsz: int = 1024,
    conf: float = 0.1,
    dpi: int = 300,
    enable_ocr: bool = False,
    load_4bit: bool = False,
    load_8bit: bool = False,
    max_pages: int = None
):
    """Complete pipeline: PDF → Single DOCX"""

    # Resolve script directory so model paths work regardless of cwd (e.g. when run from api/scripts)
    _root = Path(__file__).resolve().parent
    if model_path is None:
        model_path = str(_root / "doclayout_yolo_docstructbench_imgsz1024.pt")
    else:
        model_path = str(Path(model_path).resolve())
    if ocr_model_path is None:
        ocr_model_path = str(_root / "Qwen2.5-VL-3B")
    else:
        # Resolve if it looks like a local path (directory exists)
        p = Path(ocr_model_path)
        if p.exists():
            ocr_model_path = str(p.resolve())
        # else: keep as-is for HuggingFace model IDs (e.g. "Qwen/Qwen2.5-VL-3B")

    pdf_path = Path(pdf_path).resolve()
    if not pdf_path.exists():
        print(f"Error: PDF file not found: {pdf_path}")
        return None

    if output_docx is None:
        output_docx = pdf_path.parent / f"{pdf_path.stem}_reconstructed.docx"
    else:
        output_docx = Path(output_docx).resolve()
    
    print("=" * 80)
    print("COMPLETE PDF PROCESSING PIPELINE")
    print("=" * 80)
    print(f"PDF: {pdf_path.name}")
    print(f"Output: {output_docx.name}")
    print()
    
    # Step 1: Load YOLO model
    print("[Step 1] Loading YOLO model...")
    model = YOLOv10(str(model_path))
    device = 'cuda' if torch.cuda.is_available() else 'cpu'
    print(f"  Using device: {device}")
    print("  ✓ Model loaded")
    
    # Step 2: Load OCR models if enabled
    processor = None
    qwen_model = None
    html_parser = None
    if enable_ocr:
        print(f"\n[Step 2] Loading Qwen2.5-VL OCR model from '{ocr_model_path}'...")
        try:
            # Cấu hình quantization để giảm VRAM
            model_kwargs = {
                "device_map": "auto",
            }
            
            if load_4bit:
                from transformers import BitsAndBytesConfig
                quantization_config = BitsAndBytesConfig(
                    load_in_4bit=True,
                    bnb_4bit_compute_dtype=torch.float16,
                    bnb_4bit_quant_type="nf4",
                    bnb_4bit_use_double_quant=True,
                )
                model_kwargs["quantization_config"] = quantization_config
                print("  Using 4-bit quantization (BitsAndBytes NF4)")
            elif load_8bit:
                from transformers import BitsAndBytesConfig
                quantization_config = BitsAndBytesConfig(
                    load_in_8bit=True,
                )
                model_kwargs["quantization_config"] = quantization_config
                print("  Using 8-bit quantization (BitsAndBytes)")
            else:
                model_kwargs["torch_dtype"] = torch.float16
                print("  Using float16 (no quantization)")
            
            qwen_model = Qwen2_5_VLForConditionalGeneration.from_pretrained(
                ocr_model_path, **model_kwargs
            )
            processor = AutoProcessor.from_pretrained(ocr_model_path)
            processor.tokenizer.padding_side = 'left'
            html_parser = HtmlToDocx()
            print("  ✓ Qwen2.5-VL OCR model loaded")
        except Exception as e:
            print(f"  ⚠️  Failed to load Qwen2.5-VL OCR model, skipping OCR step. Error: {e}")
            enable_ocr = False
    
    # Step 3: Convert PDF to images
    print(f"\n[Step 3] Converting PDF to images (DPI={dpi})...")
    try:
        images = pdf_to_images(str(pdf_path), dpi=dpi)
    except Exception as e:
        print(f"  ⚠️  Failed to convert PDF to images: {e}")
        return None
    total_pages = len(images)
    if max_pages is not None and max_pages > 0:
        images = images[:max_pages]
        print(f"  ✓ Converted {total_pages} pages, processing first {len(images)} pages")
    else:
        print(f"  ✓ Converted {len(images)} pages to images")
    
    # Step 4: Create Word document
    print("\n[Step 4] Creating Word document...")
    doc = Document()
    
    # Set page margins (làm cho trang "sát" hơn so với PDF gốc)
    sections = doc.sections
    for section in sections:
        # Giảm lề trên/dưới để khoảng trắng giữa các trang nhỏ lại
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
    
    # Process each page
    print("\n[Step 5] Processing pages...")
    all_pages_bboxes = []
    
    for page_idx, image in enumerate(images):
        print(f"\n  Processing page {page_idx + 1}/{len(images)}...")
        
        # ========================================================================
        # BƯỚC 1: YOLO DETECT BBOXES (không có thứ tự)
        # ========================================================================
        # YOLO detect tất cả bboxes cùng lúc (parallel detection)
        # Kết quả KHÔNG có thứ tự - cần sort lại sau
        # ========================================================================
        print("    Detecting bboxes...")
        detections = detect_bboxes(model, image, imgsz=imgsz, conf=conf)
        boxes = detections['boxes']
        scores = detections['scores']
        class_names = detections['class_names']
        class_ids = detections['class_ids']
        
        print(f"    Found {len(boxes)} bboxes (unsorted - will be sorted later)")
        
        # ========================================================================
        # BƯỚC 1.5: LOẠI BỎ CÁC BBOXES CHỒNG LÊN NHAU (OVERLAPPING)
        # ========================================================================
        # YOLO có thể detect cùng một vùng text với nhiều class khác nhau
        # (ví dụ: "title" và "plain text" cho cùng một vùng)
        # Loại bỏ các bboxes có overlap > 80%, giữ lại bbox có confidence cao hơn
        # ========================================================================
        print("    Removing overlapping bboxes (overlap > 80%)...")
        filtered_boxes = []
        filtered_scores = []
        filtered_class_names = []
        removed_indices = set()
        
        for i, (box1, score1, class1) in enumerate(zip(boxes, scores, class_names)):
            if i in removed_indices:
                continue
            
            x1_1, y1_1, x2_1, y2_1 = box1
            is_duplicate = False
            
            for j, (box2, score2, class2) in enumerate(zip(boxes, scores, class_names)):
                if i == j or j in removed_indices:
                    continue
                
                x1_2, y1_2, x2_2, y2_2 = box2
                
                # Tính overlap
                overlap_x = max(0, min(x2_1, x2_2) - max(x1_1, x1_2))
                overlap_y = max(0, min(y2_1, y2_2) - max(y1_1, y1_2))
                overlap_area = overlap_x * overlap_y
                
                area1 = (x2_1 - x1_1) * (y2_1 - y1_1)
                area2 = (x2_2 - x1_2) * (y2_2 - y1_2)
                min_area = min(area1, area2)
                
                if min_area > 0:
                    overlap_ratio = overlap_area / min_area
                    
                    # Nếu overlap > 80%, loại bỏ bbox có confidence thấp hơn
                    if overlap_ratio > 0.8:
                        if score1 < score2:
                            is_duplicate = True
                            removed_indices.add(i)
                            break
                        else:
                            removed_indices.add(j)
            
            if not is_duplicate:
                filtered_boxes.append(box1)
                filtered_scores.append(score1)
                filtered_class_names.append(class1)
        
        boxes = filtered_boxes
        scores = filtered_scores
        class_names = filtered_class_names
        print(f"    ✓ Kept {len(boxes)} bboxes after removing overlaps (removed {len(removed_indices)} overlapping bboxes)")
        
        # ========================================================================
        # BƯỚC 2: OCR VÀ XỬ LÝ TỪNG BBOX (chưa có thứ tự)
        # ========================================================================
        # OCR từng bbox để lấy text, nhưng chưa sort theo thứ tự đọc
        # ========================================================================
        # Prepare batches for OCR
        valid_boxes_info = []
        messages_batch = []
        
        for idx, (box, score, class_name) in enumerate(zip(boxes, scores, class_names)):
            x1, y1, x2, y2 = box
            width = x2 - x1
            height = y2 - y1
            
            is_in_header_zone = y1 < image.shape[0] * 0.20
            if class_name == 'abandon':
                if (height > width * 3 or width > height * 3) and not is_in_header_zone:
                    continue
            
            bbox_image = crop_bbox(image, box, padding=10)
            if bbox_image.shape[0] == 0 or bbox_image.shape[1] == 0:
                continue
            
            if enable_ocr and qwen_model is not None and processor is not None:
                pil_image = Image.fromarray(cv2.cvtColor(bbox_image, cv2.COLOR_BGR2RGB))
                prompt_text = (
                    "Trích xuất NGUYÊN VĂN (không được bịa thêm) trong hình ảnh sang định dạng HTML.\n"
                    "YÊU CẦU CỐT LÕI:\n"
                    "- KHÔNG ĐƯỢC tự suy diễn thành danh sách (bullet/numbered list) nếu trên ảnh không THẤY RÕ từng gạch đầu dòng/dấu chấm số.\n"
                    "YÊU CẦU ĐỊNH DẠNG:\n"
                    "- Giữ nguyên định dạng gốc của văn bản. Nếu chữ trong ảnh là in đậm, in nghiêng, hoặc gạch chân, hãy dùng thẻ <b>, <i>, <u> tương ứng.\n"
                    "- Nếu là đoạn văn bình thường, dùng thẻ <p>.\n"
                    "- Nếu là danh sách có dấu đầu dòng hoặc đánh số, dùng thẻ <ul> hoặc <ol> cùng với <li>.\n"
                    "- LƯU Ý QUAN TRỌNG: Chỉ thêm các thẻ định dạng đặc biệt (như <b>, <i>, <ul>) khi bạn nhìn thấy RÕ RÀNG định dạng đó trong ảnh. Nếu nét chữ chỉ hơi đậm hơn một chút do chất lượng in, hãy coi đó là chữ bình thường.\n"
                    "Chỉ xuất ra mã HTML thuần túy, tuyệt đối không giải thích."
                )
                messages = [
                    {
                        "role": "system",
                        "content": [
                            {"type": "text", "text": (
                                "Bạn là công cụ OCR chuyển đổi ảnh văn bản bản cứng thành HTML. "
                                "Nhiệm vụ của bạn là nhận diện chính xác văn bản và cấu trúc (như in đậm, in nghiêng, danh sách) rồi xuất ra mã HTML. "
                                "Trích xuất trung thực, không tự suy diễn thêm định dạng nếu nó không thực sự xuất hiện trong ảnh."
                            )},
                        ],
                    },
                    {
                        "role": "user",
                        "content": [
                            {"type": "image", "image": pil_image},
                            {"type": "text", "text": prompt_text},
                        ],
                    }
                ]
                messages_batch.append(messages)
            
            valid_boxes_info.append({
                'idx': idx, 'box': box, 'score': score, 'class_name': class_name, 'bbox_image': bbox_image, 'text': ""
            })
            
        # Run batched OCR in chunks to prevent VRAM spikes
        if enable_ocr and qwen_model is not None and processor is not None and messages_batch:
            BATCH_SIZE = 16
            print(f"      Running batched OCR for {len(messages_batch)} bboxes (Batch size: {BATCH_SIZE})...")
            
            all_texts = []
            import math
            total_batches = math.ceil(len(messages_batch) / BATCH_SIZE)
            
            for i in range(0, len(messages_batch), BATCH_SIZE):
                batch_msgs = messages_batch[i:i+BATCH_SIZE]
                print(f"        Processing batch {i//BATCH_SIZE + 1}/{total_batches}...")
                try:
                    text_inputs = processor.apply_chat_template(batch_msgs, tokenize=False, add_generation_prompt=True)
                    image_inputs, video_inputs = process_vision_info(batch_msgs)
                    
                    inputs = processor(
                        text=text_inputs,
                        images=image_inputs,
                        videos=video_inputs,
                        padding=True,
                        return_tensors="pt",
                    )
                    inputs = inputs.to(qwen_model.device)
                    
                    with torch.no_grad():
                        # Generate with memory optimization for batching
                        generated_ids = qwen_model.generate(**inputs, max_new_tokens=2048)
                    
                    generated_ids_trimmed = [
                        out_ids[len(in_ids):] for in_ids, out_ids in zip(inputs.input_ids, generated_ids)
                    ]
                    texts = processor.batch_decode(
                        generated_ids_trimmed, skip_special_tokens=True, clean_up_tokenization_spaces=False
                    )
                    all_texts.extend(texts)
                    
                    # Clear cache to free VRAM for next batch
                    torch.cuda.empty_cache()
                except Exception as e:
                    print(f"        ⚠️ Batch {i//BATCH_SIZE + 1} failed: {e}")
                    # Append empty strings for failures to keep alignment
                    all_texts.extend([""] * len(batch_msgs))
                    torch.cuda.empty_cache()
            
            # ============================================================
            # Phát hiện output rác bằng cách so sánh với prompt
            # Thay vì hardcode từng pattern, dùng word-overlap ratio
            # ============================================================
            # Tách prompt thành tập từ đặc trưng (loại bỏ từ phổ biến)
            prompt_words = set(prompt_text.lower().split())
            # Thêm từ khóa refusal phổ biến
            refusal_keywords = {'sorry', 'cannot', 'unable', 'apologize', 'không thể', 'xin lỗi', 'không rõ'}
            
            # Map texts back to valid boxes
            for i, valid_box in enumerate(valid_boxes_info):
                if i < len(all_texts):
                    t = all_texts[i]
                    # Strip markdown/html code block wrappers
                    t = t.replace("```markdown\n", "").replace("```markdown\\n", "").replace("```html\n", "").replace("```html\\n", "").replace("```", "").strip()
                    # Strip literal prefixes
                    for prefix in ["markdown", "html", "text:"]:
                        if t.lower().startswith(prefix):
                            t = t[len(prefix):].strip()
                    
                    # Smart junk detection: so sánh output với prompt
                    t_lower = t.lower()
                    t_words = set(t_lower.split())
                    
                    is_junk = False
                    if t_words:
                        # Tính tỉ lệ từ trong output trùng với prompt
                        overlap = t_words & prompt_words
                        overlap_ratio = len(overlap) / len(t_words)
                        # Nếu > 50% từ trong output trùng với prompt → echo
                        if overlap_ratio > 0.5 and len(t_words) > 3:
                            is_junk = True
                        # Kiểm tra refusal keywords
                        if t_words & refusal_keywords and len(t_words) < 30:
                            is_junk = True
                    
                    if is_junk:
                        t = ""
                    
                    # Strip fake leading bullet "- " cho non-list text
                    if t.startswith("- ") and t.count("\n- ") == 0:
                        t = t[2:].strip()
                    
                    valid_box['text'] = t
                else:
                    valid_box['text'] = ""

                
        # Post-process valid boxes
        seen_texts_positions = {}
        page_bboxes = []
        
        for box_info in valid_boxes_info:
            idx = box_info['idx']
            box = box_info['box']
            score = box_info['score']
            class_name = box_info['class_name']
            text = box_info['text']
            x1, y1, x2, y2 = box
            width = x2 - x1
            height = y2 - y1
            
            if not text or text in ['[No text detected]', '[Empty bbox]'] or text.startswith('[OCR Error:'):
                continue
            
            # Check if this is a document number - never skip these
            text_lower = text.lower().strip()
            is_doc_number = text_lower.startswith('số:') or text_lower.startswith('số ') or ('số:' in text_lower and '/' in text_lower and ('nđ-cp' in text_lower or 'nd-cp' in text_lower))
            
            # Check for duplicates - improved logic
            text_normalized = text.lower().strip()
            
            # Normalize text for comparison (remove diacritics for "chính phủ" detection)
            def normalize_for_comparison(text):
                """Normalize text for comparison, handling both with/without diacritics"""
                # Remove diacritics for comparison
                import unicodedata
                text_no_diacritics = ''.join(c for c in unicodedata.normalize('NFD', text) 
                                           if unicodedata.category(c) != 'Mn')
                return text_no_diacritics.lower().strip()
            
            text_normalized_no_diacritics = normalize_for_comparison(text)
            
            # Special handling for header elements - only keep one instance
            is_chinh_phu = (text_normalized == 'chính phủ' or 
                           text_normalized_no_diacritics == 'chinh phu' or
                           text_normalized == 'chinh phu')
            is_quoc_hieu = ('cộng hòa' in text_normalized and 'việt nam' in text_normalized)
            is_tieu_ngu = ('độc lập' in text_normalized and 'tự do' in text_normalized)
            is_nghi_dinh = (text_normalized == 'nghị định')
            
            is_header_element = is_chinh_phu or is_quoc_hieu or is_tieu_ngu or is_nghi_dinh
            
            # Check for "CHÍNH PHỦ" duplicates (handle both with/without diacritics)
            should_skip_chinh_phu = False
            if is_chinh_phu:
                # Check if we already have a "CHÍNH PHỦ" (with or without diacritics)
                for existing_text, existing_data in list(seen_texts_positions.items()):
                    existing_normalized = normalize_for_comparison(existing_text)
                    # If it's also "chinh phu" (with or without diacritics)
                    if existing_normalized == 'chinh phu':
                        existing_coords = existing_data['coords']
                        existing_x1, existing_y1, existing_x2, existing_y2 = existing_coords
                        
                        # Check if they're in similar position (header zone, left side)
                        center_x = (x1 + x2) / 2
                        center_y = (y1 + y2) / 2
                        existing_center_x = (existing_x1 + existing_x2) / 2
                        existing_center_y = (existing_y1 + existing_y2) / 2
                        distance = ((center_x - existing_center_x) ** 2 + (center_y - existing_center_y) ** 2) ** 0.5
                        
                        # If very close (<100px) in header zone, it's a duplicate
                        if distance < 100 and y1 < image.shape[0] * 0.20:
                            # Keep the one with diacritics (CHÍNH PHỦ), skip the one without (CHINH PHU)
                            has_diacritics = 'chính phủ' in text_normalized
                            existing_has_diacritics = 'chính phủ' in existing_text.lower()
                            
                            if has_diacritics and not existing_has_diacritics:
                                # This one has diacritics, replace the old one
                                # Remove old bbox
                                page_bboxes = [b for b in page_bboxes if not (
                                    normalize_for_comparison(b['text'].lower().strip()) == 'chinh phu' and
                                    abs(b['x1'] - existing_x1) < 50 and abs(b['y1'] - existing_y1) < 50
                                )]
                                # Update tracking - remove old, will add new one later
                                seen_texts_positions.pop(existing_text, None)
                                break
                            else:
                                # This one doesn't have diacritics or old one already has diacritics, skip
                                should_skip_chinh_phu = True
                                break
            
            if should_skip_chinh_phu:
                continue
            
            # For other header elements, if already exists, skip (only keep first one)
            if is_header_element and not is_chinh_phu and text_normalized in seen_texts_positions:
                continue
            
            # Check if this text already exists (for non-header elements)
            if text_normalized in seen_texts_positions:
                existing_bbox = seen_texts_positions[text_normalized]
                existing_x1, existing_y1, existing_x2, existing_y2 = existing_bbox['coords']
                
                # Calculate overlap or distance
                overlap_x = max(0, min(x2, existing_x2) - max(x1, existing_x1))
                overlap_y = max(0, min(y2, existing_y2) - max(y1, existing_y1))
                overlap_area = overlap_x * overlap_y
                
                bbox_area = width * height
                existing_area = (existing_x2 - existing_x1) * (existing_y2 - existing_y1)
                overlap_ratio = overlap_area / min(bbox_area, existing_area) if min(bbox_area, existing_area) > 0 else 0
                
                # If significant overlap (>30%) or very close (<50px distance), it's a duplicate
                center_x = (x1 + x2) / 2
                center_y = (y1 + y2) / 2
                existing_center_x = (existing_x1 + existing_x2) / 2
                existing_center_y = (existing_y1 + existing_y2) / 2
                distance = ((center_x - existing_center_x) ** 2 + (center_y - existing_center_y) ** 2) ** 0.5
                
                is_duplicate = overlap_ratio > 0.3 or distance < 50
                
                # Skip duplicates (except document numbers)
                if is_duplicate and not is_doc_number:
                    continue
            
            # Add to tracking
            seen_texts_positions[text_normalized] = {
                'coords': (x1, y1, x2, y2),
                'score': float(score)
            }
            
            x1, y1, x2, y2 = box
            page_bboxes.append({
                'class': class_name,
                'confidence': float(score),
                'text': text,
                'x1': float(x1),
                'y1': float(y1),
                'x2': float(x2),
                'y2': float(y2),
                'center_x': (x1 + x2) / 2,
                'center_y': (y1 + y2) / 2,
            })
        
        # Merge related bboxes for "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" if split
        merged_bboxes = []
        i = 0
        while i < len(page_bboxes):
            bbox = page_bboxes[i]
            text_lower = bbox['text'].lower().strip()
            
            # Check if this is part of "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
            is_quoc_hieu_part = ('cộng hòa' in text_lower or 'nghĩa việt nam' in text_lower or 
                                'xã hội' in text_lower or 'chủ nghĩa' in text_lower)
            
            if is_quoc_hieu_part and i < len(page_bboxes) - 1:
                # Check if next bbox is also part of quoc hieu
                next_bbox = page_bboxes[i + 1]
                next_text_lower = next_bbox['text'].lower().strip()
                is_next_quoc_hieu = ('cộng hòa' in next_text_lower or 'nghĩa việt nam' in next_text_lower or 
                                     'xã hội' in next_text_lower or 'chủ nghĩa' in next_text_lower)
                
                # Check if they're on same row (similar Y position) and close horizontally
                y_diff = abs(bbox['y1'] - next_bbox['y1'])
                x_gap = next_bbox['x1'] - bbox['x2']
                
                if is_next_quoc_hieu and y_diff < 50 and x_gap < 100:
                    # Merge them
                    merged_text = bbox['text'].strip() + ' ' + next_bbox['text'].strip()
                    merged_bbox = {
                        'class': bbox['class'],
                        'confidence': max(bbox['confidence'], next_bbox['confidence']),
                        'text': merged_text,
                        'x1': min(bbox['x1'], next_bbox['x1']),
                        'y1': min(bbox['y1'], next_bbox['y1']),
                        'x2': max(bbox['x2'], next_bbox['x2']),
                        'y2': max(bbox['y2'], next_bbox['y2']),
                        'center_x': (min(bbox['x1'], next_bbox['x1']) + max(bbox['x2'], next_bbox['x2'])) / 2,
                        'center_y': (min(bbox['y1'], next_bbox['y1']) + max(bbox['y2'], next_bbox['y2'])) / 2,
                    }
                    merged_bboxes.append(merged_bbox)
                    i += 2  # Skip next bbox
                    continue
            
            merged_bboxes.append(bbox)
            i += 1
        
        # Trích header cho trang 1 và loại bỏ khỏi danh sách content
        # UPDATE: We now use absolute coordinates to place text exactly where it appears in YOLO detections.
        # Header components will also be drawn exactly where they are detected.
        header_elements = None
        content_bboxes = merged_bboxes
        
        # ========================================================================
        # BƯỚC 3: SORT BBOXES THEO THỨ TỰ ĐỌC TỰ NHIÊN
        # ========================================================================
        # Sort bboxes theo quy luật: từ trên xuống dưới (top-to-bottom),
        # từ trái qua phải (left-to-right) - như mắt người đọc sách
        # ========================================================================
        if content_bboxes:
            page_width = max(bbox['x2'] for bbox in content_bboxes)
            
            # Debug: Print Y positions before sorting (only for first page)
            if page_idx == 0:
                print(f"    Debug: Bbox Y positions before sorting:")
                for i, bbox in enumerate(content_bboxes[:10], 1):
                    text_preview = bbox['text'][:50].replace('\n', ' ')
                    print(f"      {i}. Y={bbox['y1']:.0f}: {text_preview}...")
            
            # Sort theo thứ tự đọc tự nhiên: top-to-bottom, left-to-right
            sorted_bboxes = sort_bboxes_by_position(content_bboxes, page_width)
            
            # ========================================================================
            # BƯỚC 4: ĐÁNH SỐ THỨ TỰ ĐỌC (READING ORDER)
            # ========================================================================
            # Sau khi sort xong, đánh số thứ tự 1, 2, 3, ... cho các bboxes
            # theo thứ tự đọc tự nhiên như mắt người đọc sách
            # Số thứ tự này được lưu vào bbox['reading_order'] để đảm bảo
            # OCR và xử lý text theo đúng thứ tự này
            # ========================================================================
            print(f"    Assigning reading order numbers (top-to-bottom, left-to-right)...")
            for reading_order, bbox in enumerate(sorted_bboxes, start=1):
                bbox['reading_order'] = reading_order
            
            # Debug: Print reading order and Y positions after sorting (only for first page)
            if page_idx == 0:
                print(f"    Debug: Bbox reading order and positions after sorting:")
                for bbox in sorted_bboxes[:12]:
                    reading_order = bbox.get('reading_order', '?')
                    text_preview = bbox['text'][:50].replace('\n', ' ')
                    print(f"      #{reading_order} (Y={bbox['y1']:.0f}, X={bbox['x1']:.0f}): {text_preview}...")
            
            all_pages_bboxes.append((page_idx + 1, sorted_bboxes, image.shape[1], image.shape[0], header_elements))
            print(f"    ✓ Processed {len(sorted_bboxes)} bboxes with text (merged {len(page_bboxes) - len(merged_bboxes)} related bboxes)")
        else:
            print(f"    ⚠️  No valid bboxes found for this page")
    
    # Step 6: Add all pages to Word document
    # ========================================================================
    # Sử dụng DrawingML textbox (wp:anchor) để đặt mỗi bbox vào đúng vị trí
    # trên trang Word. Mỗi bbox → 1 textbox duy nhất chứa ALL paragraphs.
    # ========================================================================
    print("\n[Step 6] Adding pages to Word document...")
    
    # EMU constants: 1 inch = 914400 EMU, 1 pt = 12700 EMU, 1 cm = 360000 EMU
    EMU_PER_PT = 12700
    
    # A4 page dimensions in points
    A4_WIDTH_PT = 595
    A4_HEIGHT_PT = 842
    
    # Margins in points (match the section margins set earlier: left=3cm, right=2cm, top=1.5cm)
    LEFT_MARGIN_PT = 3.0 * 28.3465  # ~85 pt
    TOP_MARGIN_PT = 1.5 * 28.3465   # ~42.5 pt
    
    textbox_id_counter = 100  # unique id for each textbox shape
    
    def create_textbox_element(textbox_id, x_emu, y_emu, w_emu, h_emu, content_paragraphs_xml, wrap_type="none"):
        """
        Tạo DrawingML textbox element (wp:anchor) để neo tại vị trí tuyệt đối trên trang.
        
        Parameters:
            textbox_id: unique shape id
            x_emu: horizontal position in EMU from page edge
            y_emu: vertical position in EMU from page edge  
            w_emu: width in EMU
            h_emu: height in EMU (min height, auto-extend if needed)
            content_paragraphs_xml: raw XML string of <w:p> elements to put inside
            wrap_type: 'none' (floating) or 'square' (text wrapping)
        
        Returns:
            OxmlElement for w:drawing
        """
        from lxml import etree
        
        nsmap = {
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
            'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
            'v': 'urn:schemas-microsoft-com:vml',
        }
        
        if wrap_type == 'square':
            wrap_xml = '<wp:wrapSquare wrapText="bothSides"/>'
        elif wrap_type == 'topAndBottom':
            wrap_xml = '<wp:wrapTopAndBottom/>'
        else:
            wrap_xml = '<wp:wrapNone/>'
        
        # Build the anchor XML
        drawing_xml = f'''<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                     xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                     xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                     xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
                     xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
                     xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
                     xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
          <wp:anchor distT="0" distB="0" distL="0" distR="0"
                     simplePos="0" relativeHeight="{textbox_id}"
                     behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="page">
              <wp:posOffset>{x_emu}</wp:posOffset>
            </wp:positionH>
            <wp:positionV relativeFrom="page">
              <wp:posOffset>{y_emu}</wp:posOffset>
            </wp:positionV>
            <wp:extent cx="{w_emu}" cy="{h_emu}"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            {wrap_xml}
            <wp:docPr id="{textbox_id}" name="TextBox {textbox_id}"/>
            <a:graphic>
              <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                <wps:wsp>
                  <wps:cNvSpPr txBox="1"/>
                  <wps:spPr>
                    <a:xfrm>
                      <a:off x="0" y="0"/>
                      <a:ext cx="{w_emu}" cy="{h_emu}"/>
                    </a:xfrm>
                    <a:prstGeom prst="rect">
                      <a:avLst/>
                    </a:prstGeom>
                    <a:noFill/>
                    <a:ln>
                      <a:noFill/>
                    </a:ln>
                  </wps:spPr>
                  <wps:txbx>
                    <w:txbxContent>
                      {content_paragraphs_xml}
                    </w:txbxContent>
                  </wps:txbx>
                  <wps:bodyPr wrap="square" lIns="0" tIns="0" rIns="0" bIns="0"
                              anchor="t" anchorCtr="0">
                    <a:noAutofit/>
                  </wps:bodyPr>
                </wps:wsp>
              </a:graphicData>
            </a:graphic>
          </wp:anchor>
        </w:drawing>'''
        
        drawing_element = etree.fromstring(drawing_xml)
        return drawing_element
    
    def paragraphs_to_xml(paragraphs):
        """
        Chuyển danh sách python-docx Paragraph objects thành raw XML string.
        """
        from lxml import etree
        xml_parts = []
        for p in paragraphs:
            xml_parts.append(etree.tostring(p._p, encoding='unicode'))
        return '\n'.join(xml_parts)
    
    for page_num, sorted_bboxes, page_width, page_height, header_elements in all_pages_bboxes:
        # Add page break if not first page
        if page_num > 1:
            page_break_para = doc.add_paragraph()
            page_break_para.paragraph_format.space_before = Pt(0)
            page_break_para.paragraph_format.space_after = Pt(0)
            page_break_para.paragraph_format.line_spacing = 1.0
            page_break_run = page_break_para.add_run()
            page_break_run.add_break(WD_BREAK.PAGE)
        
        # Đảm bảo section có kích thước chuẩn A4
        section = doc.sections[-1]
        section.page_width = Pt(A4_WIDTH_PT)
        section.page_height = Pt(A4_HEIGHT_PT)
        
        # Tạo 1 paragraph ẩn để neo tất cả textboxes của trang này, tránh tạo dòng trống
        page_anchor_para = doc.add_paragraph()
        page_anchor_para.paragraph_format.space_before = Pt(0)
        page_anchor_para.paragraph_format.space_after = Pt(0)
        page_anchor_para.paragraph_format.line_spacing = 1.0
        page_anchor_run = page_anchor_para.add_run()
        page_anchor_run.font.size = Pt(1)
        
        for bbox in sorted_bboxes:
            text = bbox['text'].strip()
            if not text:
                continue
            
            x1_px, y1_px = bbox['x1'], bbox['y1']
            x2_px, y2_px = bbox['x2'], bbox['y2']
            
            # ============================================================
            # Tính toạ độ EMU (English Metric Units) từ pixel coordinates
            # Pixel → tỉ lệ trên trang → điểm (pt) → EMU
            # ============================================================
            x_pt = (x1_px / page_width) * A4_WIDTH_PT
            y_pt = (y1_px / page_height) * A4_HEIGHT_PT
            width_pt = ((x2_px - x1_px) / page_width) * A4_WIDTH_PT
            height_pt = ((y2_px - y1_px) / page_height) * A4_HEIGHT_PT
            
            # Mở rộng width thêm 15% để text ít bị wrap hơn, giảm overlap
            width_pt = width_pt * 1.15
            # Giới hạn width không vượt quá lề phải trang (A4_WIDTH_PT)
            max_width = A4_WIDTH_PT - x_pt
            width_pt = min(width_pt, max_width)
            
            # Đảm bảo kích thước tối thiểu
            width_pt = max(width_pt, 30)
            height_pt = max(height_pt, 12)
            
            x_emu = int(x_pt * EMU_PER_PT)
            y_emu = int(y_pt * EMU_PER_PT)
            w_emu = int(width_pt * EMU_PER_PT)
            h_emu = int(height_pt * EMU_PER_PT)
            
            # ============================================================
            # Tạo nội dung cho textbox: convert markdown → HTML → Word paragraphs
            # Dùng tạm 1 Document phụ để HtmlToDocx render HTML ra paragraphs,
            # sau đó copy XML sang textbox content.
            # ============================================================
            # Nếu text đã là HTML (chứa tags), dùng trực tiếp
            # Nếu là plain text hoặc markdown, convert qua markdown trước
            if '<' in text and ('>' in text) and any(tag in text.lower() for tag in ['<p>', '<b>', '<i>', '<u>', '<ul>', '<ol>', '<li>', '<br', '<h']):
                html_content = text
            else:
                html_content = markdown.markdown(text)
            
            # Sanitize HTML: loại bỏ <img> tags không có src (gây crash HtmlToDocx)
            html_content = re.sub(r'<img(?![^>]*src\s*=)[^>]*/?\s*>', '', html_content, flags=re.IGNORECASE)
            # Loại bỏ các tag lạ không phải HTML chuẩn (ví dụ: <M.S.D.N.:0318788>)
            html_content = re.sub(r'<(?!/?(?:p|b|i|u|s|a|br|hr|h[1-6]|ul|ol|li|table|tr|td|th|thead|tbody|div|span|strong|em|sub|sup|pre|code|blockquote)\b)[A-Z][^>]*>', '', html_content, flags=re.IGNORECASE)
            
            # Tạo document phụ để render HTML
            temp_doc = Document()
            try:
                if html_parser is not None:
                    html_parser.add_html_to_document(html_content, temp_doc)
                else:
                    p = temp_doc.add_paragraph()
                    p.add_run(text)
            except Exception as e:
                # Fallback: nếu HTML parsing lỗi, dùng plain text
                print(f"      ⚠️ HTML parse error, using plain text: {e}")
                temp_doc = Document()
                p = temp_doc.add_paragraph()
                # Strip tất cả HTML tags để lấy plain text
                plain_text = re.sub(r'<[^>]+>', ' ', text).strip()
                plain_text = re.sub(r'\s+', ' ', plain_text)
                p.add_run(plain_text)
            
            # Lấy các paragraphs đã render (bỏ qua paragraph trống đầu tiên mặc định)
            temp_paragraphs = temp_doc.paragraphs
            if temp_paragraphs and not temp_paragraphs[0].text.strip():
                temp_paragraphs = temp_paragraphs[1:]
            
            if not temp_paragraphs:
                continue
            
            # Đặt font và spacing cho các paragraph bên trong textbox
            # GIỮ NGUYÊN bold/italic/underline mà HtmlToDocx đã set từ HTML
            for tp in temp_paragraphs:
                tp.paragraph_format.space_after = Pt(0)
                tp.paragraph_format.space_before = Pt(0)
                tp.paragraph_format.line_spacing = 1.0
                for run in tp.runs:
                    # Lưu lại formatting hiện tại trước khi set font
                    existing_bold = run.font.bold
                    existing_italic = run.font.italic
                    existing_underline = run.font.underline
                    
                    run.font.name = 'Times New Roman'
                    if run._element.rPr is not None:
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    else:
                        rPr = OxmlElement('w:rPr')
                        run._element.insert(0, rPr)
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                        rPr.append(rFonts)
                    run.font.size = Pt(10)
                    
                    # Khôi phục bold/italic/underline
                    if existing_bold is not None:
                        run.font.bold = existing_bold
                    if existing_italic is not None:
                        run.font.italic = existing_italic
                    if existing_underline is not None:
                        run.font.underline = existing_underline
            
            # Chuyển paragraph objects → XML string
            content_xml = paragraphs_to_xml(temp_paragraphs)
            
            # Tạo textbox element và gắn vào document
            textbox_id_counter += 1
            drawing_elem = create_textbox_element(
                textbox_id_counter, x_emu, y_emu, w_emu, h_emu, content_xml
            )
            
            # Gắn vào paragraph neo
            page_anchor_run._element.append(drawing_elem)
        
    # --- STEP 8 (REVISED): Render bboxes sang văn bản thường dùng indent + invisible table ---
    print("\n[Step 8] Creating clean transcript (indent + table approach)...")

    transcript_doc = Document()
    for section in transcript_doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    A4_WIDTH_PT = 595.0
    A4_HEIGHT_PT = 842.0
    CONTENT_WIDTH_CM = 15.0  # A4 trừ lề trái 3cm + lề phải 2cm
    LEFT_MARGIN_CM = 3.0

    def bbox_to_x1_cm(bbox, page_width):
        """Tọa độ x1 của bbox quy ra cm tuyệt đối trên A4 (tính từ lề trái trang)"""
        return (bbox['x1'] / page_width) * (A4_WIDTH_PT / 28.3465)

    def bbox_to_width_cm(bbox, page_width):
        return ((bbox['x2'] - bbox['x1']) / page_width) * (A4_WIDTH_PT / 28.3465)

    def get_indent_cm(bbox, page_width):
        """Indent tính từ lề content (3cm), không âm"""
        return max(0.0, bbox_to_x1_cm(bbox, page_width) - LEFT_MARGIN_CM)

    def bboxes_overlap_y(a, b, threshold=0.4):
        """Kiểm tra 2 bbox có overlap Y >= threshold * min_height không"""
        overlap = min(a['y2'], b['y2']) - max(a['y1'], b['y1'])
        min_h = min(a['y2'] - a['y1'], b['y2'] - b['y1'])
        return min_h > 0 and (overlap / min_h) >= threshold

    def bboxes_overlap_x(a, b):
        """Kiểm tra 2 bbox có overlap X không (nếu có → không thể song song)"""
        return min(a['x2'], b['x2']) - max(a['x1'], b['x1']) > 10

    def set_table_borders(table):
        """Helper to clear table borders in python-docx"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)
        
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = tblBorders.find(qn('w:' + border_name))
            if border is None:
                border = OxmlElement('w:' + border_name)
                tblBorders.append(border)
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')

    def group_into_rows(sorted_bboxes):
        """
        Nhóm sorted_bboxes thành rows. Mỗi row là list of bbox.
        Row có >1 bbox = song song (multi-column).
        Dùng overlap Y với TẤT CẢ bbox trong row hiện tại để quyết định.
        """
        rows = []
        current_row = []
        
        for bbox in sorted_bboxes:
            if not bbox['text'].strip():
                continue
            if not current_row:
                current_row.append(bbox)
            else:
                # Bbox mới có overlap Y với ít nhất 1 bbox trong row hiện tại?
                has_y_overlap = any(bboxes_overlap_y(bbox, b) for b in current_row)
                # Bbox mới có overlap X với bất kỳ bbox nào trong row không?
                has_x_overlap = any(bboxes_overlap_x(bbox, b) for b in current_row)
                
                if has_y_overlap and not has_x_overlap:
                    # Song song thật sự → thêm vào row
                    current_row.append(bbox)
                else:
                    rows.append(current_row)
                    current_row = [bbox]
        
        if current_row:
            rows.append(current_row)
        return rows

    def render_bbox_to_paragraphs(bbox, page_width, target_doc, indent_override=None):
        """Render 1 bbox thành list paragraph, gắn indent từ tọa độ x1"""
        text = bbox['text'].strip()
        if not text:
            return
        
        # Convert text → HTML → paragraphs (giữ nguyên logic hiện tại)
        if '<' in text and '>' in text and any(t in text.lower() for t in ['<p>', '<b>', '<i>', '<u>', '<ul>', '<ol>', '<li>', '<br', '<h']):
            html_content = text
        else:
            html_content = markdown.markdown(text)
        
        html_content = re.sub(r'<img(?![^>]*src\s*=)[^>]*/?\s*>', '', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'<(?!/?(?:p|b|i|u|s|a|br|hr|h[1-6]|ul|ol|li|table|tr|td|th|thead|tbody|div|span|strong|em|sub|sup|pre|code|blockquote)\b)[A-Z][^>]*>', '', html_content, flags=re.IGNORECASE)
        
        temp_doc = Document()
        try:
            if html_parser is not None:
                html_parser.add_html_to_document(html_content, temp_doc)
            else:
                temp_doc.add_paragraph().add_run(text)
        except Exception:
            temp_doc = Document()
            plain = re.sub(r'<[^>]+>', ' ', text).strip()
            temp_doc.add_paragraph().add_run(re.sub(r'\s+', ' ', plain))
        
        temp_paras = [p for p in temp_doc.paragraphs if p.text.strip()]
        if not temp_paras:
            return
        
        indent_cm = indent_override if indent_override is not None else get_indent_cm(bbox, page_width)
        
        for tp in temp_paras:
            import copy
            from docx.text.paragraph import Paragraph
            new_p = copy.deepcopy(tp._p)
            body = target_doc._element.body
            if body.sectPr is not None:
                body.sectPr.addprevious(new_p)
            else:
                body.append(new_p)
            cloned = Paragraph(new_p, target_doc)
            
            # Apply font
            for run in cloned.runs:
                saved = (run.font.bold, run.font.italic, run.font.underline)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                if run._element.rPr is not None:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                run.font.bold, run.font.italic, run.font.underline = saved
            
            # Apply indent từ tọa độ x1 thật
            if indent_cm > 0.3:
                pPr = new_p.get_or_add_pPr()
                ind = pPr.find(qn('w:ind'))
                twips = int(indent_cm * 567)
                if ind is None:
                    ind = OxmlElement('w:ind')
                    ind.set(qn('w:left'), str(twips))
                    pPr.append(ind)
                else:
                    ind.set(qn('w:left'), str(twips))

    # --- STEP 8: Create clean transcript ---
    print("\n[Step 8] Creating clean transcript (indent + table approach)...")

    transcript_doc = Document()
    for section in transcript_doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.0)

    A4_WIDTH_PT      = 595.0
    A4_HEIGHT_PT     = 842.0
    CONTENT_WIDTH_CM = 15.0
    LEFT_MARGIN_CM   = 3.0

    # ── helpers ──────────────────────────────────────────────────────────────

    def _set_table_borders_none(table):
        tbl   = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)
        for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            el = tblBorders.find(qn(f'w:{edge}'))
            if el is None:
                el = OxmlElement(f'w:{edge}')
                tblBorders.append(el)
            el.set(qn('w:val'),   'none')
            el.set(qn('w:sz'),    '0')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), 'auto')
        # Zero cell margins để tránh trang trắng thừa
        tblCellMar = tblPr.find(qn('w:tblCellMar'))
        if tblCellMar is None:
            tblCellMar = OxmlElement('w:tblCellMar')
            tblPr.append(tblCellMar)
        for side in ['top', 'left', 'bottom', 'right']:
            mar = tblCellMar.find(qn(f'w:{side}'))
            if mar is None:
                mar = OxmlElement(f'w:{side}')
                tblCellMar.append(mar)
            mar.set(qn('w:w'),    '0')
            mar.set(qn('w:type'), 'dxa')

    def _apply_font(para, space_after_pt=4.0, space_before_pt=0.0):
        """Áp font Times New Roman 10pt, giữ nguyên bold/italic/underline."""
        para.paragraph_format.space_after  = Pt(space_after_pt)
        para.paragraph_format.space_before = Pt(space_before_pt)
        para.paragraph_format.line_spacing = 1.0
        for run in para.runs:
            b, i, u = run.font.bold, run.font.italic, run.font.underline
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            if run._element.rPr is not None:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            else:
                rPr    = OxmlElement('w:rPr')
                run._element.insert(0, rPr)
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                rPr.append(rFonts)
            if b is not None: run.font.bold      = b
            if i is not None: run.font.italic    = i
            if u is not None: run.font.underline = u

    def _render_text_to_raw_paragraphs(text):
        """Chuyển text thành list paragraphs - KHÔNG convert markdown để giữ nguyên OCR output."""
        # Chỉ dùng html_parser nếu text thực sự là HTML có tags
        if '<' in text and '>' in text and any(
            t in text.lower() for t in ['<p>', '<b>', '<i>', '<u>', '<ul>', '<ol>', '<li>', '<br', '<h']
        ):
            html_content = text
            html_content = re.sub(r'<img(?![^>]*src\s*=)[^>]*/?\s*>', '', html_content, flags=re.IGNORECASE)
            html_content = re.sub(
                r'<(?!/?(?:p|b|i|u|s|a|br|hr|h[1-6]|ul|ol|li|table|tr|td|th|thead|tbody|div|span|strong|em|sub|sup|pre|code|blockquote)\b)[A-Z][^>]*>',
                '', html_content, flags=re.IGNORECASE
            )
            tmp = Document()
            try:
                if html_parser is not None:
                    html_parser.add_html_to_document(html_content, tmp)
                else:
                    tmp.add_paragraph().add_run(text)
            except Exception:
                tmp = Document()
                tmp.add_paragraph().add_run(
                    re.sub(r'\s+', ' ', re.sub(r'<[^>]+>', ' ', text)).strip()
                )
        else:
            # Plain text: tách theo dòng, mỗi dòng là 1 paragraph
            tmp = Document()
            lines = text.splitlines()
            for line in lines:
                line = line.strip()
                if line:
                    tmp.add_paragraph().add_run(line)
            if not any(p.text.strip() for p in tmp.paragraphs):
                tmp.add_paragraph().add_run(text.strip())

        return [p for p in tmp.paragraphs if p.text.strip()]

    def _append_paragraphs_to_body(raw_paras, target_doc, indent_cm=0.0, space_after_pt=4.0, right_indent_cm=0.0):
        import copy
        from docx.text.paragraph import Paragraph
        for idx, tp in enumerate(raw_paras):
            new_p  = copy.deepcopy(tp._p)
            body   = target_doc._element.body
            if body.sectPr is not None:
                body.sectPr.addprevious(new_p)
            else:
                body.append(new_p)
            cloned = Paragraph(new_p, target_doc)
            is_last = (idx == len(raw_paras) - 1)
            _apply_font(cloned, space_after_pt=space_after_pt if is_last else 1.0)

            if indent_cm > 0.1 or right_indent_cm > 0.1:
                pPr   = new_p.get_or_add_pPr()
                ind   = pPr.find(qn('w:ind'))
                if ind is None:
                    ind = OxmlElement('w:ind')
                    pPr.append(ind)
                if indent_cm > 0.1:
                    existing = int(ind.get(qn('w:left'), '0') or '0')
                    ind.set(qn('w:left'), str(existing + int(indent_cm * 567)))
                if right_indent_cm > 0.1:
                    ind.set(qn('w:right'), str(int(right_indent_cm * 567)))

    def _append_paragraphs_to_cell(raw_paras, cell, alignment):
        """Copy raw paragraphs vào cell của table."""
        import copy
        from docx.text.paragraph import Paragraph
        for tp in raw_paras:
            new_p  = copy.deepcopy(tp._p)
            cell._element.append(new_p)
            cloned = Paragraph(new_p, cell)
            cloned.alignment = alignment
            _apply_font(cloned)

    def _bboxes_overlap_y(a, b, threshold=0.4):
        overlap = min(a['y2'], b['y2']) - max(a['y1'], b['y1'])
        h_a     = a['y2'] - a['y1']
        h_b     = b['y2'] - b['y1']
        min_h   = min(h_a, h_b)
        
        if min_h <= 0:
            return False
        
        # Loại trừ: nếu 1 bbox cao gấp 2.5 lần bbox kia
        # → bbox cao đó là paragraph nhiều dòng, không phải song song thật
        height_ratio = max(h_a, h_b) / min_h
        if height_ratio > 2.5:
            return False
    
        return (overlap / min_h) >= threshold

    def _bboxes_overlap_x(a, b, tol=10):
        return (min(a['x2'], b['x2']) - max(a['x1'], b['x1'])) > tol

    def _group_into_rows(bboxes):
        """Nhóm bboxes thành rows dựa trên overlap Y. Row > 1 phần tử = song song."""
        rows    = []
        current = []
        for bbox in bboxes:
            if not bbox['text'].strip():
                continue
            if not current:
                current.append(bbox)
            else:
                has_y = any(_bboxes_overlap_y(bbox, b) for b in current)
                has_x = any(_bboxes_overlap_x(bbox, b) for b in current)
                if has_y and not has_x:
                    current.append(bbox)
                else:
                    rows.append(current)
                    current = [bbox]
        if current:
            rows.append(current)
        return rows

    # ── main loop ─────────────────────────────────────────────────────────────

    for page_num, sorted_bboxes, page_width, page_height, header_elements in all_pages_bboxes:

        # Page break
        if page_num > 1:
            pb = transcript_doc.add_paragraph()
            pb.paragraph_format.space_before = Pt(0)
            pb.paragraph_format.space_after  = Pt(0)
            pb.paragraph_format.line_spacing = 1.0
            pb.add_run().add_break(WD_BREAK.PAGE)

        # Header table
        if header_elements:
            add_header_table(transcript_doc, header_elements)

        # Phân loại mep vs main
        margin_l    = page_width * 0.05
        margin_r    = page_width * 0.95
        mep_bboxes  = []
        main_bboxes = []
        for bbox in sorted_bboxes:
            if not bbox['text'].strip():
                continue
            if bbox['x1'] < margin_l or bbox['x2'] > margin_r:
                mep_bboxes.append(bbox)
            else:
                main_bboxes.append(bbox)

        # ── Render mep bboxes: floating textbox ──────────────────────────────
        if mep_bboxes:
            anchor_para = transcript_doc.add_paragraph()
            anchor_para.paragraph_format.space_before = Pt(0)
            anchor_para.paragraph_format.space_after  = Pt(0)
            anchor_para.paragraph_format.line_spacing = 1.0
            anchor_run = anchor_para.add_run()

            for bbox in mep_bboxes:
                raw_paras = _render_text_to_raw_paragraphs(bbox['text'])
                if not raw_paras:
                    continue
                for tp in raw_paras:
                    _apply_font(tp)
                content_xml = paragraphs_to_xml(raw_paras)

                x_pt      = (bbox['x1'] / page_width)             * A4_WIDTH_PT
                y_pt      = (bbox['y1'] / page_height)            * A4_HEIGHT_PT
                width_pt  = ((bbox['x2'] - bbox['x1']) / page_width)  * A4_WIDTH_PT
                height_pt = ((bbox['y2'] - bbox['y1']) / page_height) * A4_HEIGHT_PT
                width_pt  = min(max(width_pt * 1.15, 30), A4_WIDTH_PT - x_pt)
                height_pt = max(height_pt, 12)

                textbox_id_counter += 1
                drawing_elem = create_textbox_element(
                    textbox_id_counter,
                    int(x_pt      * EMU_PER_PT),
                    int(y_pt      * EMU_PER_PT),
                    int(width_pt  * EMU_PER_PT),
                    int(height_pt * EMU_PER_PT),
                    content_xml, wrap_type="none"
                )
                anchor_run._element.append(drawing_elem)

        # ── Render main bboxes: indent + invisible table ──────────────────────
        rows = _group_into_rows(main_bboxes)

        for row_idx, row in enumerate(rows):
            row.sort(key=lambda b: b['x1'])

            # Tính space_after từ Y-gap tới row tiếp theo
            if row_idx + 1 < len(rows):
                next_row       = rows[row_idx + 1]
                cur_y2         = max(b['y2'] for b in row)
                next_y1        = min(b['y1'] for b in next_row)
                gap_px         = max(next_y1 - cur_y2, 0)
                gap_pt         = (gap_px / page_height) * A4_HEIGHT_PT
                # Clamp: tối thiểu 2pt, tối đa 40pt
                space_after_pt = max(2.0, min(gap_pt, 40.0))
            else:
                space_after_pt = 4.0

            if len(row) == 1:
                bbox      = row[0]
                raw       = _render_text_to_raw_paragraphs(bbox['text'])
                indent_cm = max(0.0, (bbox['x1'] / page_width) * (A4_WIDTH_PT / 28.3465) - LEFT_MARGIN_CM)

                # Tính right indent: nếu bbox không trải hết đến lề phải content
                bbox_x2_cm        = (bbox['x2'] / page_width) * (A4_WIDTH_PT / 28.3465)
                content_right_cm  = (A4_WIDTH_PT / 28.3465) - 2.0  # lề phải 2cm
                right_indent_cm   = max(0.0, content_right_cm - bbox_x2_cm)
                # Chỉ áp dụng nếu bbox thực sự hẹp (tránh áp sai cho paragraph full-width)
                if right_indent_cm < 0.5:
                    right_indent_cm = 0.0

                _append_paragraphs_to_body(raw, transcript_doc, indent_cm,
                                           space_after_pt=space_after_pt,
                                           right_indent_cm=right_indent_cm)

            else:
                # Song song → invisible table, width cột theo tọa độ tuyệt đối
                table = transcript_doc.add_table(rows=1, cols=len(row))
                _set_table_borders_none(table)
                table.autofit = False

                row_x1   = min(b['x1'] for b in row)
                row_x2   = max(b['x2'] for b in row)
                row_span = max(row_x2 - row_x1, 1)

                for i, bbox in enumerate(row):
                    cell     = table.cell(0, i)
                    b_span   = bbox['x2'] - bbox['x1']
                    col_w_cm = (b_span / row_span) * CONTENT_WIDTH_CM
                    cell.width = Cm(max(col_w_cm, 1.0))

                    for p in list(cell.paragraphs):
                        cell._element.remove(p._element)

                    raw       = _render_text_to_raw_paragraphs(bbox['text'])
                    alignment = determine_alignment_by_position(bbox, page_width, page_height)

                    if raw:
                        _append_paragraphs_to_cell(raw, cell, alignment)
                    else:
                        cell.add_paragraph()

                    if not cell.paragraphs:
                        cell.add_paragraph()

                # Paragraph spacer sau table để tạo khoảng cách tương đương gap
                spacer = transcript_doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(0)
                spacer.paragraph_format.space_after  = Pt(0)
                spacer.paragraph_format.line_spacing = 1.0
                spacer_run = spacer.add_run()
                spacer_run.font.size = Pt(max(1.0, min(space_after_pt * 0.5, 12.0)))
    
    # Step 9: Save documents
    print(f"\n[Step 9] Saving Word documents...")

    def save_doc_safe(d, path):
        try:
            d.save(str(path))
            print(f"  ✓ Saved to: {path}")
            return path
        except PermissionError:
            import datetime
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            new_output = path.parent / f"{path.stem}_{timestamp}.docx"
            print(f"  ⚠️  Original file is locked, saving as: {new_output.name}")
            d.save(str(new_output))
            print(f"  ✓ Saved to: {new_output}")
            return new_output

    # Create layout output path using the original path
    # layout_path = output_docx.parent / f"{output_docx.stem}_layout{output_docx.suffix}"
    
    # # Save both versions
    # layout_docx = save_doc_safe(doc, layout_path)
    transcript_docx = save_doc_safe(transcript_doc, output_docx)
    output_docx = transcript_docx
    
    print("\n" + "=" * 80)
    print("PROCESSING COMPLETE")
    print("=" * 80)
    print(f"Main output file (transcript): {output_docx}")
    # print(f"Layout output file: {layout_docx}")
    print(f"Total pages: {len(images)}")
    # all_pages_bboxes giờ là (page_num, sorted_bboxes, page_width, page_height, header_elements)
    total_bboxes = sum(len(bboxes) for _, bboxes, _, _, _ in all_pages_bboxes)
    print(f"Total bboxes: {total_bboxes}")
    
    return output_docx


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Complete PDF to DOCX pipeline")
    parser.add_argument("pdf_path", type=str, help="Path to PDF file")
    parser.add_argument("--output", type=str, default=None, help="Output DOCX file path")
    parser.add_argument(
        "--model",
        type=str,
        default=None,
        help="Path to YOLO model file (default: ./doclayout_yolo_docstructbench_imgsz1024.pt)",
    )
    parser.add_argument(
        "--ocr-model",
        type=str,
        default=None,
        help="HuggingFace ID or local directory path for the Qwen2.5-VL model (default: ./Qwen2.5-VL-3B)",
    )
    parser.add_argument(
        "--use-4bit",
        action="store_true",
        help="Enable 4-bit quantization (off by default for max quality)",
    )
    parser.add_argument(
        "--load-8bit",
        action="store_true",
        help="Use 8-bit quantization instead of 4-bit",
    )
    parser.add_argument("--imgsz", type=int, default=1024, help="Image size for inference")
    parser.add_argument("--conf", type=float, default=0.1, help="Confidence threshold")
    parser.add_argument(
        "--no-ocr",
        action="store_true",
        help="Disable OCR (only detection)",
    )
    parser.add_argument(
        "--max-pages",
        type=int,
        default=None,
        help="Maximum number of pages to process (for testing)",
    )

    args = parser.parse_args()

    process_pdf_to_docx(
        pdf_path=args.pdf_path,
        output_docx=args.output,
        model_path=args.model,
        ocr_model_path=args.ocr_model,
        imgsz=args.imgsz,
        conf=args.conf,
        enable_ocr=not args.no_ocr,
        load_4bit=args.use_4bit,
        load_8bit=args.load_8bit,
        max_pages=args.max_pages,
    )

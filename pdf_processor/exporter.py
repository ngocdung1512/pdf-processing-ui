"""
Document Exporter - Enhanced version with better format preservation
Exports to HTML, Word, PDF with improved layout and formatting
"""

import json
from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime
import re

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except ImportError:
    SimpleDocTemplate = None


class DocumentExporter:
    """
    Enhanced Document Exporter with better format preservation
    """
    
    def __init__(self):
        pass
    
    def _hex_to_rgb(self, color_value: int) -> tuple:
        """Convert color integer to RGB tuple"""
        if color_value == 0:
            return (0, 0, 0)  # Black default
        r = (color_value >> 16) & 0xFF
        g = (color_value >> 8) & 0xFF
        b = color_value & 0xFF
        return (r, g, b)
    
    def export_to_html(
        self,
        data: Dict[str, Any],
        output_path: Path,
        preserve_layout: bool = True
    ) -> Path:
        """Export to HTML with preserved layout"""
        html_parts = []
        
        html_parts.append("""<!DOCTYPE html>
<html lang="vi">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Document Export</title>
    <style>
        body {
            font-family: 'Times New Roman', serif;
            margin: 0;
            padding: 20px;
            background: #f5f5f5;
        }
        .page {
            background: white;
            margin: 20px auto;
            padding: 40px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
            position: relative;
        }
        .block {
            margin: 5px 0;
        }
        .title {
            font-size: 18px;
            font-weight: bold;
            margin: 10px 0;
            color: #1a1a1a;
        }
        .header {
            font-size: 16px;
            font-weight: bold;
            margin: 8px 0;
            color: #2c2c2c;
        }
        .heading1 {
            font-size: 16px;
            font-weight: bold;
            margin: 10px 0 5px 0;
            color: #1a1a1a;
        }
        .heading2 {
            font-size: 14px;
            font-weight: bold;
            margin: 8px 0 4px 0;
            color: #2c2c2c;
        }
        .paragraph {
            font-size: 12px;
            line-height: 1.6;
            margin: 8px 0;
            text-align: justify;
            color: #333;
        }
        .text {
            font-size: 12px;
            line-height: 1.5;
            margin: 5px 0;
            color: #444;
        }
        .list_item {
            font-size: 12px;
            line-height: 1.5;
            margin: 5px 0 5px 20px;
            color: #444;
        }
        .table {
            margin: 15px 0;
            border-collapse: collapse;
            width: 100%;
        }
        .table td, .table th {
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }
        .table th {
            background-color: #f2f2f2;
            font-weight: bold;
        }
        .image {
            margin: 15px 0;
            text-align: center;
        }
        .metadata {
            font-size: 10px;
            color: #666;
            margin-top: 20px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
        }
    </style>
</head>
<body>
""")
        
        for page_data in data.get("pages", []):
            page_num = page_data.get("page_num", 0)
            width = page_data.get("width", 800)
            height = page_data.get("height", 1000)
            blocks = page_data.get("blocks", [])
            
            width_px = int(width * 1.33)
            height_px = int(height * 1.33)
            
            html_parts.append(f'<div class="page" style="width: {width_px}px; min-height: {height_px}px;">')
            html_parts.append(f'<h2 style="text-align: center; margin-bottom: 20px;">Trang {page_num + 1}</h2>')
            
            for block in blocks:
                block_type = block.get("type", "text")
                text = block.get("text", "")
                bbox = block.get("bbox", [0, 0, 0, 0])
                
                css_class = {
                    "title": "title",
                    "header": "header",
                    "heading1": "heading1",
                    "heading2": "heading2",
                    "paragraph": "paragraph",
                    "text": "text",
                    "list_item": "list_item",
                    "table": "table",
                    "image": "image"
                }.get(block_type, "text")
                
                if block_type == "table":
                    html_parts.append('<table class="table">')
                    rows = text.split('\n')
                    for row in rows[:20]:
                        if row.strip():
                            html_parts.append('<tr>')
                            cells = row.split('\t')[:10]
                            for cell in cells:
                                html_parts.append(f'<td>{self._escape_html(cell.strip())}</td>')
                            html_parts.append('</tr>')
                    html_parts.append('</table>')
                elif block_type == "image":
                    html_parts.append(f'<div class="image">[Image placeholder - bbox: {bbox}]</div>')
                else:
                    if text.strip():
                        html_parts.append(f'<div class="block {css_class}">{self._escape_html(text)}</div>')
            
            html_parts.append('</div>')
        
        metadata = data.get("metadata", {})
        html_parts.append('<div class="metadata">')
        html_parts.append(f'<p><strong>Loại PDF:</strong> {metadata.get("pdf_type", "unknown")}</p>')
        html_parts.append(f'<p><strong>Tổng số trang:</strong> {metadata.get("total_pages", 0)}</p>')
        html_parts.append(f'<p><strong>Thời gian xử lý:</strong> {metadata.get("processing_time", 0):.2f} giây</p>')
        html_parts.append(f'<p><strong>Layout Engine:</strong> {"Có" if metadata.get("layout_engine_used") else "Không"}</p>')
        html_parts.append(f'<p><strong>Xuất lúc:</strong> {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>')
        html_parts.append('</div>')
        
        html_parts.append("""
</body>
</html>
""")
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(html_parts))
        
        return output_path
    
    def export_to_word(
        self,
        data: Dict[str, Any],
        output_path: Path
    ) -> Path:
        """
        Export to Word (DOCX) with enhanced format preservation
        Preserves fonts, sizes, colors, bold, italic, spacing, alignment
        IMPORTANT: Does NOT auto-convert numbered items to Word bullets
        """
        if not Document:
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx")
        
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        for page_data in data.get("pages", []):
            page_num = page_data.get("page_num", 0)
            blocks = page_data.get("blocks", [])
            
            if page_num > 0:
                doc.add_page_break()
            
            # Add page title (optional, can remove)
            # title_para = doc.add_paragraph(f'Trang {page_num + 1}')
            # title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # title_run = title_para.runs[0]
            # title_run.font.size = Pt(14)
            # title_run.font.bold = True
            # doc.add_paragraph()
            
            for block in blocks:
                block_type = block.get("type", "text")
                text = block.get("text", "")
                
                if not text.strip() and block_type != "table":
                    continue
                
                # Get formatting info if available
                font_name = block.get("font", "Times New Roman")
                font_size = block.get("font_size", 12)
                color = block.get("color", 0)
                is_bold = block.get("bold", False)
                is_italic = block.get("italic", False)
                
                if block_type == "title":
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.size = Pt(max(font_size, 18))
                    run.font.bold = True
                    run.font.name = font_name
                    if color > 0:
                        rgb = self._hex_to_rgb(color)
                        run.font.color.rgb = RGBColor(*rgb)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.space_after = Pt(6)
                
                elif block_type == "header":
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.size = Pt(max(font_size, 14))
                    run.font.bold = True
                    run.font.name = font_name
                    if color > 0:
                        rgb = self._hex_to_rgb(color)
                        run.font.color.rgb = RGBColor(*rgb)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.space_after = Pt(4)
                
                elif block_type == "heading1":
                    # Điều 1., Điều 2., etc.
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.size = Pt(max(font_size, 14))
                    run.font.bold = True
                    run.font.name = font_name
                    if color > 0:
                        rgb = self._hex_to_rgb(color)
                        run.font.color.rgb = RGBColor(*rgb)
                    para.space_before = Pt(6)
                    para.space_after = Pt(3)
                
                elif block_type == "heading2":
                    # Khoản: 1., 2., etc.
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.size = Pt(max(font_size, 13))
                    run.font.bold = True
                    run.font.name = font_name
                    if color > 0:
                        rgb = self._hex_to_rgb(color)
                        run.font.color.rgb = RGBColor(*rgb)
                    para.space_before = Pt(3)
                    para.space_after = Pt(2)
                
                elif block_type == "list_item":
                    # Điểm: a), b), etc.
                    # IMPORTANT: Keep as text, don't convert to Word bullet
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    run.font.bold = is_bold
                    run.font.italic = is_italic
                    if color > 0:
                        rgb = self._hex_to_rgb(color)
                        run.font.color.rgb = RGBColor(*rgb)
                    para.space_after = Pt(2)
                    # Add left indent to show hierarchy
                    para.paragraph_format.left_indent = Inches(0.25)
                
                elif block_type == "paragraph":
                    para = doc.add_paragraph(text)
                    para.space_after = Pt(6)
                    if para.runs:
                        run = para.runs[0]
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                        run.font.bold = is_bold
                        run.font.italic = is_italic
                        if color > 0:
                            rgb = self._hex_to_rgb(color)
                            run.font.color.rgb = RGBColor(*rgb)
                
                elif block_type == "table":
                    rows = text.split('\n')[:30]
                    if rows:
                        first_row_cells = rows[0].split('\t')[:15]
                        num_cols = max(len(first_row_cells), 2)
                        
                        table = doc.add_table(rows=min(len(rows), 30), cols=num_cols)
                        table.style = 'Light Grid Accent 1'
                        
                        for i, row in enumerate(rows[:30]):
                            cells = row.split('\t')[:num_cols]
                            for j, cell_text in enumerate(cells):
                                if j < num_cols:
                                    table.rows[i].cells[j].text = cell_text.strip()
                
                else:  # text
                    para = doc.add_paragraph(text)
                    para.space_after = Pt(3)
                    if para.runs:
                        run = para.runs[0]
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                        run.font.bold = is_bold
                        run.font.italic = is_italic
                        if color > 0:
                            rgb = self._hex_to_rgb(color)
                            run.font.color.rgb = RGBColor(*rgb)
        
        # Add metadata (optional)
        # doc.add_page_break()
        # doc.add_heading('Thông tin xử lý', level=1)
        # metadata = data.get("metadata", {})
        # doc.add_paragraph(f'Loại PDF: {metadata.get("pdf_type", "unknown")}')
        # doc.add_paragraph(f'Tổng số trang: {metadata.get("total_pages", 0)}')
        # doc.add_paragraph(f'Thời gian xử lý: {metadata.get("processing_time", 0):.2f} giây')
        # doc.add_paragraph(f'Layout Engine: {"Có" if metadata.get("layout_engine_used") else "Không"}')
        # doc.add_paragraph(f'Xuất lúc: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        doc.save(str(output_path))
        
        return output_path
    
    def export_to_pdf(
        self,
        data: Dict[str, Any],
        output_path: Path
    ) -> Path:
        """Export to PDF with preserved layout"""
        if not SimpleDocTemplate:
            raise RuntimeError("reportlab not installed. Install with: pip install reportlab")
        
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        story = []
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=12
        )
        
        header_style = ParagraphStyle(
            'CustomHeader',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#2c2c2c'),
            spaceAfter=8
        )
        
        heading1_style = ParagraphStyle(
            'CustomHeading1',
            parent=styles['Heading1'],
            fontSize=14,
            textColor=colors.HexColor('#1a1a1a'),
            spaceBefore=6,
            spaceAfter=3
        )
        
        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontSize=13,
            textColor=colors.HexColor('#2c2c2c'),
            spaceBefore=3,
            spaceAfter=2
        )
        
        para_style = ParagraphStyle(
            'CustomPara',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
            spaceAfter=6,
            alignment=4
        )
        
        for page_data in data.get("pages", []):
            page_num = page_data.get("page_num", 0)
            blocks = page_data.get("blocks", [])
            
            if page_num > 0:
                story.append(PageBreak())
            
            for block in blocks:
                block_type = block.get("type", "text")
                text = block.get("text", "")
                
                if not text.strip() and block_type != "table":
                    continue
                
                if block_type == "title":
                    story.append(Paragraph(self._escape_xml(text), title_style))
                    story.append(Spacer(1, 6))
                
                elif block_type == "header":
                    story.append(Paragraph(self._escape_xml(text), header_style))
                    story.append(Spacer(1, 4))
                
                elif block_type == "heading1":
                    story.append(Paragraph(self._escape_xml(text), heading1_style))
                    story.append(Spacer(1, 3))
                
                elif block_type == "heading2":
                    story.append(Paragraph(self._escape_xml(text), heading2_style))
                    story.append(Spacer(1, 2))
                
                elif block_type == "paragraph":
                    story.append(Paragraph(self._escape_xml(text), para_style))
                    story.append(Spacer(1, 6))
                
                elif block_type == "list_item":
                    story.append(Paragraph(self._escape_xml(text), styles['Normal']))
                    story.append(Spacer(1, 2))
                
                elif block_type == "table":
                    rows = text.split('\n')[:20]
                    if rows:
                        table_data = []
                        for row in rows:
                            cells = row.split('\t')[:10]
                            table_data.append([self._escape_xml(cell.strip()) for cell in cells])
                        
                        if table_data:
                            table = Table(table_data)
                            table.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                ('FONTSIZE', (0, 0), (-1, 0), 10),
                                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                                ('GRID', (0, 0), (-1, -1), 1, colors.black)
                            ]))
                            story.append(table)
                            story.append(Spacer(1, 12))
                
                else:  # text
                    story.append(Paragraph(self._escape_xml(text), styles['Normal']))
                    story.append(Spacer(1, 3))
        
        doc.build(story)
        
        return output_path
    
    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters"""
        return (text
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&#39;'))
    
    def _escape_xml(self, text: str) -> str:
        """Escape XML special characters"""
        return (text
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&#39;'))

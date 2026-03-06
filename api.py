from fastapi import FastAPI, UploadFile, File
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional
import uuid
from pathlib import Path
import shutil
import sys
import os
import re
import threading
import time
from subprocess import Popen, PIPE, STDOUT
import numpy as np

# Import PyMuPDF for PDF detection
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

# Import PaddleOCR for basic OCR (scan PDFs)
try:
    from paddleocr import PaddleOCR
except ImportError:
    PaddleOCR = None

# Global basic OCR model (lazy-loaded)
_basic_ocr_model = None

app = FastAPI()

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = Path(__file__).parent
UPLOAD_DIR = BASE_DIR / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

# ===============================
# In-memory job store
# ===============================
JOB_PROGRESS = {}   # job_id -> {current, total, percent, elapsed_time, start_time, pdf_type}
JOB_STATUS = {}     # job_id -> running | done | error | cancelled


def detect_pdf_type(pdf_path: Path) -> str:
    """Detect if PDF is scanned (image) or text-based"""
    if not fitz:
        # Default to scan if PyMuPDF not available
        print(f"[INFO] PyMuPDF not available, defaulting to scan", flush=True)
        return "scan"
    
    try:
        doc = fitz.open(pdf_path)
        total_pages = doc.page_count
        
        print(f"[INFO] PDF has {total_pages} pages", flush=True)
        
        # For very large files (>50 pages), sample more pages
        sample_pages = min(5, total_pages) if total_pages > 50 else min(3, total_pages)
        
        text_length = 0
        has_images = 0
        hybrid_pages = 0
        total_images = 0
        
        for page_num in range(sample_pages):
            page = doc[page_num]
            text = page.get_text()
            text_len = len(text.strip())
            text_length += text_len
            
            # Check if page has images
            image_list = page.get_images()
            img_count = len(image_list)
            if img_count > 0:
                has_images += 1
                total_images += img_count
            
            # Check if hybrid (both text and images)
            if text_len > 0 and img_count > 0:
                hybrid_pages += 1
        
        doc.close()
        
        # Calculate metrics
        avg_text = text_length / sample_pages if sample_pages > 0 else 0
        image_ratio = has_images / sample_pages if sample_pages > 0 else 0
        hybrid_ratio = hybrid_pages / sample_pages if sample_pages > 0 else 0
        
        print(f"[INFO] PDF Detection - Avg text: {avg_text:.0f} chars/page, Images: {has_images}/{sample_pages} pages, Hybrid: {hybrid_pages}/{sample_pages}", flush=True)
        
        # Decision logic:
        # 1. Very large files (>50 pages) with many images → use OCR (faster)
        # 2. Hybrid PDFs (text + images) → use OCR if >30% hybrid or >50 pages
        # 3. Pure text with many pages (>100) → still use convert_keep_layout
        # 4. Pure text small/medium → use convert_keep_layout
        
        is_large_file = total_pages > 50
        is_very_large = total_pages > 100
        has_many_images = total_images > sample_pages * 2  # Average >2 images per page
        
        if avg_text < 100 and image_ratio > 0.5:
            print(f"[INFO] Detected as: SCAN (low text: {avg_text:.0f}, high images: {image_ratio:.1%})", flush=True)
            return "scan"
        elif avg_text >= 100:
            # Pure text PDF detected
            if is_very_large and (hybrid_ratio > 0.3 or has_many_images):
                # Large hybrid file - OCR might be faster
                print(f"[INFO] Detected as: TEXT (but large hybrid: {total_pages} pages, {hybrid_ratio:.1%} hybrid)", flush=True)
                print(f"[INFO] Recommendation: Using OCR for better performance with large hybrid files", flush=True)
                return "scan"  # Use OCR for large hybrid files
            elif is_large_file and hybrid_ratio > 0.5:
                # Medium-large hybrid - use OCR
                print(f"[INFO] Detected as: HYBRID TEXT (large file: {total_pages} pages, {hybrid_ratio:.1%} hybrid)", flush=True)
                print(f"[INFO] Recommendation: Using OCR for better performance", flush=True)
                return "scan"
            else:
                # Pure or mostly text
                print(f"[INFO] Detected as: TEXT (high text: {avg_text:.0f}, {total_pages} pages)", flush=True)
                return "text"
        else:
            # Default to scan if unclear
            print(f"[INFO] Detected as: SCAN (unclear, defaulting to scan)", flush=True)
            return "scan"
    except Exception as e:
        print(f"[WARN] Error detecting PDF type: {e}, defaulting to scan", flush=True)
        return "scan"


def process_pdf_background(job_id: str, pdf_path: Path, output_docx: Path, pdf_type: str):
    """Process PDF in background thread"""
    start_time = time.time()
    JOB_PROGRESS[job_id]["start_time"] = start_time
    JOB_PROGRESS[job_id]["pdf_type"] = pdf_type
    
    try:
        # ===============================
        # Select script based on PDF type
        # ===============================
        scripts_dir = BASE_DIR / "scripts"
        if pdf_type == "scan":
            script_path = scripts_dir / "convert_pdf_gpu.py"
            print(f"[INFO] Job {job_id}: Using OCR script (convert_pdf_gpu.py) for SCAN PDF", flush=True)
        else:  # text
            script_path = scripts_dir / "convert_keep_layout.py"
            print(f"[INFO] Job {job_id}: Using layout-preserving script (convert_keep_layout.py) for TEXT PDF", flush=True)
        
        # Verify script exists
        if not script_path.exists():
            error_msg = f"Script not found: {script_path}. Please ensure the script exists in the scripts directory."
            print(f"[ERROR] {error_msg}", flush=True)
            raise FileNotFoundError(error_msg)
        
        # Verify PDF file exists
        if not pdf_path.exists():
            error_msg = f"PDF file not found: {pdf_path}"
            print(f"[ERROR] {error_msg}", flush=True)
            raise FileNotFoundError(error_msg)
        
        # Verify Python executable
        if not sys.executable:
            error_msg = "Python executable not found"
            print(f"[ERROR] {error_msg}", flush=True)
            raise RuntimeError(error_msg)
        
        cmd = [
            sys.executable,
            "-u",  # Unbuffered output
            str(script_path),
            str(pdf_path),
            "--output",
            str(output_docx),
        ]
        
        print(f"[INFO] Job {job_id}: Starting conversion with command: {' '.join(cmd)}", flush=True)
        print(f"[INFO] Job {job_id}: Python executable: {sys.executable}", flush=True)
        print(f"[INFO] Job {job_id}: Script path: {script_path}", flush=True)
        print(f"[INFO] Job {job_id}: PDF path: {pdf_path}", flush=True)
        print(f"[INFO] Job {job_id}: Output path: {output_docx}", flush=True)

        # ===============================
        # Regex for OCR phases (only for scan PDFs)
        # ===============================
        pdf_info_pattern = re.compile(r"--- PDF Info: (\d+) pages ---")
        layout_pattern = re.compile(r"Recognizing Layout")
        ocr_error_pattern = re.compile(r"OCR Error Detection")
        bbox_pattern = re.compile(r"Detecting bboxes")
        # Match formats like:
        # - "Recognizing Text: 0% | 0/40 [time]"
        # - "Recognizing Text: 2%|2 | 1/40 [time]"
        # - "Recognizing Text: 18%|#7 | 7/40 [time]"
        # Only match lines that start with "Recognizing Text:" and have progress bar format
        text_pattern = re.compile(r"^Recognizing Text:\s*\d+%.*?(\d+)\s*/\s*(\d+)")
        
        # For text PDFs (convert_keep_layout.py)
        layout_convert_pattern = re.compile(r"--- Starting Layout-Preserving Conversion ---")

        # Set environment to force unbuffered output and UTF-8 on Windows
        env = os.environ.copy()
        env["PYTHONUNBUFFERED"] = "1"
        env["PYTHONIOENCODING"] = "utf-8"

        # Use UTF-8 encoding to avoid "charmap codec can't decode" on Windows
        # when subprocess prints Vietnamese or other non-ASCII characters
        # Run subprocess with project root as cwd so scripts and relative paths (models, etc.) resolve correctly
        process = Popen(
            cmd,
            stdout=PIPE,
            stderr=STDOUT,
            text=True,
            encoding="utf-8",
            errors="replace",  # Replace invalid bytes instead of raising
            bufsize=0,  # Unbuffered
            universal_newlines=True,
            env=env,
            cwd=str(BASE_DIR),
        )

        for line in process.stdout:
            print(line, end="", flush=True)
            
            # Skip if job is already done
            if JOB_STATUS.get(job_id) == "done":
                continue

            # Update elapsed time
            elapsed = time.time() - start_time
            JOB_PROGRESS[job_id]["elapsed_time"] = elapsed

            if pdf_type == "scan":
                # ----- For SCAN PDFs: Parse OCR progress -----
                # Parse PDF info at start
                m_info = pdf_info_pattern.search(line)
                if m_info:
                    total_pages = int(m_info.group(1))
                    # Set total from PDF Info - this is the source of truth
                    JOB_PROGRESS[job_id]["total"] = total_pages
                    JOB_PROGRESS[job_id]["percent"] = 1
                    continue

                # Phase-based progress
                if layout_pattern.search(line):
                    JOB_PROGRESS[job_id]["percent"] = 5

                elif ocr_error_pattern.search(line):
                    JOB_PROGRESS[job_id]["percent"] = 10

                elif bbox_pattern.search(line):
                    JOB_PROGRESS[job_id]["percent"] = 15

                # Real page-based progress
                # Only process if line contains "Recognizing Text:" to avoid false matches
                if "Recognizing Text:" in line:
                    m = text_pattern.search(line.strip())  # Strip to handle line breaks
                    if m:
                        current = int(m.group(1))
                        total_from_log = int(m.group(2))
                        
                        # ALWAYS use total from PDF Info if available (source of truth)
                        current_total = JOB_PROGRESS[job_id].get("total", 0)
                        if current_total > 0:
                            total = current_total
                            current = min(current, total)
                        else:
                            total = total_from_log
                            JOB_PROGRESS[job_id]["total"] = total
                            print(f"[INFO] Job {job_id}: Using total from log: {total}", flush=True)
                        
                        # Progress: 15% (initial phases) + up to 85% based on pages
                        if total > 0:
                            percent = 15 + int((current / total) * 85)
                        else:
                            percent = 15

                        JOB_PROGRESS[job_id].update({
                            "current": current,
                            "percent": min(percent, 99)
                        })
                        print(f"[PROGRESS] Job {job_id}: {current}/{total} pages ({percent}%)", flush=True)
            
            else:
                # ----- For TEXT PDFs: Simple progress tracking -----
                # Parse PDF info if available (from convert_keep_layout.py)
                m_info = pdf_info_pattern.search(line)
                if m_info:
                    total_pages = int(m_info.group(1))
                    JOB_PROGRESS[job_id]["total"] = total_pages
                    JOB_PROGRESS[job_id]["percent"] = 20
                    JOB_PROGRESS[job_id]["current"] = 0
                    print(f"[PROGRESS] Job {job_id}: Text PDF - {total_pages} pages detected, starting conversion", flush=True)
                
                if layout_convert_pattern.search(line):
                    # Progress to 30% when conversion starts
                    JOB_PROGRESS[job_id]["percent"] = 30
                    print(f"[PROGRESS] Job {job_id}: Text PDF conversion started", flush=True)
                
                # For text PDFs, gradually increase progress based on time
                # This runs for every log line, so we update progress smoothly
                elapsed = time.time() - start_time
                current_percent = JOB_PROGRESS[job_id].get("percent", 0)
                
                # Only update if we have total pages and percent is reasonable
                total_pages = JOB_PROGRESS[job_id].get("total", 0)
                if total_pages > 0:
                    # For text PDFs, pdf2docx can be slow for large files
                    # Estimate: ~1-3 seconds per page depending on complexity
                    estimated_time_per_page = 2.0  # seconds per page (conservative for large files)
                    estimated_total_time = total_pages * estimated_time_per_page
                    
                    # Progress: 30% (start) + 65% (during conversion) + 5% (finalize)
                    if elapsed > 0 and current_percent < 95:
                        # Use logarithmic scale for progress (slower at end)
                        if elapsed < estimated_total_time:
                            progress_estimate = 30 + int((elapsed / max(estimated_total_time, 1)) * 65)
                        else:
                            # If over estimated time, show 95% (almost done, finalizing)
                            progress_estimate = 95
                        
                        progress_estimate = min(progress_estimate, 95)  # Cap at 95% until done
                        if progress_estimate > current_percent:
                            JOB_PROGRESS[job_id]["percent"] = progress_estimate
                            print(f"[PROGRESS] Job {job_id}: Text PDF progress {progress_estimate}% (elapsed: {elapsed:.1f}s, estimated: {estimated_total_time:.1f}s)", flush=True)
                elif current_percent < 30:
                    # If no total yet, at least show we're starting
                    JOB_PROGRESS[job_id]["percent"] = 20

        process.wait()

        if process.returncode == 0:
            elapsed_total = time.time() - start_time
            JOB_PROGRESS[job_id].update({
                "percent": 100,
                "elapsed_time": elapsed_total
            })
            JOB_STATUS[job_id] = "done"
            print(f"[SUCCESS] Job {job_id} completed in {elapsed_total:.2f} seconds", flush=True)
        else:
            JOB_STATUS[job_id] = "error"
            error_msg = f"Chuyển đổi thất bại (mã thoát: {process.returncode}). Kiểm tra terminal backend để xem log chi tiết."
            print(f"[ERROR] Job {job_id} failed with return code {process.returncode}", flush=True)
            JOB_PROGRESS[job_id]["error"] = error_msg
            # Try to capture stderr if available
            if hasattr(process, 'stderr') and process.stderr:
                stderr_output = process.stderr.read()
                if stderr_output:
                    print(f"[ERROR] Stderr output: {stderr_output}", flush=True)

    except FileNotFoundError as e:
        error_msg = f"Script not found: {e}"
        print(f"[ERROR] {error_msg}", flush=True)
        JOB_STATUS[job_id] = "error"
        JOB_PROGRESS[job_id]["error"] = error_msg
    except Exception as e:
        error_msg = f"OCR ERROR for job {job_id}: {e}"
        print(f"[ERROR] {error_msg}", flush=True)
        import traceback
        print(f"[ERROR] Traceback: {traceback.format_exc()}", flush=True)
        JOB_STATUS[job_id] = "error"
        JOB_PROGRESS[job_id]["error"] = str(e)


@app.post("/convert")
async def convert_pdf(file: UploadFile = File(...)):
    job_id = str(uuid.uuid4())

    pdf_path = UPLOAD_DIR / f"{job_id}.pdf"
    output_docx = UPLOAD_DIR / f"{job_id}.docx"

    # Save uploaded PDF
    with open(pdf_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    # Detect PDF type
    detection_start = time.time()
    pdf_type = detect_pdf_type(pdf_path)
    pdf_type_display = "scan" if pdf_type == "scan" else "text"
    detection_time = time.time() - detection_start
    
    print(f"[INFO] Job {job_id}: Detected PDF type: {pdf_type_display} (took {detection_time:.2f}s)", flush=True)
    print(f"[INFO] Job {job_id}: File will be processed with {'OCR (convert_pdf_gpu.py)' if pdf_type == 'scan' else 'Direct conversion (convert_keep_layout.py)'}", flush=True)

    # Initialize job status
    JOB_PROGRESS[job_id] = {
        "current": 0,
        "total": 0,
        "percent": 0,
        "elapsed_time": 0,
        "start_time": 0,
        "pdf_type": pdf_type_display
    }
    JOB_STATUS[job_id] = "running"

    # Start processing in background thread
    thread = threading.Thread(
        target=process_pdf_background,
        args=(job_id, pdf_path, output_docx, pdf_type),
        daemon=True
    )
    thread.start()

    # Return job_id immediately so frontend can start polling
    return {"job_id": job_id, "pdf_type": pdf_type_display}


@app.post("/convert_basic")
async def convert_basic_pdf(file: UploadFile = File(...)):
    """
    Basic OCR/Text extraction:
    - Chỉ trích xuất text, không giữ bố cục.
    - Hỗ trợ cả PDF text và PDF scan.
      + PDF text: dùng cùng luồng xử lý PDF text như OCR nâng cao (script convert_keep_layout.py), sau đó rút text.
      + PDF scan: dùng pipeline riêng `ocr_basic.py` (YOLO + VietOCR), KHÔNG dùng convert_pdf_gpu.py và KHÔNG fallback sang OCR nâng cao.
    """
    temp_id = str(uuid.uuid4())
    pdf_path = UPLOAD_DIR / f"basic_{temp_id}.pdf"
    output_docx = UPLOAD_DIR / f"basic_{temp_id}.docx"

    # Save uploaded PDF
    with open(pdf_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    # Detect PDF type (scan / text)
    pdf_type = detect_pdf_type(pdf_path)

    # ===== Chọn luồng xử lý theo loại PDF =====
    if pdf_type == "scan":
        # PDF scan: dùng pipeline riêng ocr_basic.py (YOLO + VietOCR), KHÔNG gọi convert_pdf_gpu.py
        script_path = BASE_DIR / "ocr_basic.py"
        if not script_path.exists():
            return JSONResponse(
                status_code=500,
                content={
                    "error": f"Không tìm thấy script OCR cơ bản cho PDF scan: {script_path}. Vui lòng kiểm tra file ocr_basic.py."
                },
            )

        temp_layout_docx = UPLOAD_DIR / f"basic_{temp_id}_ocr_basic.docx"

        cmd = [
            sys.executable,
            "-u",
            str(script_path),
            str(pdf_path),
            "--output",
            str(temp_layout_docx),
        ]

        print(f"[INFO] Basic OCR (scan) - running command: {' '.join(cmd)}", flush=True)

        env = os.environ.copy()
        env["PYTHONUNBUFFERED"] = "1"
        env["PYTHONIOENCODING"] = "utf-8"

        process = Popen(
            cmd,
            stdout=PIPE,
            stderr=STDOUT,
            text=True,
            encoding="utf-8",
            errors="replace",
            bufsize=0,
            universal_newlines=True,
            env=env,
            cwd=str(BASE_DIR),
        )

        # Log đơn giản, không cập nhật progress cho frontend
        for line in process.stdout:
            print(line, end="", flush=True)

        process.wait()

        if process.returncode != 0 or not temp_layout_docx.exists():
            print(f"[ERROR] Basic OCR (scan) via ocr_basic.py failed, return code: {process.returncode}", flush=True)
            return JSONResponse(
                status_code=500,
                content={
                    "error": "OCR cơ bản cho PDF scan thất bại (ocr_basic.py). Vui lòng thử lại hoặc dùng chế độ OCR nâng cao."
                },
            )

        # Lấy text thuần từ DOCX bố cục do ocr_basic.py tạo ra
        text = extract_text_from_docx(temp_layout_docx)
    else:
        # PDF text: dùng cùng script như OCR nâng cao (convert_keep_layout.py), sau đó rút text thuần
        scripts_dir = BASE_DIR / "scripts"
        script_path = scripts_dir / "convert_keep_layout.py"
        if not script_path.exists():
            return JSONResponse(
                status_code=500,
                content={
                    "error": f"Không tìm thấy script xử lý PDF text: {script_path}. Vui lòng kiểm tra thư mục scripts."
                },
            )

        temp_layout_docx = UPLOAD_DIR / f"basic_{temp_id}_layout.docx"

        cmd = [
            sys.executable,
            "-u",
            str(script_path),
            str(pdf_path),
            "--output",
            str(temp_layout_docx),
        ]

        print(f"[INFO] Basic OCR (text) - running command: {' '.join(cmd)}", flush=True)

        env = os.environ.copy()
        env["PYTHONUNBUFFERED"] = "1"
        env["PYTHONIOENCODING"] = "utf-8"

        process = Popen(
            cmd,
            stdout=PIPE,
            stderr=STDOUT,
            text=True,
            encoding="utf-8",
            errors="replace",
            bufsize=0,
            universal_newlines=True,
            env=env,
            cwd=str(BASE_DIR),
        )

        # Log đơn giản, không cập nhật progress cho frontend
        for line in process.stdout:
            print(line, end="", flush=True)

        process.wait()

        if process.returncode != 0 or not temp_layout_docx.exists():
            print(f"[ERROR] Basic OCR (text) failed, return code: {process.returncode}", flush=True)
            return JSONResponse(
                status_code=500,
                content={
                    "error": "OCR cơ bản cho PDF văn bản (text) thất bại. Vui lòng thử lại hoặc dùng chế độ OCR nâng cao."
                },
            )

        # Lấy text thuần từ DOCX bố cục
        text = extract_text_from_docx(temp_layout_docx)

    if not text or not text.strip():
        return JSONResponse(
            status_code=400,
            content={
                "error": "Không trích xuất được văn bản từ PDF với chế độ OCR cơ bản. Vui lòng sử dụng chế độ OCR nâng cao."
            },
        )

    try:
        create_plain_docx_from_text(text, output_docx)
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Lỗi khi tạo file DOCX đơn giản: {e}"},
        )

    filename_stem = Path(file.filename).stem if file.filename else "result"

    return FileResponse(
        output_docx,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{filename_stem}_basic.docx",
    )


@app.get("/progress/{job_id}")
def get_progress(job_id: str):
    if job_id not in JOB_STATUS:
        return JSONResponse(
            status_code=404,
            content={"error": "Job not found"}
        )

    progress_data = JOB_PROGRESS.get(job_id, {
        "current": 0, 
        "total": 0, 
        "percent": 0, 
        "elapsed_time": 0,
        "start_time": 0
    })
    
    # Update elapsed time if still running
    if JOB_STATUS[job_id] == "running" and progress_data.get("start_time"):
        progress_data["elapsed_time"] = time.time() - progress_data["start_time"]
    
    return {
        "status": JOB_STATUS[job_id],
        **progress_data
    }


@app.get("/result/{job_id}")
def get_result(job_id: str):
    output_docx = UPLOAD_DIR / f"{job_id}.docx"

    if not output_docx.exists():
        return JSONResponse(
            status_code=404,
            content={"error": "Result not ready"}
        )

    return FileResponse(
        output_docx,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="result.docx"
    )


@app.post("/cancel/{job_id}")
def cancel_job(job_id: str):
    """Đánh dấu hủy job phía backend (không chờ subprocess kết thúc)."""
    if job_id not in JOB_STATUS:
        return JSONResponse(status_code=404, content={"error": "Job not found"})

    JOB_STATUS[job_id] = "cancelled"
    progress_data = JOB_PROGRESS.get(job_id)
    if progress_data is not None:
        progress_data["status"] = "cancelled"

    return {"status": "cancelled"}


def extract_text_from_pdf(pdf_path: Path) -> str:
    """Extract text from PDF file"""
    try:
        if fitz:
            doc = fitz.open(pdf_path)
            text_parts = []
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text = page.get_text()
                text_parts.append(text)
            doc.close()
            return "\n\n".join(text_parts)
        else:
            return ""
    except Exception as e:
        print(f"[ERROR] Failed to extract text from PDF: {e}", flush=True)
        return ""


def create_plain_docx_from_text(text: str, output_path: Path) -> None:
    """Create a simple DOCX containing only plain text paragraphs."""
    try:
        from docx import Document  # Imported here to avoid global import if not needed

        doc = Document()
        # Split text into lines, keep empty lines as paragraph breaks
        for line in text.splitlines():
            doc.add_paragraph(line)
        doc.save(output_path)
    except Exception as e:
        print(f"[ERROR] Failed to create plain DOCX: {e}", flush=True)
        raise


def extract_text_from_docx(docx_path: Path) -> str:
    """Extract text from DOCX file"""
    try:
        from docx import Document
        doc = Document(docx_path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        return "\n\n".join(text_parts)
    except Exception as e:
        print(f"[ERROR] Failed to extract text from DOCX: {e}", flush=True)
        return ""


def get_basic_ocr() -> Optional["PaddleOCR"]:
    """Lazy-load và trả về instance PaddleOCR dùng cho OCR cơ bản (scan)."""
    global _basic_ocr_model

    if _basic_ocr_model is not None:
        return _basic_ocr_model

    if PaddleOCR is None:
        print("[WARN] PaddleOCR chưa được cài đặt. Vui lòng `pip install paddleocr` nếu muốn dùng OCR cơ bản cho PDF scan.", flush=True)
        return None

    try:
        # Một số phiên bản PaddleOCR không hỗ trợ tham số use_gpu, nên chỉ dùng tham số chung.
        print(f"[INFO] Khởi tạo PaddleOCR cho OCR cơ bản (lang='vi')", flush=True)
        _basic_ocr_model = PaddleOCR(
            lang="vi",
            use_angle_cls=True,
        )
        return _basic_ocr_model
    except Exception as e:
        print(f"[ERROR] Không khởi tạo được PaddleOCR cho OCR cơ bản: {e}", flush=True)
        _basic_ocr_model = None
        return None


def ocr_basic_scan_pdf_to_text(pdf_path: Path, temp_prefix: str = "basic_scan") -> str:
    """
    OCR PDF scan bằng PaddleOCR, chỉ lấy text, không giữ bố cục.

    - Render từng trang PDF sang ảnh (dpi vừa phải để cân bằng tốc độ/chất lượng).
    - Chạy PaddleOCR trên từng ảnh.
    - Ghép tất cả dòng text lại thành 1 chuỗi.

    Lưu ý: Nếu không có fitz hoặc PaddleOCR → trả về chuỗi rỗng.
    """
    if fitz is None:
        print("[WARN] PyMuPDF (fitz) chưa được cài đặt, không thể render PDF sang ảnh cho OCR cơ bản.", flush=True)
        return ""

    ocr = get_basic_ocr()
    if ocr is None:
        print("[WARN] PaddleOCR không sẵn sàng, không thể OCR PDF scan ở chế độ cơ bản.", flush=True)
        return ""

    texts: List[str] = []

    try:
        doc = fitz.open(pdf_path)
    except Exception as e:
        print(f"[ERROR] Không mở được PDF cho OCR cơ bản: {e}", flush=True)
        return ""

    try:
        total_pages = doc.page_count
        print(f"[INFO] Basic OCR (scan) - tổng số trang: {total_pages}", flush=True)

        for page_index in range(total_pages):
            page = doc[page_index]
            # Dùng dpi vừa phải để tránh quá chậm (150–180 là mức hợp lý)
            try:
                pix = page.get_pixmap(dpi=160)
            except TypeError:
                # Fallback nếu phiên bản fitz không hỗ trợ tham số dpi
                pix = page.get_pixmap()

            image_filename = f"{temp_prefix}_{uuid.uuid4().hex}_page{page_index+1}.png"
            image_path = UPLOAD_DIR / image_filename

            try:
                pix.save(str(image_path))
            except Exception as e:
                print(f"[WARN] Không lưu được ảnh trang {page_index+1} cho OCR cơ bản: {e}", flush=True)
                continue

            try:
                # PaddleOCR trả về list các dòng; mỗi dòng là list các bbox + info
                result = ocr.ocr(str(image_path), cls=True)
                if not result:
                    continue

                for line in result:
                    for item in line:
                        if len(item) < 2:
                            continue
                        info = item[1]
                        # info có thể là tuple/list (text, score) hoặc string tuỳ phiên bản
                        if isinstance(info, (list, tuple)) and len(info) >= 1:
                            txt = str(info[0])
                        else:
                            txt = str(info)

                        txt = txt.strip()
                        if txt:
                            texts.append(txt)
            except Exception as e:
                print(f"[WARN] Lỗi OCR trang {page_index+1} ở chế độ cơ bản: {e}", flush=True)
            finally:
                # Xoá file ảnh tạm để tránh đầy ổ đĩa
                try:
                    if image_path.exists():
                        image_path.unlink()
                except Exception as e:
                    print(f"[WARN] Không xoá được ảnh tạm {image_path}: {e}", flush=True)

        doc.close()
    except Exception as e:
        print(f"[ERROR] Lỗi khi OCR basic PDF scan: {e}", flush=True)
        try:
            doc.close()
        except Exception:
            pass

    combined_text = "\n".join(texts)
    print(f"[INFO] Basic OCR (scan) - tổng số dòng text trích xuất: {len(texts)}", flush=True)
    return combined_text



"""
Document Parser - Step 1: Ingestion & Structural Parsing

Parse Word/PDF files into structured elements with unique IDs.
Each paragraph and table cell gets a unique ID for later modification.

CRITICAL FIX (v2):
  Previously, paragraph.text (python-docx) was used to extract text.
  python-docx's .text property includes text from anchored textboxes/frames
  embedded inside the paragraph XML — but doc_surgery._collect_text_segments()
  correctly EXCLUDES textbox content (w:txbxContent).

  This mismatch caused layout corruption:
    - LLM sees text T1 (paragraph.text = runs + textbox text)
    - Surgery finds text T2 (runs only, no textbox text)
    - diff-based replacement computes wrong character positions
    - Wrong runs get modified → text lands in wrong alignment/position

  Fix: use lxml directly (same as doc_surgery) to extract ONLY run text,
  ensuring the ID→content mapping is 100% consistent with what the
  surgery will find and modify.

  Additionally fixed: doc.paragraphs (python-docx) includes paragraphs
  inside table cells, causing index drift when using para_idx as an
  index into doc.paragraphs after any table in the document.
"""
import os
import uuid
import zipfile
import shutil
from pathlib import Path
from dataclasses import dataclass, field, asdict
from typing import Optional, Tuple

from lxml import etree
from docx import Document

WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
XML_NS  = 'http://www.w3.org/XML/1998/namespace'


@dataclass
class DocumentElement:
    """A single element (paragraph or table cell) in a document."""
    id: str              # "Para_0", "Table_0_Cell_1_2"
    type: str            # "paragraph" | "table_cell"
    content: str         # text content (runs only, NO textbox text)
    table_id: Optional[str] = None
    row: Optional[int] = None
    col: Optional[int] = None
    metadata: dict = field(default_factory=dict)
    context: str = ""

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass
class DocumentStructure:
    """Complete parsed structure of a document."""
    doc_id: str
    file_name: str
    docx_path: str
    elements: list[DocumentElement] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "doc_id": self.doc_id,
            "file_name": self.file_name,
            "docx_path": self.docx_path,
            "elements": [e.to_dict() for e in self.elements],
        }

    def get_element_by_id(self, element_id: str) -> Optional[DocumentElement]:
        for e in self.elements:
            if e.id == element_id:
                return e
        return None

    def get_full_text(self) -> str:
        """Get full document text with element IDs for LLM context."""
        lines = []
        for e in self.elements:
            if e.type == "paragraph":
                lines.append(f"[{e.id}] {e.content}")
            elif e.type == "table_cell":
                lines.append(f"[{e.id}] (Table: {e.table_id}, Row {e.row}, Col {e.col}): {e.content}")
        return "\n".join(lines)

    def get_summary_text(self) -> str:
        para_count = sum(1 for e in self.elements if e.type == "paragraph")
        table_ids = set(e.table_id for e in self.elements if e.type == "table_cell" and e.table_id)
        return (
            f"Document: {self.file_name}\n"
            f"Paragraphs: {para_count}\n"
            f"Tables: {len(table_ids)}\n"
            f"Total elements: {len(self.elements)}"
        )

    def build_context_windows(self, window: int = 1):
        for i, element in enumerate(self.elements):
            parts = []
            start = max(0, i - window)
            for j in range(start, i):
                prev = self.elements[j]
                parts.append(f"[{prev.id}] {prev.content}")
            parts.append(f">>> [{element.id}] {element.content}")
            end = min(len(self.elements), i + window + 1)
            for j in range(i + 1, end):
                nxt = self.elements[j]
                parts.append(f"[{nxt.id}] {nxt.content}")
            element.context = "\n".join(parts)


# ─────────────────────────────────────────────────────────────
# KEY FIX: Extract text from XML runs only (mirrors doc_surgery)
# ─────────────────────────────────────────────────────────────

def _get_paragraph_text_from_runs(p_elem) -> str:
    """
    Extract text ONLY from <w:r> runs and <w:hyperlink> children —
    exactly mirroring doc_surgery._collect_text_segments().

    This deliberately EXCLUDES <w:txbxContent> (textboxes/shapes),
    so the text shown to the LLM exactly matches what doc_surgery
    will find and modify via its segment-based replacement.

    Using python-docx's paragraph.text instead would include textbox
    text, causing the diff-based replacement to compute wrong positions
    and corrupt the document layout.
    """
    parts = []
    for child in p_elem:
        local = etree.QName(child.tag).localname
        if local == 'r':
            for t in child.findall(f'{{{WORD_NS}}}t'):
                parts.append(t.text or '')
        elif local == 'hyperlink':
            for r in child.findall(f'{{{WORD_NS}}}r'):
                for t in r.findall(f'{{{WORD_NS}}}t'):
                    parts.append(t.text or '')
        # Deliberately skip: 'bookmarkStart', 'bookmarkEnd', 'ins', 'del',
        # 'sdt', 'smartTag', 'pPr' (paragraph properties), and especially
        # anything containing txbxContent (textboxes anchored in this para).
    return ''.join(parts)


def _get_cell_text_from_runs(tc_elem) -> str:
    """
    Extract text from all paragraphs within a table cell using run-only logic.
    Joins multiple paragraphs with newline.
    """
    parts = []
    for p in tc_elem.findall(f'{{{WORD_NS}}}p'):
        para_text = _get_paragraph_text_from_runs(p)
        if para_text:
            parts.append(para_text)
    return '\n'.join(parts)


# ─────────────────────────────────────────────────────────────
# Metadata extraction (formatting info, for reference only)
# ─────────────────────────────────────────────────────────────

def _extract_run_metadata_from_xml(p_elem) -> dict:
    """
    Extract formatting metadata directly from the paragraph XML element.
    Avoids using doc.paragraphs[] which has index drift issues after tables.
    """
    meta = {
        "font_name": None,
        "bold": False,
        "italic": False,
        "underline": False,
        "font_size": None,
        "alignment": None,
    }

    # Alignment from paragraph properties
    pPr = p_elem.find(f'{{{WORD_NS}}}pPr')
    if pPr is not None:
        jc = pPr.find(f'{{{WORD_NS}}}jc')
        if jc is not None:
            meta["alignment"] = jc.get(f'{{{WORD_NS}}}val')

    # Font/style from first non-empty run
    for r in p_elem.findall(f'{{{WORD_NS}}}r'):
        t_elems = r.findall(f'{{{WORD_NS}}}t')
        run_text = ''.join(t.text or '' for t in t_elems)
        if not run_text.strip():
            continue

        rPr = r.find(f'{{{WORD_NS}}}rPr')
        if rPr is None:
            break

        rFonts = rPr.find(f'{{{WORD_NS}}}rFonts')
        if rFonts is not None:
            meta["font_name"] = (
                rFonts.get(f'{{{WORD_NS}}}ascii') or
                rFonts.get(f'{{{WORD_NS}}}eastAsia') or
                rFonts.get(f'{{{WORD_NS}}}hAnsi')
            )

        bold = rPr.find(f'{{{WORD_NS}}}b')
        meta["bold"] = bold is not None and bold.get(f'{{{WORD_NS}}}val') != '0'

        italic = rPr.find(f'{{{WORD_NS}}}i')
        meta["italic"] = italic is not None and italic.get(f'{{{WORD_NS}}}val') != '0'

        underline = rPr.find(f'{{{WORD_NS}}}u')
        meta["underline"] = underline is not None and underline.get(f'{{{WORD_NS}}}val') not in (None, 'none')

        sz = rPr.find(f'{{{WORD_NS}}}sz')
        if sz is not None:
            val = sz.get(f'{{{WORD_NS}}}val')
            if val:
                meta["font_size"] = str(int(val) / 2)  # half-points → points

        break

    return meta


# ─────────────────────────────────────────────────────────────
# Main parse function
# ─────────────────────────────────────────────────────────────

def parse_docx(docx_path: str, doc_id: str = None, file_name: str = None) -> DocumentStructure:
    """
    Parse a .docx file into a DocumentStructure with unique element IDs.

    Uses lxml directly on the raw document.xml to extract paragraph text,
    ensuring the ID→content mapping is 100% consistent with doc_surgery.

    Element IDs are assigned based on the element's position in the body XML:
      - Para_N  → Nth <w:p> direct child of <w:body> (0-indexed)
      - Table_T_Cell_R_C → cell at row R, col C in the Tth <w:tbl>

    Only elements with non-empty run-text are included (empty paragraphs,
    structural spacers, and paragraphs whose only text is in textboxes are
    excluded from the LLM context but their XML positions are still counted).
    """
    docx_path = Path(docx_path).resolve()
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    if doc_id is None:
        doc_id = str(uuid.uuid4())
    if file_name is None:
        file_name = docx_path.name

    # ── Read and parse document.xml directly with lxml ──────────────────
    with zipfile.ZipFile(str(docx_path), 'r') as z:
        doc_xml = z.read('word/document.xml')
    root = etree.fromstring(doc_xml)

    body = root.find(f'{{{WORD_NS}}}body')
    if body is None:
        raise ValueError("No <w:body> found in document.xml")

    elements = []
    para_idx  = 0  # counts ALL direct <w:p> children of <w:body>
    table_idx = 0  # counts ALL direct <w:tbl> children of <w:body>

    for block in body:
        local = etree.QName(block.tag).localname

        if local == 'p':
            # ── Body-level paragraph ────────────────────────────────────
            text = _get_paragraph_text_from_runs(block).strip()

            if text:
                element = DocumentElement(
                    id=f"Para_{para_idx}",
                    type="paragraph",
                    content=text,
                    metadata=_extract_run_metadata_from_xml(block),
                )
                elements.append(element)
                # (para_idx is still incremented below even for empty paras)

            para_idx += 1

        elif local == 'tbl':
            # ── Table ───────────────────────────────────────────────────
            table_id = f"Table_{table_idx}"

            rows = block.findall(f'{{{WORD_NS}}}tr')
            for row_idx, row in enumerate(rows):
                cells = row.findall(f'{{{WORD_NS}}}tc')
                for col_idx, cell in enumerate(cells):
                    cell_text = _get_cell_text_from_runs(cell).strip()
                    element = DocumentElement(
                        id=f"{table_id}_Cell_{row_idx}_{col_idx}",
                        type="table_cell",
                        content=cell_text,
                        table_id=table_id,
                        row=row_idx,
                        col=col_idx,
                        metadata={},
                    )
                    elements.append(element)

            table_idx += 1

        # sectPr and other body-level elements are intentionally ignored.

    doc_structure = DocumentStructure(
        doc_id=doc_id,
        file_name=file_name,
        docx_path=str(docx_path),
        elements=elements,
    )
    doc_structure.build_context_windows(window=1)
    return doc_structure


# ─────────────────────────────────────────────────────────────
# PDF conversion helpers (unchanged)
# ─────────────────────────────────────────────────────────────

def _env_truthy(name: str, default: str = "false") -> bool:
    v = os.environ.get(name, default)
    return str(v).strip().lower() in ("1", "true", "yes", "on")


def _ocr_quantization_from_env() -> Tuple[bool, bool]:
    """
    Qwen2.5-VL VRAM control for scanned PDFs (AnythingLLM bridge / ingest).

    Env:
      PDF_PIPELINE_OCR_LOAD_4BIT — default false (same as before env support;
        avoids Windows bitsandbytes/CUDA access violations on some setups).
      PDF_PIPELINE_OCR_LOAD_8BIT — default false; if true, 4bit is disabled.
    For low VRAM, set PDF_PIPELINE_OCR_LOAD_4BIT=true (start-pdf-extract-bridge.bat does this).
    """
    load_8bit = _env_truthy("PDF_PIPELINE_OCR_LOAD_8BIT")
    load_4bit = (not load_8bit) and _env_truthy(
        "PDF_PIPELINE_OCR_LOAD_4BIT", "false"
    )
    return load_4bit, load_8bit


def convert_pdf_to_docx(pdf_path: str, output_dir: str = None) -> str:
    """
    Convert a PDF file to DOCX using the advanced auto_process_pdf pipeline.
    Automatically detects if PDF is scanned/image-based and routes to YOLO+OCR,
    otherwise uses fast PyMuPDF extraction.
    """
    pdf_path = Path(pdf_path).resolve()
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")

    if output_dir is None:
        output_dir = pdf_path.parent
    else:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

    docx_path = output_dir / f"{pdf_path.stem}_converted.docx"

    # Add src directory to sys.path to resolve pdf_processing module imports
    import sys
    src_dir = str(Path(__file__).resolve().parent.parent)
    if src_dir not in sys.path:
        sys.path.insert(0, src_dir)

    try:
        from pdf_processing.auto_process_pdf import analyze_pdf, convert_regular_pdf
        from pdf_processing.processs_pdf_to_docs import process_pdf_to_docx

        print(f"[{pdf_path.name}] Auto-detecting PDF type...")
        text_ratio, _, _ = analyze_pdf(pdf_path)

        if text_ratio > 0.8:
            print("  ✓ Classification: TEXT-BASED PDF -> Using PyMuPDF fast extraction")
            convert_regular_pdf(pdf_path, docx_path)
        else:
            print("  ✓ Classification: SCANNED/HYBRID PDF -> Using advanced OCR pipeline (YOLO + Qwen2.5-VL)")
            load_4bit, load_8bit = _ocr_quantization_from_env()
            process_pdf_to_docx(
                pdf_path=str(pdf_path),
                output_docx=str(docx_path),
                enable_ocr=True,
                load_4bit=load_4bit,
                load_8bit=load_8bit,
            )
            print(f"  ✓ Advanced OCR pipeline finished: {docx_path}")
    except Exception as e:
        import traceback
        print(f"⚠️ Error using advanced PDF pipeline: {e}")
        traceback.print_exc()
        print("Falling back to basic PyMuPDF text extraction...")
        _convert_pdf_with_pymupdf(str(pdf_path), str(docx_path))

    return str(docx_path)


def _convert_pdf_with_pymupdf(pdf_path: str, docx_path: str):
    """Fallback: simple text extraction using PyMuPDF."""
    import fitz
    from docx import Document
    from docx.shared import Cm

    doc = fitz.open(pdf_path)
    docx = Document()

    for section in docx.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    for page_num in range(len(doc)):
        page = doc[page_num]
        blocks = page.get_text("blocks")
        blocks.sort(key=lambda b: (b[1], b[0]))

        for b in blocks:
            if len(b) >= 7 and b[6] == 0:
                text = b[4].strip()
                if text:
                    docx.add_paragraph(text)

        if page_num < len(doc) - 1:
            docx.add_page_break()

    docx.save(docx_path)
    doc.close()
    print(f"✓ Converted PDF to DOCX (PyMuPDF fallback): {docx_path}")


def ingest_file(file_path: str, upload_dir: str = None) -> DocumentStructure:
    """
    Main entry point: ingest a PDF or DOCX file.

    1. If PDF → convert to DOCX first
    2. Parse DOCX → structured elements with IDs (using lxml, run-text only)
    3. Return DocumentStructure
    """
    file_path = Path(file_path).resolve()
    suffix = file_path.suffix.lower()
    doc_id = str(uuid.uuid4())

    if suffix == '.pdf':
        docx_path = convert_pdf_to_docx(str(file_path), output_dir=upload_dir)
        return parse_docx(docx_path, doc_id=doc_id, file_name=file_path.name)

    elif suffix in ('.docx', '.doc'):
        if upload_dir:
            dest = Path(upload_dir) / file_path.name
            if not dest.exists():
                shutil.copy2(file_path, dest)
            docx_path = str(dest)
        else:
            docx_path = str(file_path)
        return parse_docx(docx_path, doc_id=doc_id, file_name=file_path.name)

    else:
        raise ValueError(f"Unsupported file type: {suffix}. Only .pdf, .docx, .doc are supported.")
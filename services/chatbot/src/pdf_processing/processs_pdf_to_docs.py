"""
Complete PDF processing pipeline: PDF → Single DOCX file
Combines all steps: Detection → OCR → Reconstruction → Word Export
"""
import sys
import argparse
import re
from pathlib import Path

for _d in Path(__file__).resolve().parents:
    if (_d / "ocr_app").is_dir() and (_d / "package.json").is_file():
        _rs = str(_d)
        if _rs not in sys.path:
            sys.path.insert(0, _rs)
        break

from repo_layout import find_monorepo_root, resolve_yolo_weights
from docx import Document
from docx.shared import Pt, Cm

# Add the 'src' directory to sys.path to allow running as script directly
src_path = str(Path(__file__).resolve().parent.parent)
if src_path not in sys.path:
    sys.path.insert(0, src_path)
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from pdf_processing.yolo_detect import (
    pdf_to_images, detect_bboxes, 
    crop_bbox
)
from doclayout_yolo import YOLOv10
import torch
import cv2
import numpy as np
from PIL import Image
from transformers import Qwen2_5_VLForConditionalGeneration, AutoProcessor
from qwen_vl_utils import process_vision_info
import markdown
from htmldocx import HtmlToDocx


def determine_alignment_by_position(bbox: dict, page_width: float, page_height: float = None):
    """
    Phân tích vị trí để xác định alignment như con người nhìn vào văn bản
    Dựa vào vị trí X (horizontal) để quyết định căn trái/giữa/phải
    """
    x1 = bbox['x1']
    x2 = bbox['x2']
    center_x = bbox['center_x']
    y1 = bbox['y1']
    
    page_center = page_width / 2
    left_margin_threshold = page_width * 0.25  # 25% từ trái
    right_margin_threshold = page_width * 0.75  # 75% từ trái
    center_tolerance = page_width * 0.15  # 15% tolerance cho center
    
    # Phân tích vị trí X để xác định alignment
    # Nếu bbox bắt đầu gần lề trái và kết thúc trước center -> LEFT
    if x1 < left_margin_threshold and x2 < page_center + center_tolerance:
        return WD_ALIGN_PARAGRAPH.LEFT
    # Nếu bbox kết thúc gần lề phải và bắt đầu sau center -> RIGHT
    elif x2 > right_margin_threshold and x1 > page_center - center_tolerance:
        return WD_ALIGN_PARAGRAPH.RIGHT
    # Nếu center_x gần với page_center -> CENTER
    elif abs(center_x - page_center) < center_tolerance:
        return WD_ALIGN_PARAGRAPH.CENTER
    # Mặc định: LEFT
    else:
        return WD_ALIGN_PARAGRAPH.LEFT


def sort_bboxes_by_position(bboxes, page_width=None):
    """Sort bboxes by reading order - improved header handling for multi-column layout"""
    if not bboxes:
        return []
    
    if page_width is None:
        max_x = max(bbox['x2'] for bbox in bboxes if bbox.get('x2'))
        page_width = max_x if max_x > 0 else 2000
    
    max_y = max(bbox['y2'] for bbox in bboxes if bbox.get('y2'))
    page_height = max_y if max_y > 0 else 3000
    
    # Calculate median bbox height for better tolerance calculation
    bbox_heights = [bbox['y2'] - bbox['y1'] for bbox in bboxes if 'y1' in bbox and 'y2' in bbox]
    if bbox_heights:
        median_height = sorted(bbox_heights)[len(bbox_heights) // 2]
    else:
        median_height = 30  # Default fallback
    
    def sort_body_bboxes(body_bboxes):
        """
        Sort body bboxes: simple top-to-bottom, left-to-right (natural reading order)
        No priority system - just sort by Y position, then X position within each row

        FIX: Sort by center_y (not y1) so that a tall multi-line bbox whose y1 happens
        to sit slightly above a shorter single-line bbox is still placed AFTER it in the
        reading order.  Using y1 as the primary key caused, e.g., a 2-line paragraph
        starting at y1=550 to sort before a 1-line heading at y1=560 even though the
        paragraph is visually lower on the page (its center_y=585 > heading center_y=575).
        """
        if not body_bboxes:
            return []
        
        # Step 1: Sort all bboxes by center_y (vertical midpoint) then x1.
        # center_y is more robust than y1 for bboxes of varying heights because it
        # represents the visual "middle" of each element rather than its top edge.
        sorted_by_y = sorted(body_bboxes, key=lambda b: (b['center_y'], b['x1']))
        
        # Step 2: Group into strips (same row) based on center_y tolerance
        # Use stricter tolerance to avoid grouping different rows together
        strips = []
        current_strip = []
        last_center_y = None  # Track strip reference by center_y, NOT y1
        
        for bbox in sorted_by_y:
            bbox_center_y = bbox['center_y']
            bbox_height = bbox.get('y2', bbox['y1']) - bbox['y1']
            
            if last_center_y is None:
                # First bbox
                current_strip.append(bbox)
                last_center_y = bbox_center_y
            else:
                # Calculate center_y difference
                y_diff = abs(bbox_center_y - last_center_y)
                
                # Use MUCH stricter tolerance: max(median_height * 0.3, 10px)
                # Chỉ nhóm vào cùng dòng nếu center_y difference rất nhỏ (< 30% of median height)
                strict_tolerance = max(median_height * 0.3, 10)
                
                if y_diff <= strict_tolerance:
                    # Same row (within very strict tolerance) - add to current strip
                    current_strip.append(bbox)
                    # Update reference to average center of current strip
                    # (do NOT use min/max which can drift the reference)
                    last_center_y = sum(b['center_y'] for b in current_strip) / len(current_strip)
                else:
                    # New row - finalize current strip
                    if current_strip:
                        # Sort within strip: left to right
                        current_strip.sort(key=lambda b: (b['x1'], b['center_x']))
                        strips.append(current_strip)
                    # Start new strip
                    current_strip = [bbox]
                    last_center_y = bbox_center_y
        
        # Handle last strip
        if current_strip:
            current_strip.sort(key=lambda b: (b['x1'], b['center_x']))
            strips.append(current_strip)
        
        # Step 3: Flatten strips (already sorted top-to-bottom, left-to-right within each strip)
        result = []
        for strip in strips:
            result.extend(strip)
        
        return result
    
    # Sort tất cả bboxes theo center_y (top-to-bottom, left-to-right)
    return sort_body_bboxes(bboxes)


def process_pdf_to_docx(
    pdf_path: str,
    output_docx: str = None,
    model_path: str = None,
    ocr_model_path: str = None,
    imgsz: int = 1024,
    conf: float = 0.1,
    dpi: int = 300,
    enable_ocr: bool = False,
    load_4bit: bool = False,
    load_8bit: bool = False,
    max_pages: int = None
):
    """Complete pipeline: PDF → Single DOCX"""

    _root = find_monorepo_root(Path(__file__))
    model_path = resolve_yolo_weights(_root, model_path)
    if ocr_model_path is None:
        ocr_model_path = str(_root / "Qwen2.5-VL-3B")
    else:
        # Resolve if it looks like a local path (directory exists)
        p = Path(ocr_model_path)
        if p.exists():
            ocr_model_path = str(p.resolve())
        # else: keep as-is for HuggingFace model IDs (e.g. "Qwen/Qwen2.5-VL-3B")

    pdf_path = Path(pdf_path).resolve()
    if not pdf_path.exists():
        print(f"Error: PDF file not found: {pdf_path}")
        return None

    if output_docx is None:
        output_docx = pdf_path.parent / f"{pdf_path.stem}_reconstructed.docx"
    else:
        output_docx = Path(output_docx).resolve()
    
    print("=" * 80)
    print("COMPLETE PDF PROCESSING PIPELINE")
    print("=" * 80)
    print(f"PDF: {pdf_path.name}")
    print(f"Output: {output_docx.name}")
    print()
    
    # Step 1: Load YOLO model
    print("[Step 1] Loading YOLO model...")
    model = YOLOv10(str(model_path))
    device = 'cuda' if torch.cuda.is_available() else 'cpu'
    print(f"  Using device: {device}")
    print("  ✓ Model loaded")
    
    # Step 2: Load OCR models if enabled
    processor = None
    qwen_model = None
    html_parser = None
    if enable_ocr:
        print(f"\n[Step 2] Loading Qwen2.5-VL OCR model from '{ocr_model_path}'...")
        try:
            # Cấu hình quantization để giảm VRAM
            model_kwargs = {
                "device_map": "auto",
            }
            
            if load_4bit:
                from transformers import BitsAndBytesConfig
                quantization_config = BitsAndBytesConfig(
                    load_in_4bit=True,
                    bnb_4bit_compute_dtype=torch.float16,
                    bnb_4bit_quant_type="nf4",
                    bnb_4bit_use_double_quant=True,
                )
                model_kwargs["quantization_config"] = quantization_config
                print("  Using 4-bit quantization (BitsAndBytes NF4)")
            elif load_8bit:
                from transformers import BitsAndBytesConfig
                quantization_config = BitsAndBytesConfig(
                    load_in_8bit=True,
                )
                model_kwargs["quantization_config"] = quantization_config
                print("  Using 8-bit quantization (BitsAndBytes)")
            else:
                model_kwargs["torch_dtype"] = torch.float16
                print("  Using float16 (no quantization)")
            
            qwen_model = Qwen2_5_VLForConditionalGeneration.from_pretrained(
                ocr_model_path, **model_kwargs
            )
            processor = AutoProcessor.from_pretrained(ocr_model_path)
            processor.tokenizer.padding_side = 'left'
            html_parser = HtmlToDocx()
            print("  ✓ Qwen2.5-VL OCR model loaded")
        except Exception as e:
            print(f"  ⚠️  Failed to load Qwen2.5-VL OCR model, skipping OCR step. Error: {e}")
            enable_ocr = False
    
    # Step 3: Convert PDF to images
    print(f"\n[Step 3] Converting PDF to images (DPI={dpi})...")
    try:
        images = pdf_to_images(str(pdf_path), dpi=dpi)
    except Exception as e:
        print(f"  ⚠️  Failed to convert PDF to images: {e}")
        return None
    total_pages = len(images)
    if max_pages is not None and max_pages > 0:
        images = images[:max_pages]
        print(f"  ✓ Converted {total_pages} pages, processing first {len(images)} pages")
    else:
        print(f"  ✓ Converted {len(images)} pages to images")
    
    # Step 4: Create Word document
    print("\n[Step 4] Creating Word document...")
    doc = Document()
    
    # Set page margins (làm cho trang "sát" hơn so với PDF gốc)
    sections = doc.sections
    for section in sections:
        # Giảm lề trên/dưới để khoảng trắng giữa các trang nhỏ lại
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
    
    # Process each page
    print("\n[Step 5] Processing pages...")
    all_pages_bboxes = []
    
    for page_idx, image in enumerate(images):
        print(f"\n  Processing page {page_idx + 1}/{len(images)}...")
        
        # ========================================================================
        # BƯỚC 1: YOLO DETECT BBOXES (không có thứ tự)
        # ========================================================================
        # YOLO detect tất cả bboxes cùng lúc (parallel detection)
        # Kết quả KHÔNG có thứ tự - cần sort lại sau
        # ========================================================================
        print("    Detecting bboxes...")
        detections = detect_bboxes(model, image, imgsz=imgsz, conf=conf)
        boxes = detections['boxes']
        scores = detections['scores']
        class_names = detections['class_names']
        class_ids = detections['class_ids']
        
        print(f"    Found {len(boxes)} bboxes (unsorted - will be sorted later)")
        
        # ========================================================================
        # BƯỚC 1.5: LOẠI BỎ CÁC BBOXES CHỒNG LÊN NHAU (OVERLAPPING)
        # ========================================================================
        # YOLO có thể detect cùng một vùng text với nhiều class khác nhau
        # (ví dụ: "title" và "plain text" cho cùng một vùng)
        # Loại bỏ các bboxes có overlap > 80%, giữ lại bbox có confidence cao hơn
        # ========================================================================
        print("    Removing overlapping bboxes (overlap > 80%)...")
        filtered_boxes = []
        filtered_scores = []
        filtered_class_names = []
        removed_indices = set()
        
        for i, (box1, score1, class1) in enumerate(zip(boxes, scores, class_names)):
            if i in removed_indices:
                continue
            
            x1_1, y1_1, x2_1, y2_1 = box1
            is_duplicate = False
            
            for j, (box2, score2, class2) in enumerate(zip(boxes, scores, class_names)):
                if i == j or j in removed_indices:
                    continue
                
                x1_2, y1_2, x2_2, y2_2 = box2
                
                # Tính overlap
                overlap_x = max(0, min(x2_1, x2_2) - max(x1_1, x1_2))
                overlap_y = max(0, min(y2_1, y2_2) - max(y1_1, y1_2))
                overlap_area = overlap_x * overlap_y
                
                area1 = (x2_1 - x1_1) * (y2_1 - y1_1)
                area2 = (x2_2 - x1_2) * (y2_2 - y1_2)
                min_area = min(area1, area2)
                
                if min_area > 0:
                    overlap_ratio = overlap_area / min_area
                    
                    # Nếu overlap > 80%, loại bỏ bbox có confidence thấp hơn
                    if overlap_ratio > 0.8:
                        if score1 < score2:
                            is_duplicate = True
                            removed_indices.add(i)
                            break
                        else:
                            removed_indices.add(j)
            
            if not is_duplicate:
                filtered_boxes.append(box1)
                filtered_scores.append(score1)
                filtered_class_names.append(class1)
        
        boxes = filtered_boxes
        scores = filtered_scores
        class_names = filtered_class_names
        print(f"    ✓ Kept {len(boxes)} bboxes after removing overlaps (removed {len(removed_indices)} overlapping bboxes)")
        
        # ========================================================================
        # BƯỚC 2: OCR VÀ XỬ LÝ TỪNG BBOX (chưa có thứ tự)
        # ========================================================================
        # OCR từng bbox để lấy text, nhưng chưa sort theo thứ tự đọc
        # ========================================================================
        # Prepare batches for OCR
        valid_boxes_info = []
        messages_batch = []
        
        for idx, (box, score, class_name) in enumerate(zip(boxes, scores, class_names)):
            x1, y1, x2, y2 = box
            width = x2 - x1
            height = y2 - y1
            
            is_in_header_zone = y1 < image.shape[0] * 0.20
            if class_name == 'abandon':
                if (height > width * 3 or width > height * 3) and not is_in_header_zone:
                    continue
            
            bbox_image = crop_bbox(image, box, padding=10)
            if bbox_image.shape[0] == 0 or bbox_image.shape[1] == 0:
                continue
            
            if enable_ocr and qwen_model is not None and processor is not None:
                pil_image = Image.fromarray(cv2.cvtColor(bbox_image, cv2.COLOR_BGR2RGB))
                prompt_text = (
                    "Trích xuất NGUYÊN VĂN (không được bịa thêm) trong hình ảnh sang định dạng HTML.\n"
                    "YÊU CẦU CỐT LÕI:\n"
                    "- KHÔNG ĐƯỢC tự suy diễn thành danh sách (bullet/numbered list) nếu trên ảnh không THẤY RÕ từng gạch đầu dòng/dấu chấm số.\n"
                    "YÊU CẦU ĐỊNH DẠNG:\n"
                    "- Giữ nguyên định dạng gốc của văn bản. Nếu chữ trong ảnh là in đậm, in nghiêng, hoặc gạch chân, hãy dùng thẻ <b>, <i>, <u> tương ứng.\n"
                    "- Nếu là đoạn văn bình thường, dùng thẻ <p>.\n"
                    "- Nếu là danh sách có dấu đầu dòng hoặc đánh số, dùng thẻ <ul> hoặc <ol> cùng với <li>.\n"
                    "- Nếu trong ảnh có BẢNG (hàng, cột, ô), dùng <table>, mỗi hàng <tr>, mỗi ô <td> hoặc <th> (tiêu đề). Không chuyển bảng thành danh sách hay đoạn.\n"
                    "- LƯU Ý: Chỉ thêm thẻ (b, i, ul, table...) khi thấy RÕ RÀNG trong ảnh.\n"
                    "Chỉ xuất ra mã HTML thuần túy, tuyệt đối không giải thích."
                )
                messages = [
                    {
                        "role": "system",
                        "content": [
                            {"type": "text", "text": (
                                "Bạn là công cụ OCR chuyển ảnh văn bản thành HTML. Nhận diện văn bản và cấu trúc (in đậm, in nghiêng, danh sách, bảng). Khi có bảng phải dùng <table>, <tr>, <td>, <th>. Trích xuất trung thực."
                            )},
                        ],
                    },
                    {
                        "role": "user",
                        "content": [
                            {"type": "image", "image": pil_image},
                            {"type": "text", "text": prompt_text},
                        ],
                    }
                ]
                messages_batch.append(messages)
            
            valid_boxes_info.append({
                'idx': idx, 'box': box, 'score': score, 'class_name': class_name, 'bbox_image': bbox_image, 'text': ""
            })
            
        # Run batched OCR in chunks to prevent VRAM spikes
        if enable_ocr and qwen_model is not None and processor is not None and messages_batch:
            BATCH_SIZE = 16
            print(f"      Running batched OCR for {len(messages_batch)} bboxes (Batch size: {BATCH_SIZE})...")
            
            all_texts = []
            import math
            total_batches = math.ceil(len(messages_batch) / BATCH_SIZE)
            
            for i in range(0, len(messages_batch), BATCH_SIZE):
                batch_msgs = messages_batch[i:i+BATCH_SIZE]
                print(f"        Processing batch {i//BATCH_SIZE + 1}/{total_batches}...")
                try:
                    text_inputs = processor.apply_chat_template(batch_msgs, tokenize=False, add_generation_prompt=True)
                    image_inputs, video_inputs = process_vision_info(batch_msgs)
                    
                    inputs = processor(
                        text=text_inputs,
                        images=image_inputs,
                        videos=video_inputs,
                        padding=True,
                        return_tensors="pt",
                    )
                    inputs = inputs.to(qwen_model.device)
                    
                    with torch.no_grad():
                        # Generate with memory optimization for batching
                        generated_ids = qwen_model.generate(**inputs, max_new_tokens=2048)
                    
                    generated_ids_trimmed = [
                        out_ids[len(in_ids):] for in_ids, out_ids in zip(inputs.input_ids, generated_ids)
                    ]
                    texts = processor.batch_decode(
                        generated_ids_trimmed, skip_special_tokens=True, clean_up_tokenization_spaces=False
                    )
                    all_texts.extend(texts)
                    
                    # Clear cache to free VRAM for next batch
                    torch.cuda.empty_cache()
                except Exception as e:
                    print(f"        ⚠️ Batch {i//BATCH_SIZE + 1} failed: {e}")
                    # Append empty strings for failures to keep alignment
                    all_texts.extend([""] * len(batch_msgs))
                    torch.cuda.empty_cache()
            
            # ============================================================
            # Phát hiện output rác bằng cách so sánh với prompt
            # Thay vì hardcode từng pattern, dùng word-overlap ratio
            # ============================================================
            # Tách prompt thành tập từ đặc trưng (loại bỏ từ phổ biến)
            prompt_words = set(prompt_text.lower().split())
            # Thêm từ khóa refusal phổ biến
            refusal_keywords = {'sorry', 'cannot', 'unable', 'apologize', 'không thể', 'xin lỗi', 'không rõ'}
            
            # Map texts back to valid boxes
            for i, valid_box in enumerate(valid_boxes_info):
                if i < len(all_texts):
                    t = all_texts[i]
                    # Strip markdown/html code block wrappers
                    t = t.replace("```markdown\n", "").replace("```markdown\\n", "").replace("```html\n", "").replace("```html\\n", "").replace("```", "").strip()
                    # Strip literal prefixes
                    for prefix in ["markdown", "html", "text:"]:
                        if t.lower().startswith(prefix):
                            t = t[len(prefix):].strip()
                    
                    # Smart junk detection: so sánh output với prompt
                    t_lower = t.lower()
                    t_words = set(t_lower.split())
                    
                    is_junk = False
                    if t_words:
                        # Tính tỉ lệ từ trong output trùng với prompt
                        overlap = t_words & prompt_words
                        overlap_ratio = len(overlap) / len(t_words)
                        # Nếu > 50% từ trong output trùng với prompt → echo
                        if overlap_ratio > 0.5 and len(t_words) > 3:
                            is_junk = True
                        # Kiểm tra refusal keywords
                        if t_words & refusal_keywords and len(t_words) < 30:
                            is_junk = True
                    
                    if is_junk:
                        t = ""
                    
                    # Strip fake leading bullet "- " cho non-list text
                    if t.startswith("- ") and t.count("\n- ") == 0:
                        t = t[2:].strip()
                    
                    valid_box['text'] = t
                else:
                    valid_box['text'] = ""

                
        # Post-process valid boxes
        seen_texts_positions = {}
        page_bboxes = []
        
        for box_info in valid_boxes_info:
            idx = box_info['idx']
            box = box_info['box']
            score = box_info['score']
            class_name = box_info['class_name']
            text = box_info['text']
            x1, y1, x2, y2 = box
            width = x2 - x1
            height = y2 - y1
            
            if not text or text in ['[No text detected]', '[Empty bbox]'] or text.startswith('[OCR Error:'):
                continue

            # Generic duplicate detection by position overlap or proximity
            text_normalized = text.lower().strip()
            if text_normalized in seen_texts_positions:
                existing_bbox = seen_texts_positions[text_normalized]
                existing_x1, existing_y1, existing_x2, existing_y2 = existing_bbox['coords']

                overlap_x    = max(0, min(x2, existing_x2) - max(x1, existing_x1))
                overlap_y    = max(0, min(y2, existing_y2) - max(y1, existing_y1))
                overlap_area = overlap_x * overlap_y
                bbox_area    = width * height
                existing_area = (existing_x2 - existing_x1) * (existing_y2 - existing_y1)
                overlap_ratio = overlap_area / min(bbox_area, existing_area) if min(bbox_area, existing_area) > 0 else 0

                center_x = (x1 + x2) / 2
                center_y = (y1 + y2) / 2
                existing_center_x = (existing_x1 + existing_x2) / 2
                existing_center_y = (existing_y1 + existing_y2) / 2
                distance = ((center_x - existing_center_x) ** 2 + (center_y - existing_center_y) ** 2) ** 0.5

                if overlap_ratio > 0.3 or distance < 50:
                    continue

            seen_texts_positions[text_normalized] = {
                'coords': (x1, y1, x2, y2),
                'score': float(score)
            }

            x1, y1, x2, y2 = box
            page_bboxes.append({
                'class': class_name,
                'confidence': float(score),
                'text': text,
                'x1': float(x1),
                'y1': float(y1),
                'x2': float(x2),
                'y2': float(y2),
                'center_x': (x1 + x2) / 2,
                'center_y': (y1 + y2) / 2,
            })

        content_bboxes = page_bboxes
        
        # ========================================================================
        # BƯỚC 3: SORT BBOXES THEO THỨ TỰ ĐỌC TỰ NHIÊN
        # ========================================================================
        # Sort bboxes theo quy luật: từ trên xuống dưới (top-to-bottom),
        # từ trái qua phải (left-to-right) - như mắt người đọc sách
        # ========================================================================
        if content_bboxes:
            page_width = max(bbox['x2'] for bbox in content_bboxes)
            
            # Debug: Print Y positions before sorting (only for first page)
            if page_idx == 0:
                print(f"    Debug: Bbox Y positions before sorting:")
                for i, bbox in enumerate(content_bboxes[:10], 1):
                    text_preview = bbox['text'][:50].replace('\n', ' ')
                    print(f"      {i}. Y={bbox['y1']:.0f}: {text_preview}...")
            
            # Sort theo thứ tự đọc tự nhiên: top-to-bottom, left-to-right
            sorted_bboxes = sort_bboxes_by_position(content_bboxes, page_width)
            
            # ========================================================================
            # BƯỚC 4: ĐÁNH SỐ THỨ TỰ ĐỌC (READING ORDER)
            # ========================================================================
            # Sau khi sort xong, đánh số thứ tự 1, 2, 3, ... cho các bboxes
            # theo thứ tự đọc tự nhiên như mắt người đọc sách
            # Số thứ tự này được lưu vào bbox['reading_order'] để đảm bảo
            # OCR và xử lý text theo đúng thứ tự này
            # ========================================================================
            print(f"    Assigning reading order numbers (top-to-bottom, left-to-right)...")
            for reading_order, bbox in enumerate(sorted_bboxes, start=1):
                bbox['reading_order'] = reading_order
            
            # Debug: Print reading order and Y positions after sorting (only for first page)
            if page_idx == 0:
                print(f"    Debug: Bbox reading order and positions after sorting:")
                for bbox in sorted_bboxes[:12]:
                    reading_order = bbox.get('reading_order', '?')
                    text_preview = bbox['text'][:50].replace('\n', ' ')
                    print(f"      #{reading_order} (Y={bbox['y1']:.0f}, X={bbox['x1']:.0f}): {text_preview}...")
            
            all_pages_bboxes.append((page_idx + 1, sorted_bboxes, image.shape[1], image.shape[0]))
            print(f"    ✓ Processed {len(sorted_bboxes)} bboxes with text")
        else:
            print(f"    ⚠️  No valid bboxes found for this page")
    
    # Step 6: Add all pages to Word document
    # ========================================================================
    # Sử dụng DrawingML textbox (wp:anchor) để đặt mỗi bbox vào đúng vị trí
    # trên trang Word. Mỗi bbox → 1 textbox duy nhất chứa ALL paragraphs.
    # ========================================================================
    print("\n[Step 6] Adding pages to Word document...")
    
    # EMU constants: 1 inch = 914400 EMU, 1 pt = 12700 EMU, 1 cm = 360000 EMU
    EMU_PER_PT = 12700
    
    # A4 page dimensions in points
    A4_WIDTH_PT = 595
    A4_HEIGHT_PT = 842
    
    # Margins in points (match the section margins set earlier: left=3cm, right=2cm, top=1.5cm)
    LEFT_MARGIN_PT = 3.0 * 28.3465  # ~85 pt
    TOP_MARGIN_PT = 1.5 * 28.3465   # ~42.5 pt
    
    textbox_id_counter = 100  # unique id for each textbox shape
    
    def create_textbox_element(textbox_id, x_emu, y_emu, w_emu, h_emu, content_paragraphs_xml, wrap_type="none"):
        """
        Tạo DrawingML textbox element (wp:anchor) để neo tại vị trí tuyệt đối trên trang.
        
        Parameters:
            textbox_id: unique shape id
            x_emu: horizontal position in EMU from page edge
            y_emu: vertical position in EMU from page edge  
            w_emu: width in EMU
            h_emu: height in EMU (min height, auto-extend if needed)
            content_paragraphs_xml: raw XML string of <w:p> elements to put inside
            wrap_type: 'none' (floating) or 'square' (text wrapping)
        
        Returns:
            OxmlElement for w:drawing
        """
        from lxml import etree
        
        nsmap = {
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
            'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
            'v': 'urn:schemas-microsoft-com:vml',
        }
        
        if wrap_type == 'square':
            wrap_xml = '<wp:wrapSquare wrapText="bothSides"/>'
        elif wrap_type == 'topAndBottom':
            wrap_xml = '<wp:wrapTopAndBottom/>'
        else:
            wrap_xml = '<wp:wrapNone/>'
        
        # Build the anchor XML
        drawing_xml = f'''<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                     xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                     xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                     xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
                     xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
                     xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
                     xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
          <wp:anchor distT="0" distB="0" distL="0" distR="0"
                     simplePos="0" relativeHeight="{textbox_id}"
                     behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="page">
              <wp:posOffset>{x_emu}</wp:posOffset>
            </wp:positionH>
            <wp:positionV relativeFrom="page">
              <wp:posOffset>{y_emu}</wp:posOffset>
            </wp:positionV>
            <wp:extent cx="{w_emu}" cy="{h_emu}"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            {wrap_xml}
            <wp:docPr id="{textbox_id}" name="TextBox {textbox_id}"/>
            <a:graphic>
              <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                <wps:wsp>
                  <wps:cNvSpPr txBox="1"/>
                  <wps:spPr>
                    <a:xfrm>
                      <a:off x="0" y="0"/>
                      <a:ext cx="{w_emu}" cy="{h_emu}"/>
                    </a:xfrm>
                    <a:prstGeom prst="rect">
                      <a:avLst/>
                    </a:prstGeom>
                    <a:noFill/>
                    <a:ln>
                      <a:noFill/>
                    </a:ln>
                  </wps:spPr>
                  <wps:txbx>
                    <w:txbxContent>
                      {content_paragraphs_xml}
                    </w:txbxContent>
                  </wps:txbx>
                  <wps:bodyPr wrap="square" lIns="0" tIns="0" rIns="0" bIns="0"
                              anchor="t" anchorCtr="0">
                    <a:noAutofit/>
                  </wps:bodyPr>
                </wps:wsp>
              </a:graphicData>
            </a:graphic>
          </wp:anchor>
        </w:drawing>'''
        
        drawing_element = etree.fromstring(drawing_xml)
        return drawing_element
    
    def paragraphs_to_xml(paragraphs):
        """
        Chuyển danh sách python-docx Paragraph objects thành raw XML string.
        """
        from lxml import etree
        xml_parts = []
        for p in paragraphs:
            xml_parts.append(etree.tostring(p._p, encoding='unicode'))
        return '\n'.join(xml_parts)
    
    for page_num, sorted_bboxes, page_width, page_height in all_pages_bboxes:
        if page_num > 1:
            page_break_para = doc.add_paragraph()
            page_break_para.paragraph_format.space_before = Pt(0)
            page_break_para.paragraph_format.space_after = Pt(0)
            page_break_para.paragraph_format.line_spacing = 1.0
            page_break_run = page_break_para.add_run()
            page_break_run.add_break(WD_BREAK.PAGE)
        
        # Đảm bảo section có kích thước chuẩn A4
        section = doc.sections[-1]
        section.page_width = Pt(A4_WIDTH_PT)
        section.page_height = Pt(A4_HEIGHT_PT)
        
        # Tạo 1 paragraph ẩn để neo tất cả textboxes của trang này, tránh tạo dòng trống
        page_anchor_para = doc.add_paragraph()
        page_anchor_para.paragraph_format.space_before = Pt(0)
        page_anchor_para.paragraph_format.space_after = Pt(0)
        page_anchor_para.paragraph_format.line_spacing = 1.0
        page_anchor_run = page_anchor_para.add_run()
        page_anchor_run.font.size = Pt(1)
        
        for bbox in sorted_bboxes:
            text = bbox['text'].strip()
            if not text:
                continue
            
            x1_px, y1_px = bbox['x1'], bbox['y1']
            x2_px, y2_px = bbox['x2'], bbox['y2']
            
            # ============================================================
            # Tính toạ độ EMU (English Metric Units) từ pixel coordinates
            # Pixel → tỉ lệ trên trang → điểm (pt) → EMU
            # ============================================================
            x_pt = (x1_px / page_width) * A4_WIDTH_PT
            y_pt = (y1_px / page_height) * A4_HEIGHT_PT
            width_pt = ((x2_px - x1_px) / page_width) * A4_WIDTH_PT
            height_pt = ((y2_px - y1_px) / page_height) * A4_HEIGHT_PT
            
            # Mở rộng width thêm 15% để text ít bị wrap hơn, giảm overlap
            width_pt = width_pt * 1.15
            # Giới hạn width không vượt quá lề phải trang (A4_WIDTH_PT)
            max_width = A4_WIDTH_PT - x_pt
            width_pt = min(width_pt, max_width)
            
            # Đảm bảo kích thước tối thiểu
            width_pt = max(width_pt, 30)
            height_pt = max(height_pt, 12)
            
            x_emu = int(x_pt * EMU_PER_PT)
            y_emu = int(y_pt * EMU_PER_PT)
            w_emu = int(width_pt * EMU_PER_PT)
            h_emu = int(height_pt * EMU_PER_PT)
            
            # ============================================================
            # Tạo nội dung cho textbox: convert markdown → HTML → Word paragraphs
            # Dùng tạm 1 Document phụ để HtmlToDocx render HTML ra paragraphs,
            # sau đó copy XML sang textbox content.
            # ============================================================
            # Nếu text đã là HTML (chứa tags), dùng trực tiếp
            # Nếu là plain text hoặc markdown, convert qua markdown trước
            if '<' in text and ('>' in text) and any(tag in text.lower() for tag in ['<p>', '<b>', '<i>', '<u>', '<ul>', '<ol>', '<li>', '<br', '<h']):
                html_content = text
            else:
                html_content = markdown.markdown(text)
            
            # Sanitize HTML: loại bỏ <img> tags không có src (gây crash HtmlToDocx)
            html_content = re.sub(r'<img(?![^>]*src\s*=)[^>]*/?\s*>', '', html_content, flags=re.IGNORECASE)
            # Loại bỏ các tag lạ không phải HTML chuẩn (ví dụ: <M.S.D.N.:0318788>)
            html_content = re.sub(r'<(?!/?(?:p|b|i|u|s|a|br|hr|h[1-6]|ul|ol|li|table|tr|td|th|thead|tbody|div|span|strong|em|sub|sup|pre|code|blockquote)\b)[A-Z][^>]*>', '', html_content, flags=re.IGNORECASE)
            
            # Tạo document phụ để render HTML
            temp_doc = Document()
            try:
                if html_parser is not None:
                    html_parser.add_html_to_document(html_content, temp_doc)
                else:
                    p = temp_doc.add_paragraph()
                    p.add_run(text)
            except Exception as e:
                # Fallback: nếu HTML parsing lỗi, dùng plain text
                print(f"      ⚠️ HTML parse error, using plain text: {e}")
                temp_doc = Document()
                p = temp_doc.add_paragraph()
                # Strip tất cả HTML tags để lấy plain text
                plain_text = re.sub(r'<[^>]+>', ' ', text).strip()
                plain_text = re.sub(r'\s+', ' ', plain_text)
                p.add_run(plain_text)
            
            # Lấy các paragraphs đã render (bỏ qua paragraph trống đầu tiên mặc định)
            temp_paragraphs = temp_doc.paragraphs
            if temp_paragraphs and not temp_paragraphs[0].text.strip():
                temp_paragraphs = temp_paragraphs[1:]
            
            if not temp_paragraphs:
                continue
            
            # Đặt font và spacing cho các paragraph bên trong textbox
            # GIỮ NGUYÊN bold/italic/underline mà HtmlToDocx đã set từ HTML
            for tp in temp_paragraphs:
                tp.paragraph_format.space_after = Pt(0)
                tp.paragraph_format.space_before = Pt(0)
                tp.paragraph_format.line_spacing = 1.0
                for run in tp.runs:
                    # Lưu lại formatting hiện tại trước khi set font
                    existing_bold = run.font.bold
                    existing_italic = run.font.italic
                    existing_underline = run.font.underline
                    
                    run.font.name = 'Times New Roman'
                    if run._element.rPr is not None:
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    else:
                        rPr = OxmlElement('w:rPr')
                        run._element.insert(0, rPr)
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                        rPr.append(rFonts)
                    run.font.size = Pt(10)
                    
                    # Khôi phục bold/italic/underline
                    if existing_bold is not None:
                        run.font.bold = existing_bold
                    if existing_italic is not None:
                        run.font.italic = existing_italic
                    if existing_underline is not None:
                        run.font.underline = existing_underline
            
            # Chuyển paragraph objects → XML string
            content_xml = paragraphs_to_xml(temp_paragraphs)
            
            # Tạo textbox element và gắn vào document
            textbox_id_counter += 1
            drawing_elem = create_textbox_element(
                textbox_id_counter, x_emu, y_emu, w_emu, h_emu, content_xml
            )
            
            # Gắn vào paragraph neo
            page_anchor_run._element.append(drawing_elem)
        
    # --- STEP 8 (REVISED): Render bboxes sang văn bản thường dùng indent + invisible table ---
    print("\n[Step 8] Creating clean transcript (indent + table approach)...")

    transcript_doc = Document()
    for section in transcript_doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    A4_WIDTH_PT = 595.0
    A4_HEIGHT_PT = 842.0
    CONTENT_WIDTH_CM = 15.0  # A4 trừ lề trái 3cm + lề phải 2cm
    LEFT_MARGIN_CM = 3.0

    # --- STEP 8: Create clean transcript ---
    print("\n[Step 8] Creating clean transcript (indent + table approach)...")

    transcript_doc = Document()
    for section in transcript_doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.0)

    A4_WIDTH_PT      = 595.0
    A4_HEIGHT_PT     = 842.0
    CONTENT_WIDTH_CM = 15.0
    LEFT_MARGIN_CM   = 3.0

    # ── helpers ──────────────────────────────────────────────────────────────

    def _set_table_borders_none(table):
        tbl   = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)
        for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            el = tblBorders.find(qn(f'w:{edge}'))
            if el is None:
                el = OxmlElement(f'w:{edge}')
                tblBorders.append(el)
            el.set(qn('w:val'),   'none')
            el.set(qn('w:sz'),    '0')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), 'auto')
        # Zero cell margins để tránh trang trắng thừa
        tblCellMar = tblPr.find(qn('w:tblCellMar'))
        if tblCellMar is None:
            tblCellMar = OxmlElement('w:tblCellMar')
            tblPr.append(tblCellMar)
        for side in ['top', 'left', 'bottom', 'right']:
            mar = tblCellMar.find(qn(f'w:{side}'))
            if mar is None:
                mar = OxmlElement(f'w:{side}')
                tblCellMar.append(mar)
            mar.set(qn('w:w'),    '0')
            mar.set(qn('w:type'), 'dxa')

    def _apply_font(para, space_after_pt=4.0, space_before_pt=0.0):
        """Áp font Times New Roman 10pt, giữ nguyên bold/italic/underline."""
        para.paragraph_format.space_after  = Pt(space_after_pt)
        para.paragraph_format.space_before = Pt(space_before_pt)
        para.paragraph_format.line_spacing = 1.0
        for run in para.runs:
            b, i, u = run.font.bold, run.font.italic, run.font.underline
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            if run._element.rPr is not None:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            else:
                rPr    = OxmlElement('w:rPr')
                run._element.insert(0, rPr)
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                rPr.append(rFonts)
            if b is not None: run.font.bold      = b
            if i is not None: run.font.italic    = i
            if u is not None: run.font.underline = u

    def _render_text_to_raw_paragraphs(text):
        """Chuyển text thành list paragraphs - KHÔNG convert markdown để giữ nguyên OCR output."""
        # Chỉ dùng html_parser nếu text thực sự là HTML có tags
        if '<' in text and '>' in text and any(
            t in text.lower() for t in ['<p>', '<b>', '<i>', '<u>', '<ul>', '<ol>', '<li>', '<br', '<h']
        ):
            html_content = text
            html_content = re.sub(r'<img(?![^>]*src\s*=)[^>]*/?\s*>', '', html_content, flags=re.IGNORECASE)
            html_content = re.sub(
                r'<(?!/?(?:p|b|i|u|s|a|br|hr|h[1-6]|ul|ol|li|table|tr|td|th|thead|tbody|div|span|strong|em|sub|sup|pre|code|blockquote)\b)[A-Z][^>]*>',
                '', html_content, flags=re.IGNORECASE
            )
            tmp = Document()
            try:
                if html_parser is not None:
                    html_parser.add_html_to_document(html_content, tmp)
                else:
                    tmp.add_paragraph().add_run(text)
            except Exception:
                tmp = Document()
                tmp.add_paragraph().add_run(
                    re.sub(r'\s+', ' ', re.sub(r'<[^>]+>', ' ', text)).strip()
                )
        else:
            # Plain text: tách theo dòng, mỗi dòng là 1 paragraph
            tmp = Document()
            lines = text.splitlines()
            for line in lines:
                line = line.strip()
                if line:
                    tmp.add_paragraph().add_run(line)
            if not any(p.text.strip() for p in tmp.paragraphs):
                tmp.add_paragraph().add_run(text.strip())

        return [p for p in tmp.paragraphs if p.text.strip()]

    def _append_paragraphs_to_body(raw_paras, target_doc, indent_cm=0.0, space_after_pt=4.0, right_indent_cm=0.0):
        import copy
        from docx.text.paragraph import Paragraph
        for idx, tp in enumerate(raw_paras):
            new_p  = copy.deepcopy(tp._p)
            body   = target_doc._element.body
            if body.sectPr is not None:
                body.sectPr.addprevious(new_p)
            else:
                body.append(new_p)
            cloned = Paragraph(new_p, target_doc)
            is_last = (idx == len(raw_paras) - 1)
            _apply_font(cloned, space_after_pt=space_after_pt if is_last else 1.0)

            if indent_cm > 0.1 or right_indent_cm > 0.1:
                pPr   = new_p.get_or_add_pPr()
                ind   = pPr.find(qn('w:ind'))
                if ind is None:
                    ind = OxmlElement('w:ind')
                    pPr.append(ind)
                if indent_cm > 0.1:
                    existing = int(ind.get(qn('w:left'), '0') or '0')
                    ind.set(qn('w:left'), str(existing + int(indent_cm * 567)))
                if right_indent_cm > 0.1:
                    ind.set(qn('w:right'), str(int(right_indent_cm * 567)))

    def _append_paragraphs_to_cell(raw_paras, cell, alignment):
        """Copy raw paragraphs vào cell của table."""
        import copy
        from docx.text.paragraph import Paragraph
        for tp in raw_paras:
            new_p  = copy.deepcopy(tp._p)
            cell._element.append(new_p)
            cloned = Paragraph(new_p, cell)
            cloned.alignment = alignment
            _apply_font(cloned)

    def _parse_html_table_to_grid(text):
        """Parse HTML <table><tr><td>...</td></tr>...</table> into list of rows (list of list of cell text).
        Returns None if text is not a table, else [[cell1, cell2,...], ...].
        """
        if not text or '<table' not in text.lower() or '</tr>' not in text.lower():
            return None
        # Strip tbody for easier regex
        s = re.sub(r'</?tbody\s*>', '', text, flags=re.IGNORECASE)
        rows = []
        for tr_match in re.finditer(r'<tr[^>]*>(.*?)</tr>', s, re.DOTALL | re.IGNORECASE):
            row_html = tr_match.group(1)
            cells = re.findall(r'<t[hd][^>]*>(.*?)</t[hd]>', row_html, re.DOTALL | re.IGNORECASE)
            cell_texts = [re.sub(r'<[^>]+>', ' ', c).strip() for c in cells]
            cell_texts = [re.sub(r'\s+', ' ', t).strip() for t in cell_texts]
            if cell_texts:
                rows.append(cell_texts)
        if not rows:
            return None
        return rows

    def _insert_word_table_from_grid(doc, grid, space_after_pt=4.0):
        """Create a Word table (Insert table style) and fill each cell with text from grid. Apply font."""
        if not grid:
            return
        num_rows = len(grid)
        num_cols = max(len(r) for r in grid)
        if num_cols == 0:
            return
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        for r_idx, row_cells in enumerate(grid):
            for c_idx, cell_text in enumerate(row_cells):
                if c_idx >= num_cols:
                    break
                cell = table.cell(r_idx, c_idx)
                cell.text = cell_text
                for para in cell.paragraphs:
                    _apply_font(para, space_after_pt=0)
        # Spacer paragraph after table
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(space_after_pt)
        spacer.paragraph_format.line_spacing = 1.0

    def _bboxes_overlap_y(a, b, threshold=0.4):
        overlap = min(a['y2'], b['y2']) - max(a['y1'], b['y1'])
        h_a     = a['y2'] - a['y1']
        h_b     = b['y2'] - b['y1']
        min_h   = min(h_a, h_b)
        
        if min_h <= 0:
            return False
        
        # Loại trừ: nếu 1 bbox cao gấp 2.5 lần bbox kia
        # → bbox cao đó là paragraph nhiều dòng, không phải song song thật
        height_ratio = max(h_a, h_b) / min_h
        if height_ratio > 2.5:
            return False
    
        return (overlap / min_h) >= threshold

    def _bboxes_overlap_x(a, b, tol=10):
        return (min(a['x2'], b['x2']) - max(a['x1'], b['x1'])) > tol

    def _group_into_rows(bboxes):
        """Nhóm bboxes thành rows dựa trên overlap Y. Row > 1 phần tử = song song."""
        rows    = []
        current = []
        for bbox in bboxes:
            if not bbox['text'].strip():
                continue
            if not current:
                current.append(bbox)
            else:
                has_y = any(_bboxes_overlap_y(bbox, b) for b in current)
                has_x = any(_bboxes_overlap_x(bbox, b) for b in current)
                if has_y and not has_x:
                    current.append(bbox)
                else:
                    rows.append(current)
                    current = [bbox]
        if current:
            rows.append(current)
        return rows

    # ── main loop ─────────────────────────────────────────────────────────────

    for page_num, sorted_bboxes, page_width, page_height in all_pages_bboxes:

        # Page break
        if page_num > 1:
            pb = transcript_doc.add_paragraph()
            pb.paragraph_format.space_before = Pt(0)
            pb.paragraph_format.space_after  = Pt(0)
            pb.paragraph_format.line_spacing = 1.0
            pb.add_run().add_break(WD_BREAK.PAGE)

        # Phân loại mep vs main
        margin_l    = page_width * 0.05
        margin_r    = page_width * 0.95
        mep_bboxes  = []
        main_bboxes = []
        for bbox in sorted_bboxes:
            if not bbox['text'].strip():
                continue
            if bbox['x1'] < margin_l or bbox['x2'] > margin_r:
                mep_bboxes.append(bbox)
            else:
                main_bboxes.append(bbox)

        # ── Render mep bboxes: floating textbox ──────────────────────────────
        if mep_bboxes:
            anchor_para = transcript_doc.add_paragraph()
            anchor_para.paragraph_format.space_before = Pt(0)
            anchor_para.paragraph_format.space_after  = Pt(0)
            anchor_para.paragraph_format.line_spacing = 1.0
            anchor_run = anchor_para.add_run()

            for bbox in mep_bboxes:
                raw_paras = _render_text_to_raw_paragraphs(bbox['text'])
                if not raw_paras:
                    continue
                for tp in raw_paras:
                    _apply_font(tp)
                content_xml = paragraphs_to_xml(raw_paras)

                x_pt      = (bbox['x1'] / page_width)             * A4_WIDTH_PT
                y_pt      = (bbox['y1'] / page_height)            * A4_HEIGHT_PT
                width_pt  = ((bbox['x2'] - bbox['x1']) / page_width)  * A4_WIDTH_PT
                height_pt = ((bbox['y2'] - bbox['y1']) / page_height) * A4_HEIGHT_PT
                width_pt  = min(max(width_pt * 1.15, 30), A4_WIDTH_PT - x_pt)
                height_pt = max(height_pt, 12)

                textbox_id_counter += 1
                drawing_elem = create_textbox_element(
                    textbox_id_counter,
                    int(x_pt      * EMU_PER_PT),
                    int(y_pt      * EMU_PER_PT),
                    int(width_pt  * EMU_PER_PT),
                    int(height_pt * EMU_PER_PT),
                    content_xml, wrap_type="none"
                )
                anchor_run._element.append(drawing_elem)

        # ── Render main bboxes: indent + invisible table ──────────────────────
        rows = _group_into_rows(main_bboxes)

        for row_idx, row in enumerate(rows):
            row.sort(key=lambda b: b['x1'])

            # Tính space_after từ Y-gap tới row tiếp theo
            if row_idx + 1 < len(rows):
                next_row       = rows[row_idx + 1]
                cur_y2         = max(b['y2'] for b in row)
                next_y1        = min(b['y1'] for b in next_row)
                gap_px         = max(next_y1 - cur_y2, 0)
                gap_pt         = (gap_px / page_height) * A4_HEIGHT_PT
                # Clamp: tối thiểu 2pt, tối đa 40pt
                space_after_pt = max(2.0, min(gap_pt, 40.0))
            else:
                space_after_pt = 4.0

            if len(row) == 1:
                bbox      = row[0]
                text      = bbox['text'] or ''
                # If content is HTML table, create Word table and fill cells (like Insert Table + type text)
                grid = _parse_html_table_to_grid(text)
                if grid:
                    _insert_word_table_from_grid(transcript_doc, grid, space_after_pt=space_after_pt)
                else:
                    raw       = _render_text_to_raw_paragraphs(text)
                    indent_cm = max(0.0, (bbox['x1'] / page_width) * (A4_WIDTH_PT / 28.3465) - LEFT_MARGIN_CM)
                    bbox_x2_cm        = (bbox['x2'] / page_width) * (A4_WIDTH_PT / 28.3465)
                    content_right_cm  = (A4_WIDTH_PT / 28.3465) - 2.0
                    right_indent_cm   = max(0.0, content_right_cm - bbox_x2_cm)
                    if right_indent_cm < 0.5:
                        right_indent_cm = 0.0
                    _append_paragraphs_to_body(raw, transcript_doc, indent_cm,
                                               space_after_pt=space_after_pt,
                                               right_indent_cm=right_indent_cm)

            else:
                # Song song → invisible table, width cột theo tọa độ tuyệt đối
                table = transcript_doc.add_table(rows=1, cols=len(row))
                _set_table_borders_none(table)
                table.autofit = False

                row_x1   = min(b['x1'] for b in row)
                row_x2   = max(b['x2'] for b in row)
                row_span = max(row_x2 - row_x1, 1)

                for i, bbox in enumerate(row):
                    cell     = table.cell(0, i)
                    b_span   = bbox['x2'] - bbox['x1']
                    col_w_cm = (b_span / row_span) * CONTENT_WIDTH_CM
                    cell.width = Cm(max(col_w_cm, 1.0))

                    for p in list(cell.paragraphs):
                        cell._element.remove(p._element)

                    text_bbox = bbox['text'] or ''
                    grid     = _parse_html_table_to_grid(text_bbox)
                    if grid:
                        num_rows, num_cols = len(grid), max(len(r) for r in grid)
                        if num_cols > 0:
                            inner = transcript_doc.add_table(rows=num_rows, cols=num_cols)
                            inner.style = 'Table Grid'
                            tbl_el = inner._tbl
                            transcript_doc._element.body.remove(tbl_el)
                            cell._element.append(tbl_el)
                            for r_idx, row_cells in enumerate(grid):
                                for c_idx, cell_text in enumerate(row_cells):
                                    if c_idx < num_cols:
                                        inner.cell(r_idx, c_idx).text = cell_text
                                        for para in inner.cell(r_idx, c_idx).paragraphs:
                                            _apply_font(para)
                    else:
                        raw       = _render_text_to_raw_paragraphs(text_bbox)
                        alignment = determine_alignment_by_position(bbox, page_width, page_height)
                        if raw:
                            _append_paragraphs_to_cell(raw, cell, alignment)
                        else:
                            cell.add_paragraph()
                    if not list(cell._element):
                        cell.add_paragraph()

                # Paragraph spacer sau table để tạo khoảng cách tương đương gap
                spacer = transcript_doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(0)
                spacer.paragraph_format.space_after  = Pt(0)
                spacer.paragraph_format.line_spacing = 1.0
                spacer_run = spacer.add_run()
                spacer_run.font.size = Pt(max(1.0, min(space_after_pt * 0.5, 12.0)))
    
    # Step 9: Save documents
    print(f"\n[Step 9] Saving Word documents...")

    def save_doc_safe(d, path):
        try:
            d.save(str(path))
            print(f"  ✓ Saved to: {path}")
            return path
        except PermissionError:
            import datetime
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            new_output = path.parent / f"{path.stem}_{timestamp}.docx"
            print(f"  ⚠️  Original file is locked, saving as: {new_output.name}")
            d.save(str(new_output))
            print(f"  ✓ Saved to: {new_output}")
            return new_output

    # Create layout output path using the original path
    # layout_path = output_docx.parent / f"{output_docx.stem}_layout{output_docx.suffix}"
    
    # # Save both versions
    # layout_docx = save_doc_safe(doc, layout_path)
    transcript_docx = save_doc_safe(transcript_doc, output_docx)
    output_docx = transcript_docx
    
    print("\n" + "=" * 80)
    print("PROCESSING COMPLETE")
    print("=" * 80)
    print(f"Main output file (transcript): {output_docx}")
    # print(f"Layout output file: {layout_docx}")
    print(f"Total pages: {len(images)}")
    total_bboxes = sum(len(bboxes) for _, bboxes, _, _ in all_pages_bboxes)
    print(f"Total bboxes: {total_bboxes}")
    
    return output_docx


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Complete PDF to DOCX pipeline")
    parser.add_argument("pdf_path", type=str, help="Path to PDF file")
    parser.add_argument("--output", type=str, default=None, help="Output DOCX file path")
    parser.add_argument(
        "--model",
        type=str,
        default=None,
        help="Path to YOLO model file (default: models/doclayout_yolo_docstructbench_imgsz1024.pt)",
    )
    parser.add_argument(
        "--ocr-model",
        type=str,
        default=None,
        help="HuggingFace ID or local directory path for the Qwen2.5-VL model (default: ./Qwen2.5-VL-3B)",
    )
    parser.add_argument(
        "--use-4bit",
        action="store_true",
        help="Enable 4-bit quantization (off by default for max quality)",
    )
    parser.add_argument(
        "--load-8bit",
        action="store_true",
        help="Use 8-bit quantization instead of 4-bit",
    )
    parser.add_argument("--imgsz", type=int, default=1024, help="Image size for inference")
    parser.add_argument("--conf", type=float, default=0.1, help="Confidence threshold")
    parser.add_argument(
        "--no-ocr",
        action="store_true",
        help="Disable OCR (only detection)",
    )
    parser.add_argument(
        "--max-pages",
        type=int,
        default=None,
        help="Maximum number of pages to process (for testing)",
    )

    args = parser.parse_args()

    process_pdf_to_docx(
        pdf_path=args.pdf_path,
        output_docx=args.output,
        model_path=args.model,
        ocr_model_path=args.ocr_model,
        imgsz=args.imgsz,
        conf=args.conf,
        enable_ocr=not args.no_ocr,
        load_4bit=args.use_4bit,
        load_8bit=args.load_8bit,
        max_pages=args.max_pages,
    )